#!/usr/bin/env python
"""
基于Unstructured库的文档解析和生成服务
"""
import os
import re
import logging
from typing import Dict, List, Any, Optional
from unstructured.partition.docx import partition_docx
from unstructured.partition.doc import partition_doc
from unstructured.partition.pdf import partition_pdf
from unstructured.partition.html import partition_html
from unstructured.chunking.title import chunk_by_title
from unstructured.documents.elements import Title, NarrativeText, ListItem, Table, Header, Footer

logger = logging.getLogger(__name__)


class UnstructuredDocumentService:
    """基于Unstructured库的文档解析和生成服务"""
    
    def __init__(self):
        """初始化服务"""
        self.supported_formats = ['.docx', '.doc', '.pdf', '.html', '.txt']
    
    def parse_template_file(self, filepath: str) -> str:
        """
        解析模板文件，提取内容和基本格式信息
        
        Args:
            filepath: 模板文件路径
            
        Returns:
            解析后的文档内容（带基本格式标签）
        """
        try:
            if not os.path.exists(filepath):
                logger.error(f"模板文件不存在: {filepath}")
                return ""
            
            # 根据文件扩展名选择解析方法
            if not filepath:
                logger.error("文件路径为空")
                return ""
            file_ext = os.path.splitext(filepath)[1].lower()
            
            # 简化处理，优先使用文本读取
            if file_ext == '.txt':
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                return self._format_text_content(content)
            
            # 对于Word文档，尝试使用Unstructured，如果失败则回退到简单文本提取
            try:
                if file_ext == '.docx':
                    elements = partition_docx(filepath)
                elif file_ext == '.doc':
                    elements = partition_doc(filepath)
                elif file_ext == '.pdf':
                    elements = partition_pdf(filepath)
                elif file_ext == '.html':
                    elements = partition_html(filepath)
                else:
                    logger.warning(f"不支持的文件格式: {file_ext}")
                    return self._fallback_text_extraction(filepath)
                
                # 将元素转换为带格式的文本
                formatted_content = self._elements_to_formatted_text(elements)
                if formatted_content and formatted_content.strip():
                    logger.info(f"成功解析模板文件: {filepath}")
                    return formatted_content
                else:
                    logger.warning(f"Unstructured解析结果为空，使用备用方法")
                    return self._fallback_text_extraction(filepath)
                
            except Exception as e:
                logger.warning(f"Unstructured解析失败，使用备用方法: {e}")
                # 回退到简单的文本提取
                return self._fallback_text_extraction(filepath)
            
        except Exception as e:
            logger.error(f"解析模板文件失败: {e}")
            return ""
    
    def _fallback_text_extraction(self, filepath: str) -> str:
        """
        备用文本提取方法
        
        Args:
            filepath: 文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            if not filepath:
                logger.error("文件路径为空")
                return ""
            file_ext = os.path.splitext(filepath)[1].lower()
            
            if file_ext == '.docx':
                # 使用python-docx库
                from docx import Document
                doc = Document(filepath)
                content_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        content_parts.append(para.text.strip())
                return self._format_text_content('\n'.join(content_parts))
            
            elif file_ext == '.doc':
                # 对于.doc文件，使用多种方法尝试解析
                content = self._parse_doc_file_robust(filepath)
                if content:
                    return self._format_text_content(content)
                return ""
            
            else:
                # 尝试直接读取为文本
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return self._format_text_content(content)
                
        except Exception as e:
            logger.error(f"备用文本提取失败: {e}")
            return ""
    
    def _elements_to_formatted_text(self, elements: List[Any]) -> str:
        """
        将Unstructured元素转换为带格式的文本
        
        Args:
            elements: Unstructured解析的元素列表
            
        Returns:
            格式化的文本内容
        """
        formatted_parts = []
        
        for element in elements:
            try:
                element_type = element.category if hasattr(element, 'category') else 'Unknown'
                text = element.text if hasattr(element, 'text') else str(element)
                
                if not text or not text.strip():
                    continue
                
                # 根据元素类型添加格式标签
                if element_type == 'Title':
                    # 标题 - 添加加粗和居中标签
                    formatted_text = f"**{text.strip()}**"
                    formatted_parts.append(f"<div style='text-align: center;'>{formatted_text}</div>")
                    
                elif element_type == 'Header':
                    # 页眉
                    formatted_text = f"**{text.strip()}**"
                    formatted_parts.append(f"<div style='text-align: center;'>{formatted_text}</div>")
                    
                elif element_type == 'Footer':
                    # 页脚
                    formatted_text = f"<div style='text-align: right;'>{text.strip()}</div>"
                    formatted_parts.append(formatted_text)
                    
                elif element_type == 'NarrativeText':
                    # 段落文本
                    formatted_text = self._format_paragraph_text(text.strip())
                    formatted_parts.append(formatted_text)
                    
                elif element_type == 'ListItem':
                    # 列表项
                    formatted_text = f"• {text.strip()}"
                    formatted_parts.append(formatted_text)
                    
                elif element_type == 'Table':
                    # 表格 - 简化处理
                    formatted_text = f"<table>{text.strip()}</table>"
                    formatted_parts.append(formatted_text)
                    
                else:
                    # 其他类型，直接添加
                    formatted_parts.append(text.strip())
                    
            except Exception as e:
                logger.warning(f"处理元素时出错: {e}")
                continue
        
        return '\n'.join(formatted_parts)
    
    def _format_paragraph_text(self, text: str) -> str:
        """
        格式化段落文本，添加基本的缩进和格式
        
        Args:
            text: 原始文本
            
        Returns:
            格式化后的文本
        """
        # 检测是否为特殊段落（如诉讼请求、事实与理由等）
        if any(keyword in text for keyword in ['诉讼请求', '事实与理由', '请求事项', '此致', '具状人', '申请人']):
            return f"<div style='padding-left: 28px;'><strong>{text}</strong></div>"
        
        # 检测是否为普通段落
        if len(text) > 10:  # 长段落
            return f"<div style='padding-left: 28px;'>{text}</div>"
        else:  # 短段落
            return f"<div style='padding-left: 28px;'>{text}</div>"
    
    def _parse_doc_file_robust(self, filepath: str) -> str:
        """
        健壮的.doc文件解析方法
        
        Args:
            filepath: 文件路径
            
        Returns:
            解析的文本内容
        """
        # 方法1: 尝试使用docx2txt
        try:
            import docx2txt
            content = docx2txt.process(filepath)
            if content and content.strip():
                logger.info("使用docx2txt成功解析.doc文件")
                return content
        except Exception as e:
            logger.warning(f"docx2txt解析失败: {e}")
        
        # 方法2: 尝试使用python-docx（将.doc当作.docx处理）
        try:
            from docx import Document
            doc = Document(filepath)
            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text.strip())
            if content_parts:
                logger.info("使用python-docx成功解析.doc文件")
                return '\n'.join(content_parts)
        except Exception as e:
            logger.warning(f"python-docx解析失败: {e}")
        
        # 方法3: 尝试使用win32com（仅Windows）
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            doc = word.Documents.Open(filepath, ReadOnly=True)
            content = doc.Content.Text
            doc.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            
            if content and content.strip():
                logger.info("使用win32com成功解析.doc文件")
                return content
        except Exception as e:
            logger.warning(f"win32com解析失败: {e}")
        
        # 方法4: 尝试使用antiword（如果可用）
        try:
            import subprocess
            result = subprocess.run(['antiword', filepath], capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                logger.info("使用antiword成功解析.doc文件")
                return result.stdout
        except Exception as e:
            logger.warning(f"antiword解析失败: {e}")
        
        # 方法5: 二进制读取并过滤文本
        try:
            with open(filepath, 'rb') as f:
                content = f.read()
            
            # 提取可打印字符
            text_content = ""
            for byte in content:
                if 32 <= byte <= 126 or byte in [9, 10, 13]:  # 可打印字符和制表符、换行符
                    text_content += chr(byte)
                elif byte == 0:
                    text_content += " "  # 空字节替换为空格
            
            # 清理文本
            lines = text_content.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line and len(line) > 3:  # 过滤掉太短的行
                    cleaned_lines.append(line)
            
            if cleaned_lines:
                logger.info("使用二进制读取成功解析.doc文件")
                return '\n'.join(cleaned_lines)
        except Exception as e:
            logger.warning(f"二进制读取解析失败: {e}")
        
        # 方法6: 尝试不同的编码读取
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                if content and content.strip():
                    logger.info(f"使用{encoding}编码成功解析.doc文件")
                    return content
            except Exception as e:
                logger.warning(f"使用{encoding}编码解析失败: {e}")
        
        logger.error("所有.doc文件解析方法都失败了")
        return ""
    
    def _format_text_content(self, content: str) -> str:
        """
        格式化纯文本内容
        
        Args:
            content: 原始文本内容
            
        Returns:
            格式化后的文本
        """
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                formatted_lines.append('<br>')
                continue
            
            # 检测标题（通常较短且包含特定关键词）
            if len(line) < 50 and any(keyword in line for keyword in ['起诉状', '申请书', '答辩状', '代理词']):
                formatted_lines.append(f"<div style='text-align: center;'><strong>{line}</strong></div>")
            else:
                formatted_lines.append(f"<div style='padding-left: 28px;'>{line}</div>")
        
        return '\n'.join(formatted_lines)
    
    def load_template_files(self) -> List[Dict[str, str]]:
        """
        加载模板文件列表
        
        Returns:
            模板文件信息列表
        """
        templates = []
        
        try:
            from .models import DocumentTemplate
            template_records = DocumentTemplate.objects.filter(is_active=True, is_deleted=False)
            
            if not template_records.exists():
                logger.warning("数据库中没有找到启用的模板记录")
                return templates
            
            for template_record in template_records:
                filepath = template_record.full_file_path
                if not filepath:
                    logger.warning(f"模板文件路径为空: {template_record.template_name}")
                    continue
                    
                filename = os.path.basename(filepath)
                
                if not os.path.exists(filepath):
                    logger.warning(f"模板文件不存在: {filepath}")
                    continue
                
                # 解析模板内容
                content = self.parse_template_file(filepath)
                if not content:
                    logger.warning(f"解析模板文件失败: {filename}")
                    continue
                
                templates.append({
                    "name": filename,
                    "content": content,
                    "is_binary": False,
                    "file_type": self._get_file_type(filename),
                    "template_type": template_record.template_type,
                    "template_id": template_record.id
                })
            
            logger.info(f"从数据库成功加载 {len(templates)} 个模板文件")
            return templates
            
        except Exception as e:
            logger.error(f"加载模板文件失败: {e}")
            return templates
    
    def _get_file_type(self, filename: str) -> str:
        """
        根据文件名获取文件类型
        
        Args:
            filename: 文件名
            
        Returns:
            文件类型
        """
        if not filename:
            return 'unknown'
            
        ext = os.path.splitext(filename)[1].lower()
        type_map = {
            '.txt': 'text',
            '.doc': 'word_old',
            '.docx': 'word_new',
            '.pdf': 'pdf',
            '.html': 'html'
        }
        return type_map.get(ext, 'unknown')
    
    def extract_document_type_from_filename(self, filename: str) -> str:
        """
        从文件名提取文档类型
        
        Args:
            filename: 文件名
            
        Returns:
            文档类型
        """
        if not filename:
            return "未知文档"
            
        name_without_ext = os.path.splitext(filename)[0]
        
        # 匹配 "数字、文档类型（其他内容）" 格式
        match = re.match(r'^\d+、(.+?)(?:（.*?）)?$', name_without_ext)
        if match:
            doc_type = match.group(1).strip()
            return doc_type[:20] if len(doc_type) > 20 else doc_type
        
        # 匹配 "数字.文档类型" 格式
        match = re.match(r'^\d+\.(.+?)$', name_without_ext)
        if match:
            doc_type = match.group(1).strip()
            return doc_type[:20] if len(doc_type) > 20 else doc_type
        
        # 匹配 "数字-文档类型" 格式
        match = re.match(r'^\d+-(.+?)$', name_without_ext)
        if match:
            doc_type = match.group(1).strip()
            return doc_type[:20] if len(doc_type) > 20 else doc_type
        
        # 匹配 "数字_文档类型" 格式
        match = re.match(r'^\d+_(.+?)$', name_without_ext)
        if match:
            doc_type = match.group(1).strip()
            return doc_type[:20] if len(doc_type) > 20 else doc_type
        
        # 如果没有匹配到特定格式，直接返回原文件名（去掉扩展名）
        return name_without_ext[:20] if len(name_without_ext) > 20 else name_without_ext
    
    def generate_document_with_ai(self, case_data: Dict[str, Any], template_content: str, document_type: str) -> str:
        """
        使用AI生成文档内容
        
        Args:
            case_data: 案例数据
            template_content: 模板内容
            document_type: 文档类型
            
        Returns:
            生成的文档内容
        """
        try:
            from .direct_langchain_ai_service import get_chat_model
            
            # 构建提示词
            system_prompt = f"""
你是一个专业的法律文书生成助手。请根据提供的案例信息和模板内容，生成符合法律要求的{document_type}。

案例信息：
- 原告：{case_data.get('plaintiff_name', '待填写')}
- 原告住所地：{case_data.get('plaintiff_address', '待填写')}
- 原告统一社会信用代码：{case_data.get('plaintiff_credit_code', '待填写')}
- 原告法定代表人：{case_data.get('plaintiff_legal_representative', '待填写')}
- 被告：{case_data.get('defendant_name', '待填写')}
- 被告住所地：{case_data.get('defendant_address', '待填写')}
- 被告统一社会信用代码：{case_data.get('defendant_credit_code', '待填写')}
- 被告法定代表人：{case_data.get('defendant_legal_representative', '待填写')}
- 合同金额：{case_data.get('contract_amount', '待填写')}
- 律师费：{case_data.get('lawyer_fee', '待填写')}
- 案件描述：{case_data.get('case_description', '待填写')}

模板内容：
{template_content}

请严格按照模板的格式和结构生成文档，保持HTML标签和样式不变，只替换其中的示例数据为实际的案例信息。
"""
            
            # 获取AI模型
            model = get_chat_model()
            
            # 生成内容
            response = model.invoke(system_prompt)
            
            if hasattr(response, 'content'):
                return response.content
            else:
                return str(response)
                
        except Exception as e:
            logger.error(f"AI生成文档失败: {e}")
            return template_content  # 返回原始模板内容作为备用
    
    def generate_all_documents(self, case_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        生成所有文档
        
        Args:
            case_data: 案例数据
            
        Returns:
            生成结果
        """
        try:
            # 加载模板文件
            templates = self.load_template_files()
            
            if not templates:
                return {
                    'success': False,
                    'message': '没有找到可用的模板文件',
                    'documents': []
                }
            
            generated_documents = []
            success_count = 0
            
            for template in templates:
                try:
                    # 生成文档内容
                    document_content = self.generate_document_with_ai(
                        case_data, 
                        template['content'], 
                        self.extract_document_type_from_filename(template['name'])
                    )
                    
                    # 保存文档到数据库
                    from .models import CaseDocument
                    from .models import CaseManagement
                    
                    case_id = case_data.get('case_id')
                    if not case_id:
                        logger.error("案例ID为空")
                        continue
                    
                    try:
                        case = CaseManagement.objects.get(id=case_id)
                        logger.info(f"找到案例: {case.case_name}")
                    except CaseManagement.DoesNotExist:
                        logger.error(f"案例不存在: {case_id}")
                        continue
                    
                    # 获取文档类型，确保在choices范围内
                    doc_type = self.extract_document_type_from_filename(template['name'])
                    if doc_type not in ['word', 'pdf', 'excel']:
                        doc_type = 'word'  # 默认为word类型
                    
                    document = CaseDocument.objects.create(
                        case=case,
                        document_name=template['name'],
                        document_type=doc_type,
                        document_content=document_content,
                        generation_method='manual',  # 使用manual而不是unstructured_ai
                        template_used=template['name'],
                        file_path='',  # 添加file_path字段
                        file_size=0   # 添加file_size字段
                    )
                    
                    generated_documents.append({
                        'document_name': template['name'],
                        'template_name': template['name'],
                        'content': document_content,
                        'success': True
                    })
                    
                    success_count += 1
                    logger.info(f"成功生成文档: {template['name']}")
                    
                except Exception as e:
                    logger.error(f"生成文档失败 {template['name']}: {e}")
                    continue
            
            return {
                'success': True,
                'documents': generated_documents,
                'total_count': len(templates),
                'success_count': success_count,
                'error_count': len(templates) - success_count,
                'error': None
            }
            
        except Exception as e:
            logger.error(f"生成所有文档失败: {e}")
            return {
                'success': False,
                'documents': [],
                'total_count': 0,
                'success_count': 0,
                'error_count': 0,
                'error': str(e)
            }


# 创建全局实例
unstructured_service = UnstructuredDocumentService()
