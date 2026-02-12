"""
文档转换器模块 - 将Word文档转换为带格式的Markdown/XML
"""

import os
import re
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xml.etree.ElementTree as ET
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentConverter:
    """文档转换器 - 将Word文档转换为带格式的Markdown/XML"""
    
    def __init__(self, template_path: str):
        """
        初始化文档转换器
        
        Args:
            template_path: Word模板文件路径
        """
        self.template_path = template_path
        self.doc = Document(template_path)
    
    def to_markdown(self) -> str:
        """
        将Word文档转换为带格式的Markdown
        
        Returns:
            转换后的Markdown内容
        """
        markdown_content = []
        
        # 处理段落
        for para in self.doc.paragraphs:
            if para.text.strip():
                # 根据样式添加Markdown格式
                style_name = para.style.name if para.style else 'Normal'
                
                if 'Heading' in style_name:
                    level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
                    markdown_content.append(f"{'#' * level} {para.text}")
                elif 'Title' in style_name:
                    markdown_content.append(f"# {para.text}")
                elif 'Subtitle' in style_name:
                    markdown_content.append(f"## {para.text}")
                else:
                    # 处理粗体、斜体等格式
                    formatted_text = self._format_text(para)
                    markdown_content.append(formatted_text)
                
                markdown_content.append("")  # 空行分隔
        
        # 处理表格
        for table in self.doc.tables:
            markdown_content.append(self._table_to_markdown(table))
            markdown_content.append("")
        
        return "\n".join(markdown_content)
    
    def to_xml(self) -> str:
        """
        将Word文档转换为XML格式
        
        Returns:
            转换后的XML内容
        """
        xml_content = ["<document>"]
        
        # 处理段落
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.strip():
                style_name = para.style.name if para.style else 'Normal'
                formatted_text = self._format_text(para)
                
                xml_content.append(f'  <paragraph index="{i}" style="{style_name}">')
                xml_content.append(f"    {formatted_text}")
                xml_content.append("  </paragraph>")
        
        # 处理表格
        for i, table in enumerate(self.doc.tables):
            xml_content.append(f'  <table index="{i}">')
            for row_idx, row in enumerate(table.rows):
                xml_content.append(f'    <row index="{row_idx}">')
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = " ".join([p.text for p in cell.paragraphs])
                    xml_content.append(f'      <cell index="{cell_idx}">{cell_text}</cell>')
                xml_content.append("    </row>")
            xml_content.append("  </table>")
        
        xml_content.append("</document>")
        return "\n".join(xml_content)
    
    def to_structured_markdown(self) -> str:
        """
        将Word文档转换为结构化Markdown（保留更多格式信息）
        
        Returns:
            转换后的结构化Markdown内容
        """
        markdown_content = []
        
        # 处理段落，保留完整格式信息
        for paragraph in self.doc.paragraphs:
            if not paragraph.text.strip():
                markdown_content.append("")  # 保留空行
                continue
                
            # 检查段落样式
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # 构建格式化的段落内容
            formatted_text = paragraph.text.strip()
            
            # 检查段落整体格式
            is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
            is_italic = any(run.italic for run in paragraph.runs if run.text.strip())
            is_underline = any(run.underline for run in paragraph.runs if run.text.strip())
            
            # 提取字体信息
            font_name = None
            font_size = None
            if paragraph.runs:
                # 使用第一个run的字体信息作为段落代表
                first_run = next((run for run in paragraph.runs if run.text.strip()), None)
                if first_run:
                    font_name = first_run.font.name
                    font_size = first_run.font.size
            
            # 提取段落格式信息
            para_format = paragraph.paragraph_format
            left_indent = para_format.left_indent
            first_line_indent = para_format.first_line_indent
            space_before = para_format.space_before
            space_after = para_format.space_after
            line_spacing = para_format.line_spacing
            alignment = paragraph.alignment
            
            # 构建完整的格式标记
            format_tags = []
            
            # 字体大小标记
            if font_size:
                size_pt = font_size.pt if hasattr(font_size, 'pt') else font_size / 10000
                format_tags.append(f"<size:{int(size_pt)}>")
            
            # 字体名称标记
            if font_name and font_name != 'Normal':
                format_tags.append(f"<font:{font_name}>")
            
            # 段落缩进标记
            if first_line_indent and first_line_indent > 0:
                indent_pt = first_line_indent.pt if hasattr(first_line_indent, 'pt') else first_line_indent / 10000
                format_tags.append(f"<indent:{int(indent_pt)}>")
            
            # 段落间距标记
            if space_before and space_before > 0:
                before_pt = space_before.pt if hasattr(space_before, 'pt') else space_before / 10000
                format_tags.append(f"<space_before:{int(before_pt)}>")
            
            if space_after and space_after > 0:
                after_pt = space_after.pt if hasattr(space_after, 'pt') else space_after / 10000
                format_tags.append(f"<space_after:{int(after_pt)}>")
            
            # 行距标记
            if line_spacing and line_spacing != 1.0:
                if hasattr(line_spacing, 'pt'):
                    spacing_value = line_spacing.pt / 12
                else:
                    spacing_value = line_spacing / 10000 / 12
                spacing_value = min(max(spacing_value, 0.5), 3.0)
                format_tags.append(f"<line_spacing:{spacing_value}>")
            
            # 对齐方式标记
            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                format_tags.append("<align:center>")
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                format_tags.append("<align:right>")
            elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                format_tags.append("<align:justify>")
            
            # 应用格式标记
            if format_tags:
                formatted_text = ''.join(format_tags) + formatted_text
            
            # 根据格式添加标记
            if is_bold:
                formatted_text = f"**{formatted_text}**"
            if is_italic:
                formatted_text = f"*{formatted_text}*"
            if is_underline:
                formatted_text = f"<u>{formatted_text}</u>"
            
            # 根据段落样式添加标记
            if "Heading" in style_name or "Title" in style_name:
                level = 1
                if "Heading" in style_name:
                    try:
                        level = int(style_name.split()[-1])
                    except:
                        level = 1
                formatted_text = f"{'#' * level} {formatted_text}"
            elif style_name == "List Paragraph":
                formatted_text = f"- {formatted_text}"
            elif left_indent and left_indent > 0:
                indent_level = int(left_indent / Inches(0.5))
                formatted_text = "  " * indent_level + formatted_text
            
            markdown_content.append(formatted_text)
        
        # 提取表格内容，保留格式
        for table in self.doc.tables:
            markdown_content.append("")  # 表格前空行
            markdown_content.append("| " + " | ".join(["列" + str(i+1) for i in range(len(table.columns))]) + " |")
            markdown_content.append("| " + " | ".join(["---" for _ in range(len(table.columns))]) + " |")
            
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_content = ""
                    for run in cell.paragraphs[0].runs:
                        text = run.text
                        if run.bold:
                            text = f"**{text}**"
                        if run.italic:
                            text = f"*{text}*"
                        cell_content += text
                    row_text.append(cell_content.strip())
                if any(row_text):
                    markdown_content.append("| " + " | ".join(row_text) + " |")
            markdown_content.append("")  # 表格后空行
        
        return "\n".join(markdown_content)
    
    def _format_text(self, para) -> str:
        """
        格式化段落文本，保留粗体、斜体等格式
        
        Args:
            para: Word段落对象
            
        Returns:
            格式化后的文本
        """
        text = para.text
        formatted_text = text
        
        # 处理粗体
        for run in para.runs:
            if run.bold:
                formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
            if run.italic:
                formatted_text = formatted_text.replace(run.text, f"*{run.text}*")
            if run.underline:
                formatted_text = formatted_text.replace(run.text, f"<u>{run.text}</u>")
        
        return formatted_text
    
    def _table_to_markdown(self, table) -> str:
        """
        将表格转换为Markdown格式
        
        Args:
            table: Word表格对象
            
        Returns:
            转换后的Markdown表格
        """
        if not table.rows:
            return ""
        
        markdown_table = []
        
        # 表头
        header_row = []
        for cell in table.rows[0].cells:
            header_row.append(cell.text.strip())
        markdown_table.append("| " + " | ".join(header_row) + " |")
        
        # 分隔线
        separator = "| " + " | ".join(["---"] * len(header_row)) + " |"
        markdown_table.append(separator)
        
        # 数据行
        for row in table.rows[1:]:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            markdown_table.append("| " + " | ".join(row_data) + " |")
        
        return "\n".join(markdown_table)
    
    def get_document_structure(self) -> Dict:
        """
        获取文档结构信息
        
        Returns:
            文档结构信息字典
        """
        structure = {
            "paragraphs": [],
            "tables": [],
            "total_elements": 0
        }
        
        # 段落信息
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.strip():
                structure["paragraphs"].append({
                    "index": i,
                    "text": para.text,
                    "style": para.style.name if para.style else 'Normal',
                    "length": len(para.text)
                })
        
        # 表格信息
        for i, table in enumerate(self.doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_info["data"].append(row_data)
            
            structure["tables"].append(table_info)
        
        structure["total_elements"] = len(structure["paragraphs"]) + len(structure["tables"])
        return structure

def convert_docx_to_markdown(file_path: str) -> str:
    """
    将docx文件转换为Markdown格式的便捷函数
    
    Args:
        file_path: docx文件路径
        
    Returns:
        转换后的Markdown内容
    """
    try:
        converter = DocumentConverter(file_path)
        return converter.to_structured_markdown()
    except Exception as e:
        logger.error(f"转换docx文件失败: {e}")
        return ""

def convert_docx_to_xml(file_path: str) -> str:
    """
    将docx文件转换为XML格式的便捷函数
    
    Args:
        file_path: docx文件路径
        
    Returns:
        转换后的XML内容
    """
    try:
        converter = DocumentConverter(file_path)
        return converter.to_xml()
    except Exception as e:
        logger.error(f"转换docx文件失败: {e}")
        return ""

if __name__ == "__main__":
    # 测试文档转换器
    template_path = "backend/template/起诉状.docx"
    if os.path.exists(template_path):
        converter = DocumentConverter(template_path)
        
        # 转换为Markdown
        markdown_content = converter.to_structured_markdown()
        print("Markdown内容:")
        print(markdown_content[:500] + "..." if len(markdown_content) > 500 else markdown_content)
        
        # 转换为XML
        xml_content = converter.to_xml()
        print("\nXML内容:")
        print(xml_content[:500] + "..." if len(xml_content) > 500 else xml_content)
        
        # 获取文档结构
        structure = converter.get_document_structure()
        print(f"\n文档结构: {structure['total_elements']} 个元素")
    else:
        print(f"模板文件不存在: {template_path}")
