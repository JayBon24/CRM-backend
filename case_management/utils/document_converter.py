"""
文档转换服务模块
实现 DOCX 与 HTML 格式互转功能
"""
import os
import re
import uuid
import logging
from typing import Dict, List, Optional
from io import BytesIO
from urllib.parse import urlparse
import base64

import mammoth
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import bleach
# from html2docx import HTML2Docx  # 该库API可能不同，使用python-docx直接实现
import requests
from PIL import Image

logger = logging.getLogger(__name__)


class DocumentConverter:
    """文档转换器"""
    
    # 允许的HTML标签
    ALLOWED_TAGS = [
        'p', 'br', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'b', 'strong', 'i', 'em', 'u', 's', 'strike',
        'ul', 'ol', 'li',
        'table', 'tr', 'td', 'th', 'thead', 'tbody',
        'img', 'a',
        'span', 'div',
        'blockquote', 'pre', 'code'
    ]
    
    # 允许的HTML属性
    ALLOWED_ATTRIBUTES = {
        '*': ['style', 'class'],
        'img': ['src', 'alt', 'width', 'height', 'title'],
        'a': ['href', 'title', 'target'],
        'table': ['border', 'cellpadding', 'cellspacing', 'style'],
        'td': ['colspan', 'rowspan', 'style'],
        'th': ['colspan', 'rowspan', 'style'],
    }
    
    # 允许的CSS样式
    ALLOWED_STYLES = [
        'color', 'background-color',
        'font-family', 'font-size', 'font-weight', 'font-style',
        'text-align', 'text-decoration', 'text-indent',
        'line-height', 'margin', 'padding', 'border'
    ]
    
    def __init__(self):
        """初始化转换器"""
        self.temp_dir = None
    
    def docx_to_html(self, docx_path: str, image_output_dir: Optional[str] = None) -> Dict:
        """
        DOCX转HTML（完全使用python-docx解析，高保真格式保留）
        
        Args:
            docx_path: DOCX文件路径
            image_output_dir: 图片输出目录（可选）
        
        Returns:
            {
                'html': str,
                'images': list,
                'title': str
            }
        """
        try:
            # 使用python-docx直接解析文档
            doc = Document(docx_path)
            
            # 提取图片
            images = []
            if image_output_dir:
                images = self._extract_images_from_docx_file(doc, image_output_dir)
            
            # 直接使用python-docx解析并生成HTML（完全保留格式）
            html = self._docx_to_html_with_full_formatting(doc, image_output_dir)
            
            # 获取文档标题（从文件名）
            title = os.path.splitext(os.path.basename(docx_path))[0]
            
            return {
                'html': html,
                'images': images,
                'title': title
            }
            
        except Exception as e:
            logger.error(f"DOCX转HTML失败: {e}", exc_info=True)
            raise
    
    def html_to_docx(self, html_content: str, output_path: str, 
                     image_base_dir: Optional[str] = None) -> str:
        """
        HTML转DOCX
        
        Args:
            html_content: HTML内容字符串
            output_path: 输出DOCX文件路径
            image_base_dir: 图片基础目录（用于相对路径图片）
        
        Returns:
            文件路径
        """
        try:
            # 清理和校验HTML
            html_content = self.clean_html(html_content)
            
            # 创建Document对象
            doc = Document()
            
            # 设置默认样式
            self._apply_default_styles(doc)
            
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 处理HTML内容并转换为DOCX
            self._convert_html_elements_to_docx(soup, doc, image_base_dir)
            
            # 保存文档
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"HTML转DOCX失败: {e}", exc_info=True)
            raise
    
    def clean_html(self, html: str) -> str:
        """
        清理HTML，移除不安全的标签和脚本
        
        Args:
            html: HTML字符串
        
        Returns:
            清理后的HTML
        """
        # 使用bleach清理HTML（兼容bleach 4.0+版本，不支持styles参数）
        try:
            # 尝试使用新版本API（bleach 4.0+）
            cleaned = bleach.clean(
                html,
                tags=self.ALLOWED_TAGS,
                attributes=self.ALLOWED_ATTRIBUTES,
                strip=True
            )
        except TypeError:
            # 兼容旧版本（bleach < 4.0）
            cleaned = bleach.clean(
                html,
                tags=self.ALLOWED_TAGS,
                attributes=self.ALLOWED_ATTRIBUTES,
                styles=self.ALLOWED_STYLES,
                strip=True
            )
        
        # 使用BeautifulSoup进一步处理
        soup = BeautifulSoup(cleaned, 'html.parser')
        
        # 移除script和style标签
        for tag in soup.find_all(['script', 'style']):
            tag.decompose()
        
        # 移除危险的属性
        for tag in soup.find_all(True):
            for attr in list(tag.attrs.keys()):
                if attr.startswith('on'):  # 移除onclick等事件处理器
                    del tag.attrs[attr]
        
        # 清理style属性中的不安全样式（手动过滤）
        # 但保留所有格式相关的样式（text-align, font-size等）
        for tag in soup.find_all(True):
            if 'style' in tag.attrs:
                style_value = tag.attrs['style']
                # 清理危险内容，但保留所有格式样式
                if 'expression' in style_value.lower() or 'javascript:' in style_value.lower():
                    # 如果包含危险内容，进行过滤
                    safe_styles = []
                    for style in style_value.split(';'):
                        style = style.strip()
                        if style and 'expression' not in style.lower() and 'javascript:' not in style.lower():
                            safe_styles.append(style)
                    if safe_styles:
                        tag.attrs['style'] = '; '.join(safe_styles)
                    else:
                        del tag.attrs['style']
                # 如果没有危险内容，保留所有样式（包括text-align等）
                # 不进行过滤，直接保留
        
        return str(soup)
    
    def _enhance_html_with_docx_formatting(self, html: str, doc: Document, docx_path: str) -> str:
        """
        使用python-docx获取的格式信息增强HTML
        
        Args:
            html: mammoth转换的HTML
            doc: python-docx Document对象
            docx_path: 文档路径
        
        Returns:
            增强后的HTML
        """
        soup = BeautifulSoup(html, 'html.parser')
        
        # 获取所有段落元素（HTML）
        html_paragraphs = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        
        # 获取Word文档的所有段落
        docx_paragraphs = doc.paragraphs
        
        # 匹配段落并应用格式（处理数量不一致的情况）
        min_para_count = min(len(html_paragraphs), len(docx_paragraphs))
        for i in range(min_para_count):
            html_p = html_paragraphs[i]
            docx_p = docx_paragraphs[i]
            styles = []
            
            # 段落对齐方式
            if docx_p.paragraph_format.alignment:
                if docx_p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    styles.append('text-align: center;')
                elif docx_p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    styles.append('text-align: right;')
                elif docx_p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    styles.append('text-align: justify;')
            
            # 行间距
            if docx_p.paragraph_format.line_spacing:
                line_spacing = docx_p.paragraph_format.line_spacing
                if hasattr(line_spacing, 'pt'):
                    styles.append(f'line-height: {line_spacing.pt / 12}em;')
                elif isinstance(line_spacing, (float, int)):
                    styles.append(f'line-height: {float(line_spacing)};')
            
            # 段落间距
            if docx_p.paragraph_format.space_after:
                space_after = docx_p.paragraph_format.space_after
                if hasattr(space_after, 'pt'):
                    styles.append(f'margin-bottom: {space_after.pt}pt;')
            if docx_p.paragraph_format.space_before:
                space_before = docx_p.paragraph_format.space_before
                if hasattr(space_before, 'pt'):
                    styles.append(f'margin-top: {space_before.pt}pt;')
            
            # 缩进
            if docx_p.paragraph_format.first_line_indent:
                indent = docx_p.paragraph_format.first_line_indent
                if hasattr(indent, 'pt'):
                    styles.append(f'text-indent: {indent.pt}pt;')
            if docx_p.paragraph_format.left_indent:
                left_indent = docx_p.paragraph_format.left_indent
                if hasattr(left_indent, 'pt'):
                    styles.append(f'padding-left: {left_indent.pt}pt;')
            
            # 应用样式
            if styles:
                existing_style = html_p.get('style', '')
                if existing_style and not existing_style.endswith(';'):
                    existing_style += ';'
                html_p['style'] = existing_style + ' '.join(styles)
            
            # 处理段落内的文本格式（runs）
            # 获取段落中的所有文本节点和格式元素
            html_text_content = html_p.get_text()
            docx_text_content = ''.join([run.text for run in docx_p.runs])
            
            # 如果文本内容匹配，尝试匹配runs
            if html_text_content.strip() == docx_text_content.strip() and docx_p.runs:
                # 构建文本节点到HTML元素的映射
                html_runs = html_p.find_all(['span', 'strong', 'em', 'u', 'b', 'i'])
                docx_runs = docx_p.runs
                
                # 只处理前min(len(html_runs), len(docx_runs))个runs
                run_count = min(len(html_runs), len(docx_runs))
                for j in range(run_count):
                    html_run_elem = html_runs[j]
                    docx_run = docx_runs[j]
                    run_styles = []
                    
                    # 字体
                    if docx_run.font.name:
                        font_name = docx_run.font.name
                        run_styles.append(f"font-family: '{font_name}';")
                    
                    # 字号
                    if docx_run.font.size:
                        size_pt = docx_run.font.size.pt if hasattr(docx_run.font.size, 'pt') else docx_run.font.size
                        run_styles.append(f'font-size: {size_pt}pt;')
                    
                    # 颜色
                    if docx_run.font.color and docx_run.font.color.rgb:
                        color = docx_run.font.color.rgb
                        if isinstance(color, RGBColor):
                            run_styles.append(f'color: rgb({color.r}, {color.g}, {color.b});')
                    
                    # 加粗
                    if docx_run.bold:
                        if html_run_elem.name not in ['strong', 'b']:
                            run_styles.append('font-weight: bold;')
                    
                    # 斜体
                    if docx_run.italic:
                        if html_run_elem.name not in ['em', 'i']:
                            run_styles.append('font-style: italic;')
                    
                    # 下划线
                    if docx_run.underline:
                        if html_run_elem.name != 'u':
                            run_styles.append('text-decoration: underline;')
                    
                    # 删除线
                    if docx_run.font.strike:
                        run_styles.append('text-decoration: line-through;')
                    
                    # 高亮（背景色）
                    if docx_run.font.highlight_color:
                        highlight = docx_run.font.highlight_color
                        if hasattr(highlight, 'rgb') and highlight.rgb:
                            bg_color = highlight.rgb
                            if isinstance(bg_color, RGBColor):
                                run_styles.append(f'background-color: rgb({bg_color.r}, {bg_color.g}, {bg_color.b});')
                    
                    # 应用样式
                    if run_styles:
                        if html_run_elem.name in ['strong', 'em', 'u', 'b', 'i']:
                            # 在格式标签内添加span
                            span = soup.new_tag('span')
                            span['style'] = ' '.join(run_styles)
                            # 移动内容到span
                            for child in list(html_run_elem.children):
                                span.append(child)
                            html_run_elem.clear()
                            html_run_elem.append(span)
                        else:
                            # 直接在元素上添加样式
                            existing_style = html_run_elem.get('style', '')
                            if existing_style and not existing_style.endswith(';'):
                                existing_style += ';'
                            html_run_elem['style'] = existing_style + ' '.join(run_styles)
        
        # 处理表格格式
        for table in soup.find_all('table'):
            table_styles = ['border-collapse: collapse;', 'width: 100%;']
            existing_style = table.get('style', '')
            if existing_style and not existing_style.endswith(';'):
                existing_style += ';'
            table['style'] = existing_style + ' '.join(table_styles)
            
            # 处理表格单元格
            for td in soup.find_all(['td', 'th']):
                td_styles = ['border: 1px solid #ddd;', 'padding: 8px;']
                existing_style = td.get('style', '')
                if existing_style and not existing_style.endswith(';'):
                    existing_style += ';'
                td['style'] = existing_style + ' '.join(td_styles)
        
        return str(soup)
    
    def _docx_to_html_with_full_formatting(self, doc: Document, image_output_dir: Optional[str] = None) -> str:
        """
        使用python-docx直接解析并生成HTML，完全保留格式
        
        Args:
            doc: python-docx Document对象
            image_output_dir: 图片输出目录
        
        Returns:
            HTML字符串
        """
        html_parts = []
        
        # python-docx的paragraphs和tables是分开的，需要按文档顺序合并
        # 这里简化处理：先处理所有段落，再处理所有表格
        # 如果需要更精确的顺序，需要访问底层element
        
        # 处理文档的每个段落
        for para in doc.paragraphs:
            para_html = self._convert_paragraph_to_html(para, image_output_dir)
            if para_html:
                html_parts.append(para_html)
        
        # 处理表格
        for table in doc.tables:
            table_html = self._convert_table_to_html(table, image_output_dir)
            if table_html:
                html_parts.append(table_html)
        
        # 组合HTML
        html = '\n'.join(html_parts)
        
        # 清理HTML（保留所有格式）
        html = self.clean_html(html)
        
        return html
    
    def _convert_paragraph_to_html(self, para, image_output_dir: Optional[str] = None) -> str:
        """将段落转换为HTML"""
        # 检查是否为空段落（简化检查，避免xpath问题）
        if not para.text.strip():
            # 可以返回<br>来保留空行
            return '<p><br></p>'
        
        # 确定段落标签（根据样式判断是标题还是普通段落）
        tag = 'p'
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if 'heading' in style_name or 'heading 1' in style_name:
                tag = 'h1'
            elif 'heading 2' in style_name:
                tag = 'h2'
            elif 'heading 3' in style_name:
                tag = 'h3'
            elif 'heading 4' in style_name:
                tag = 'h4'
            elif 'heading 5' in style_name:
                tag = 'h5'
            elif 'heading 6' in style_name:
                tag = 'h6'
        
        # 构建段落样式
        para_styles = []
        
        # 对齐方式 - 使用 para.alignment 或 para.paragraph_format.alignment
        # python-docx中，para.alignment 可能更方便
        alignment = None
        # 先尝试 para.alignment
        if hasattr(para, 'alignment') and para.alignment is not None:
            alignment = para.alignment
        # 再尝试 para.paragraph_format.alignment
        elif para.paragraph_format.alignment is not None:
            alignment = para.paragraph_format.alignment
        
        # 应用对齐方式
        if alignment:
            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                para_styles.append('text-align: center;')
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                para_styles.append('text-align: right;')
            elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                para_styles.append('text-align: justify;')
            elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
                para_styles.append('text-align: left;')
        # 如果没有明确的对齐方式，根据样式判断
        elif para.style and para.style.name:
            style_name = para.style.name.lower()
            # 标题样式通常居中（可根据实际情况调整）
            if 'title' in style_name or 'heading' in style_name:
                # 有些标题可能是居中的，这里可以根据实际情况判断
                # 暂时不自动居中，需要明确设置
                pass
        
        # 行间距（更精确的处理）
        if para.paragraph_format.line_spacing:
            line_spacing = para.paragraph_format.line_spacing
            if hasattr(line_spacing, 'pt'):
                # 转换为点值
                pt_value = line_spacing.pt
                # 12pt = 1em，但实际可能需要根据字体大小调整
                # 直接使用pt更精确
                para_styles.append(f'line-height: {pt_value:.2f}pt;')
            elif isinstance(line_spacing, (float, int)):
                # 如果是倍数，转换为倍数形式
                para_styles.append(f'line-height: {float(line_spacing):.2f};')
            elif hasattr(line_spacing, '__class__') and 'LineSpacing' in str(type(line_spacing)):
                # 尝试获取原始值
                try:
                    if hasattr(line_spacing, 'value'):
                        para_styles.append(f'line-height: {line_spacing.value};')
                except:
                    pass
        
        # 段落间距
        if para.paragraph_format.space_after:
            space_after = para.paragraph_format.space_after
            if hasattr(space_after, 'pt'):
                para_styles.append(f'margin-bottom: {space_after.pt}pt;')
        if para.paragraph_format.space_before:
            space_before = para.paragraph_format.space_before
            if hasattr(space_before, 'pt'):
                para_styles.append(f'margin-top: {space_before.pt}pt;')
        
        # 缩进
        if para.paragraph_format.first_line_indent:
            indent = para.paragraph_format.first_line_indent
            if hasattr(indent, 'pt'):
                para_styles.append(f'text-indent: {indent.pt}pt;')
        if para.paragraph_format.left_indent:
            left_indent = para.paragraph_format.left_indent
            if hasattr(left_indent, 'pt'):
                para_styles.append(f'padding-left: {left_indent.pt}pt;')
        if para.paragraph_format.right_indent:
            right_indent = para.paragraph_format.right_indent
            if hasattr(right_indent, 'pt'):
                para_styles.append(f'padding-right: {right_indent.pt}pt;')
        
        # 处理段落内的runs（文本片段）
        inner_html = []
        for run in para.runs:
            run_html = self._convert_run_to_html(run, image_output_dir)
            if run_html:
                inner_html.append(run_html)
        
        # 如果没有内容但有图片，处理图片
        if not inner_html:
            # 检查是否有图片
            if hasattr(para, '_element'):
                # 这里可以提取图片，暂时跳过
                pass
        
        content = ''.join(inner_html) if inner_html else para.text
        
        # 构建段落HTML
        # 处理样式字符串，确保格式正确
        if para_styles:
            style_attr = ' '.join(para_styles)
            # 确保每个样式都有分号
            if not style_attr.endswith(';'):
                style_attr += ';'
            return f'<{tag} style="{style_attr}">{content}</{tag}>'
        else:
            return f'<{tag}>{content}</{tag}>'
    
    def _convert_run_to_html(self, run, image_output_dir: Optional[str] = None) -> str:
        """将文本run转换为HTML"""
        text = run.text
        
        # 构建run的样式
        run_styles = []
        
        # 字体
        if run.font.name:
            font_name = run.font.name
            # 处理中文字体
            if '宋体' in font_name or 'SimSun' in font_name:
                run_styles.append("font-family: '宋体', SimSun, serif;")
            elif '黑体' in font_name or 'SimHei' in font_name:
                run_styles.append("font-family: '黑体', SimHei, sans-serif;")
            elif '微软雅黑' in font_name or 'Microsoft YaHei' in font_name:
                run_styles.append("font-family: '微软雅黑', 'Microsoft YaHei', sans-serif;")
            else:
                run_styles.append(f"font-family: '{font_name}';")
        
        # 字号
        if run.font.size:
            size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
            run_styles.append(f'font-size: {size_pt}pt;')
        
        # 颜色
        if run.font.color and run.font.color.rgb:
            color = run.font.color.rgb
            if isinstance(color, RGBColor):
                run_styles.append(f'color: rgb({color.r}, {color.g}, {color.b});')
        
        # 背景色（高亮）
        if run.font.highlight_color:
            highlight = run.font.highlight_color
            if hasattr(highlight, 'rgb') and highlight.rgb:
                bg_color = highlight.rgb
                if isinstance(bg_color, RGBColor):
                    run_styles.append(f'background-color: rgb({bg_color.r}, {bg_color.g}, {bg_color.b});')
        
        # 确定格式标签
        format_tags = []
        need_wrap = False
        
        if run.bold:
            format_tags.append('strong')
        if run.italic:
            format_tags.append('em')
        if run.underline:
            format_tags.append('u')
        if run.font.strike:
            format_tags.append('s')
        
        # 如果有样式但不需要格式标签，需要span包装
        if run_styles and not format_tags:
            need_wrap = True
        
        # 构建HTML
        inner_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # 应用格式标签
        for tag in reversed(format_tags):  # 从内到外
            inner_text = f'<{tag}>{inner_text}</{tag}>'
        
        # 应用样式（如果有）
        if run_styles:
            style_attr = ' '.join(run_styles)
            if format_tags:
                # 如果已经有格式标签，在外面包一层span
                inner_text = f'<span style="{style_attr}">{inner_text}</span>'
            else:
                # 直接用span
                inner_text = f'<span style="{style_attr}">{inner_text}</span>'
        
        return inner_text
    
    def _convert_table_to_html(self, table, image_output_dir: Optional[str] = None) -> str:
        """将表格转换为HTML"""
        html_parts = ['<table style="border-collapse: collapse; width: 100%;">']
        
        for row in table.rows:
            html_parts.append('<tr>')
            for cell in row.cells:
                # 获取单元格样式
                cell_styles = ['border: 1px solid #000;', 'padding: 5px;']
                
                # 处理单元格内容
                cell_content = []
                for para in cell.paragraphs:
                    para_html = self._convert_paragraph_to_html(para, image_output_dir)
                    cell_content.append(para_html)
                
                content = ''.join(cell_content) if cell_content else cell.text
                if not content.strip():
                    content = '&nbsp;'
                
                style_attr = ' '.join(cell_styles)
                html_parts.append(f'<td style="{style_attr}">{content}</td>')
            
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        return ''.join(html_parts)
    
    def _extract_images_from_docx_file(self, doc: Document, output_dir: str) -> List[Dict]:
        """从python-docx文档中提取图片"""
        images = []
        
        try:
            # 提取文档关系中的图片
            # 注意：python-docx不直接支持图片提取，需要使用底层API
            # 这里简化处理，主要依赖mammoth的图片处理
            
            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)
            
        except Exception as e:
            logger.warning(f"提取图片失败: {e}")
        
        return images
    
    def _get_style_map(self) -> str:
        """获取mammoth样式映射（备用）"""
        return """
        p[style-name='Heading 1'] => h1:fresh
        p[style-name='Heading 2'] => h2:fresh
        p[style-name='Heading 3'] => h3:fresh
        p[style-name='Heading 4'] => h4:fresh
        p[style-name='Heading 5'] => h5:fresh
        p[style-name='Heading 6'] => h6:fresh
        """
    
    def _convert_image_handler(self, image):
        """图片转换处理器"""
        # 返回一个函数，用于处理图片
        with image.open() as image_bytes:
            # 转换为Base64（临时方案）
            image_data = base64.b64encode(image_bytes.read()).decode('utf-8')
            return {
                "src": f"data:image/{image.content_type};base64,{image_data}"
            }
    
    def _extract_images_from_docx(self, docx_bytes: BytesIO, output_dir: str) -> List[Dict]:
        """
        从DOCX中提取图片
        
        Args:
            docx_bytes: DOCX文件字节流
            output_dir: 图片输出目录
        
        Returns:
            图片列表
        """
        images = []
        
        try:
            doc = Document(BytesIO(docx_bytes))
            relationships = doc.part.rels
            
            os.makedirs(output_dir, exist_ok=True)
            
            for rel_id, rel in relationships.items():
                if 'image' in rel.target_ref or (hasattr(rel, 'target_part') and 
                                                 hasattr(rel.target_part, 'content_type') and 
                                                 rel.target_part.content_type.startswith('image/')):
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        content_type = image_part.content_type
                        
                        # 获取图片扩展名
                        ext = content_type.split('/')[-1]
                        if ext == 'jpeg':
                            ext = 'jpg'
                        
                        # 生成唯一文件名
                        image_filename = f"{uuid.uuid4()}.{ext}"
                        image_path = os.path.join(output_dir, image_filename)
                        
                        # 保存图片
                        with open(image_path, 'wb') as f:
                            f.write(image_bytes)
                        
                        images.append({
                            'id': rel_id,
                            'url': f"/media/images/documents/{image_filename}",
                            'path': image_path,
                            'size': len(image_bytes),
                            'content_type': content_type
                        })
                    except Exception as e:
                        logger.error(f"提取图片失败 (rel_id={rel_id}): {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"提取DOCX图片失败: {e}")
        
        return images
    
    def _apply_default_styles(self, document: Document) -> None:
        """应用默认Word样式"""
        try:
            # 设置默认字体
            style = document.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 设置段落格式
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.5
            
        except Exception as e:
            logger.warning(f"应用默认样式失败: {e}")
    
    def _convert_html_elements_to_docx(self, soup: BeautifulSoup, document: Document,
                                      image_base_dir: Optional[str] = None) -> None:
        """将HTML元素转换为DOCX"""
        body = soup.find('body') or soup
        
        for element in body.children:
            if hasattr(element, 'name'):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(element.name[1])
                    para = document.add_heading(element.get_text(), level=level)
                    self._apply_text_formatting(element, para.runs[-1] if para.runs else None)
                
                elif element.name in ['p', 'div']:
                    para = document.add_paragraph()
                    self._process_inline_content(element, para, image_base_dir)
                
                elif element.name in ['ul', 'ol']:
                    is_ordered = element.name == 'ol'
                    for li in element.find_all('li', recursive=False):
                        para = document.add_paragraph(
                            style='List Bullet' if not is_ordered else 'List Number'
                        )
                        self._process_inline_content(li, para, image_base_dir)
                
                elif element.name == 'table':
                    self._process_table(element, document, image_base_dir)
                
                elif element.name == 'br':
                    document.add_paragraph()
                
                elif element.name == 'img':
                    # 处理独立的图片
                    self._add_image_to_docx(element, document, image_base_dir)
    
    def _process_inline_content(self, element, paragraph, image_base_dir: Optional[str] = None):
        """处理段落中的内联内容"""
        for content in element.descendants:
            if isinstance(content, str):
                if content.strip():
                    paragraph.add_run(content)
            elif hasattr(content, 'name'):
                if content.name in ['strong', 'b']:
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name in ['em', 'i']:
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'u':
                    run = paragraph.add_run(content.get_text())
                    run.underline = True
                elif content.name == 'img':
                    self._add_image_to_paragraph(content, paragraph, image_base_dir)
                elif content.name == 'a':
                    run = paragraph.add_run(content.get_text())
                    run.font.color.rgb = RGBColor(0, 0, 255)
    
    def _process_table(self, table_element, document: Document, image_base_dir: Optional[str] = None):
        """处理表格"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        col_count = max(len(row.find_all(['td', 'th'])) for row in rows)
        if col_count == 0:
            return
        
        table = document.add_table(rows=len(rows), cols=col_count)
        table.style = 'Light Grid Accent 1'
        
        for row_idx, html_row in enumerate(rows):
            cells = html_row.find_all(['td', 'th'])
            for col_idx, html_cell in enumerate(cells):
                if col_idx < col_count:
                    word_cell = table.rows[row_idx].cells[col_idx]
                    word_cell.text = ''
                    para = word_cell.paragraphs[0]
                    self._process_inline_content(html_cell, para, image_base_dir)
    
    def _add_image_to_paragraph(self, img_element, paragraph, image_base_dir: Optional[str] = None):
        """在段落中添加图片"""
        src = img_element.get('src', '')
        if not src:
            return
        
        try:
            image_stream = self._get_image_stream(src, image_base_dir)
            if image_stream:
                run = paragraph.add_run()
                width = Inches(4)
                if img_element.get('width'):
                    try:
                        width = Inches(float(img_element.get('width')) / 96)  # 假设96 DPI
                    except:
                        pass
                run.add_picture(image_stream, width=width)
        except Exception as e:
            logger.error(f"添加图片到段落失败: {e}")
    
    def _add_image_to_docx(self, img_element, document: Document, image_base_dir: Optional[str] = None):
        """在文档中添加图片（独立段落）"""
        para = document.add_paragraph()
        self._add_image_to_paragraph(img_element, para, image_base_dir)
    
    def _apply_text_formatting(self, element, run):
        """应用文本格式"""
        if not run:
            return
        
        style = element.get('style', '')
        if 'font-weight: bold' in style or 'font-weight:700' in style:
            run.bold = True
        if 'font-style: italic' in style:
            run.italic = True
    
    def _get_image_stream(self, src: str, image_base_dir: Optional[str] = None) -> Optional[BytesIO]:
        """获取图片流"""
        # Base64图片
        if src.startswith('data:image'):
            image_data = re.match(r'data:image/[^;]+;base64,(.+)', src)
            if image_data:
                try:
                    image_bytes = base64.b64decode(image_data.group(1))
                    return BytesIO(image_bytes)
                except Exception as e:
                    logger.error(f"处理Base64图片失败: {e}")
                    return None
        
        # 网络图片
        elif src.startswith('http://') or src.startswith('https://'):
            try:
                response = requests.get(src, timeout=10)
                if response.status_code == 200:
                    return BytesIO(response.content)
            except Exception as e:
                logger.error(f"下载网络图片失败: {e}")
                return None
        
        # 本地图片
        else:
            image_path = src
            if image_base_dir and not os.path.isabs(image_path):
                image_path = os.path.join(image_base_dir, image_path)
            
            if os.path.exists(image_path):
                try:
                    with open(image_path, 'rb') as f:
                        return BytesIO(f.read())
                except Exception as e:
                    logger.error(f"读取本地图片失败: {e}")
                    return None
        
        return None

