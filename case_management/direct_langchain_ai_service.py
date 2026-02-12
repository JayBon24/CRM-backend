"""
直接使用LangChain调用各种大模型的AI服务
"""

import os
import json
import re
from typing import Dict, Any, List
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 尝试导入各种LangChain模型
try:
    from langchain_openai import ChatOpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from langchain_community.llms import Ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False

try:
    from langchain_google_genai import ChatGoogleGenerativeAI
    GOOGLE_AVAILABLE = True
except ImportError:
    GOOGLE_AVAILABLE = False

try:
    from langchain_anthropic import ChatAnthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

# 默认不提供任何硬编码密钥，要求从环境变量读取
DEFAULT_DEEPSEEK_API_KEY = ""

# 模型配置
MODEL_CONFIGS = {
    # DeepSeek 模型
    "deepseek-chat": {
        "provider": "openai",
        "model": "deepseek-chat",
        "api_key_env": "DEEPSEEK_API_KEY",
        "api_key_default": DEFAULT_DEEPSEEK_API_KEY,
        "base_url": "https://api.deepseek.com/v1",
        "available": OPENAI_AVAILABLE
    },
    
    # OpenAI 模型
    "gpt-3.5-turbo": {
        "provider": "openai",
        "model": "gpt-3.5-turbo",
        "api_key_env": "OPENAI_API_KEY",
        "available": OPENAI_AVAILABLE
    },
    "gpt-4": {
        "provider": "openai", 
        "model": "gpt-4",
        "api_key_env": "OPENAI_API_KEY",
        "available": OPENAI_AVAILABLE
    },
    
    # Google 模型
    "gemini-pro": {
        "provider": "google",
        "model": "gemini-pro",
        "api_key_env": "GOOGLE_API_KEY",
        "available": GOOGLE_AVAILABLE
    },
    
    # Anthropic 模型
    "claude-3-sonnet": {
        "provider": "anthropic",
        "model": "claude-3-sonnet-20240229",
        "api_key_env": "ANTHROPIC_API_KEY",
        "available": ANTHROPIC_AVAILABLE
    },
    
    # Ollama 本地模型
    "llama3": {
        "provider": "ollama",
        "model": "llama3", 
        "api_key_env": None,
        "available": OLLAMA_AVAILABLE
    }
}

# 按优先级排序的模型列表
PRIORITY_MODELS = [
    "deepseek-chat",      # DeepSeek Chat - 推荐
    "gpt-3.5-turbo",      # OpenAI GPT-3.5
    "gemini-pro",         # Google Gemini Pro - 免费
    "claude-3-sonnet",    # Anthropic Claude 3 Sonnet
    "gpt-4"               # OpenAI GPT-4
    # 注意：Ollama模型需要本地服务运行，暂时移除
]

def get_chat_model(model_name: str = None):
    """获取LangChain模型实例"""
    if model_name is None:
        model_name = get_best_available_model()
    
    if model_name not in MODEL_CONFIGS:
        raise ValueError(f"不支持的模型: {model_name}")
    
    config = MODEL_CONFIGS[model_name]
    
    if not config["available"]:
        raise ValueError(f"模型 {model_name} 不可用，请安装相应的依赖")
    
    # 检查API密钥
    if config["api_key_env"]:
        api_key = os.getenv(config["api_key_env"])
        if not api_key:
            # 尝试使用默认API密钥
            if "api_key_default" in config:
                api_key = config["api_key_default"]
                # 设置环境变量以确保后续调用也能使用
                os.environ[config["api_key_env"]] = api_key
            else:
                raise ValueError(f"请设置环境变量 {config['api_key_env']}")
    
    try:
        if config["provider"] == "openai":
            # 检查是否有自定义 base_url
            base_url = config.get("base_url")
            if base_url:
                return ChatOpenAI(
                    model=config["model"],
                    openai_api_key=os.getenv(config["api_key_env"]),
                    openai_api_base=base_url,
                    temperature=0.7,
                    max_tokens=4000,
                    timeout=120
                )
            else:
                return ChatOpenAI(
                    model=config["model"],
                    openai_api_key=os.getenv(config["api_key_env"]),
                    temperature=0.7,
                    max_tokens=4000,
                    timeout=120
                )
        elif config["provider"] == "google":
            return ChatGoogleGenerativeAI(
                model=config["model"],
                google_api_key=os.getenv(config["api_key_env"]),
                temperature=0.7,
                max_tokens=4000,
                timeout=30
            )
        elif config["provider"] == "anthropic":
            return ChatAnthropic(
                model=config["model"],
                anthropic_api_key=os.getenv(config["api_key_env"]),
                temperature=0.7,
                max_tokens=4000,
                timeout=30
            )
        elif config["provider"] == "ollama":
            return Ollama(
                model=config["model"],
                temperature=0.7,
                timeout=30
            )
        else:
            raise ValueError(f"不支持的提供商: {config['provider']}")
    except Exception as e:
        logger.error(f"创建模型 {model_name} 失败: {e}")
        raise

def get_best_available_model() -> str:
    """获取最佳可用模型"""
    for model_name in PRIORITY_MODELS:
        config = MODEL_CONFIGS[model_name]
        if config["available"]:
            # 检查API密钥
            if config["api_key_env"]:
                api_key = os.getenv(config["api_key_env"])
                if api_key:
                    logger.info(f"选择模型: {model_name}")
                    return model_name
            else:
                # 跳过没有API密钥要求的模型（如Ollama）
                continue
    
    # 如果没有找到可用模型，返回第一个
    logger.warning("没有找到可用的模型，使用默认模型")
    return PRIORITY_MODELS[0]


def parse_docx_file(filepath: str) -> str:
    """解析.docx文件 - 提取完整的格式和结构信息"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document(filepath)
        
        content_parts = []
        
        # 提取段落内容，保留完整格式信息
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                content_parts.append("")  # 保留空行
                continue
                
            # 检查段落样式
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # 构建格式化的段落内容
            formatted_text = paragraph.text.strip()
            
            # 检查段落整体格式
            is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
            is_italic = any(run.italic for run in paragraph.runs if run.text.strip())
            is_underline = any(run.underline for run in paragraph.runs if run.text.strip())
            
            # 提取字体信息
            font_name = None
            font_size = None
            if paragraph.runs:
                # 使用第一个run的字体信息作为段落代表
                first_run = next((run for run in paragraph.runs if run.text.strip()), None)
                if first_run:
                    font_name = first_run.font.name
                    font_size = first_run.font.size
            
            # 提取段落格式信息
            para_format = paragraph.paragraph_format
            left_indent = para_format.left_indent
            first_line_indent = para_format.first_line_indent
            space_before = para_format.space_before
            space_after = para_format.space_after
            line_spacing = para_format.line_spacing
            alignment = paragraph.alignment
            
            # 构建完整的格式标记
            format_tags = []
            
            # 字体大小标记
            if font_size:
                size_pt = font_size.pt if hasattr(font_size, 'pt') else font_size / 10000
                format_tags.append(f"<size:{int(size_pt)}>")
            
            # 字体名称标记
            if font_name and font_name != 'Normal':
                format_tags.append(f"<font:{font_name}>")
            
            # 段落缩进标记
            if first_line_indent and first_line_indent > 0:
                indent_pt = first_line_indent.pt if hasattr(first_line_indent, 'pt') else first_line_indent / 10000
                format_tags.append(f"<indent:{int(indent_pt)}>")
            
            # 段落间距标记
            if space_before and space_before > 0:
                before_pt = space_before.pt if hasattr(space_before, 'pt') else space_before / 10000
                format_tags.append(f"<space_before:{int(before_pt)}>")
            
            if space_after and space_after > 0:
                after_pt = space_after.pt if hasattr(space_after, 'pt') else space_after / 10000
                format_tags.append(f"<space_after:{int(after_pt)}>")
            
            # 行距标记
            if line_spacing and line_spacing != 1.0:
                if hasattr(line_spacing, 'pt'):
                    spacing_value = line_spacing.pt / 12
                else:
                    spacing_value = line_spacing / 10000 / 12
                spacing_value = min(max(spacing_value, 0.5), 3.0)
                format_tags.append(f"<line_spacing:{spacing_value}>")
            
            # 对齐方式标记
            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                format_tags.append("<align:center>")
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                format_tags.append("<align:right>")
            elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                format_tags.append("<align:justify>")
            
            # 应用格式标记
            if format_tags:
                formatted_text = ''.join(format_tags) + formatted_text
            
            # 根据格式添加标记
            if is_bold:
                formatted_text = f"**{formatted_text}**"
            if is_italic:
                formatted_text = f"*{formatted_text}*"
            if is_underline:
                formatted_text = f"<u>{formatted_text}</u>"
            
            # 根据段落样式添加标记
            if "Heading" in style_name or "Title" in style_name:
                level = 1
                if "Heading" in style_name:
                    try:
                        level = int(style_name.split()[-1])
                    except:
                        level = 1
                formatted_text = f"{'#' * level} {formatted_text}"
            elif style_name == "List Paragraph":
                formatted_text = f"- {formatted_text}"
            elif left_indent and left_indent > 0:
                indent_level = int(left_indent / Inches(0.5))
                formatted_text = "  " * indent_level + formatted_text
            
            content_parts.append(formatted_text)
        
        # 提取表格内容，保留格式
        for table in doc.tables:
            content_parts.append("")  # 表格前空行
            content_parts.append("| " + " | ".join(["列" + str(i+1) for i in range(len(table.columns))]) + " |")
            content_parts.append("| " + " | ".join(["---" for _ in range(len(table.columns))]) + " |")
            
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_content = ""
                    for run in cell.paragraphs[0].runs:
                        text = run.text
                        if run.bold:
                            text = f"**{text}**"
                        if run.italic:
                            text = f"*{text}*"
                        cell_content += text
                    row_text.append(cell_content.strip())
                if any(row_text):
                    content_parts.append("| " + " | ".join(row_text) + " |")
            content_parts.append("")  # 表格后空行
        
        return "\n".join(content_parts)
    except Exception as e:
        logger.error(f"解析.docx文件失败: {e}")
        return ""

def parse_doc_file(filepath: str) -> str:
    """解析.doc文件 - 尝试保留格式信息"""
    try:
        # 方法1: 尝试使用win32com（仅Windows）- 可以保留更多格式
        try:
            import win32com.client
            import os
            import pythoncom
            import time
            
            # 重试机制，最多尝试3次
            for attempt in range(3):
                try:
                    # 检查是否已经初始化COM接口
                    try:
                        pythoncom.CoInitialize()
                        com_initialized = True
                    except pythoncom.com_error:
                        # 如果已经初始化，继续使用
                        com_initialized = False
                    
                    word = None
                    doc = None
                    
                    try:
                        # 转换为绝对路径
                        abs_filepath = os.path.abspath(filepath)
                        
                        # 创建Word应用程序实例
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        word.DisplayAlerts = False  # 禁用警告
                        
                        # 等待Word完全启动
                        time.sleep(0.5)
                        
                        # 打开文档
                        doc = word.Documents.Open(abs_filepath, ReadOnly=True)
                        
                        # 尝试获取带格式的内容
                        content_parts = []
                        
                        # 遍历文档中的每个段落
                        for i in range(1, doc.Paragraphs.Count + 1):
                            try:
                                para = doc.Paragraphs(i)
                                text = para.Range.Text.strip()
                                
                                if text and text not in ['\r', '\n', '\t']:  # 过滤掉纯空白段落
                                    # 检查段落格式
                                    style_name = para.Style.NameLocal if hasattr(para.Style, 'NameLocal') else "Normal"
                                    
                                    # 简化格式处理，避免重复标签
                                    formatted_text = text
                                    
                                    # 检查段落整体格式
                                    if para.Range.Bold:
                                        formatted_text = f"**{formatted_text}**"
                                    if para.Range.Italic:
                                        formatted_text = f"*{formatted_text}*"
                                    if para.Range.Underline:
                                        formatted_text = f"<u>{formatted_text}</u>"
                                    
                                    # 根据样式添加标记
                                    if "标题" in style_name or "Heading" in style_name:
                                        formatted_text = f"# {formatted_text}"
                                    elif "列表" in style_name or "List" in style_name:
                                        formatted_text = f"- {formatted_text}"
                                    
                                    content_parts.append(formatted_text)
                            except Exception as e:
                                logger.warning(f"处理段落 {i} 时出错: {e}")
                                continue
                        
                        if content_parts:
                            logger.info(f"使用Word COM接口成功解析.doc文件 (尝试 {attempt + 1}/3)")
                            return "\n".join(content_parts)
                        else:
                            logger.warning(f"Word COM接口解析成功但未获取到内容 (尝试 {attempt + 1}/3)")
                            
                    finally:
                        # 确保资源被正确释放
                        try:
                            if doc:
                                doc.Close()
                        except:
                            pass
                        
                        try:
                            if word:
                                word.Quit()
                        except:
                            pass
                        
                        # 清理COM接口
                        if com_initialized:
                            try:
                                pythoncom.CoUninitialize()
                            except:
                                pass
                    
                    # 如果成功，跳出重试循环
                    break
                    
                except Exception as e:
                    logger.warning(f"Word COM接口第 {attempt + 1} 次尝试失败: {e}")
                    if attempt < 2:  # 不是最后一次尝试
                        time.sleep(1)  # 等待1秒后重试
                        continue
                    else:
                        raise e  # 最后一次尝试失败，抛出异常
            
        except ImportError:
            logger.warning("win32com库未安装，无法使用Word COM接口")
        except Exception as e:
            logger.warning(f"Word COM接口解析失败: {e}")
        
        # 方法2: 使用python-docx2txt作为备用方案
        try:
            import docx2txt
            content = docx2txt.process(filepath)
            if content and content.strip():
                logger.info("使用docx2txt成功解析.doc文件")
                return content
        except ImportError:
            logger.warning("docx2txt库未安装")
        except Exception as e:
            logger.warning(f"docx2txt解析失败: {e}")
        
        # 方法3: 使用python-docx作为备用方案（尝试将.doc当作.docx处理）
        try:
            from docx import Document
            # 尝试直接读取.doc文件
            doc = Document(filepath)
            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text.strip())
            if content_parts:
                logger.info("使用python-docx成功解析.doc文件")
                return "\n".join(content_parts)
        except ImportError:
            logger.warning("python-docx库未安装")
        except Exception as e:
            logger.warning(f"python-docx解析失败: {e}")
        
        # 方法4: 使用文本提取作为最后备用方案
        try:
            import subprocess
            import tempfile
            import os
            
            # 尝试使用antiword（如果可用）
            result = subprocess.run(['antiword', filepath], capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                logger.info("使用antiword成功解析.doc文件")
                return result.stdout
        except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
            logger.warning(f"antiword解析失败: {e}")
        
        # 方法5: 使用二进制读取作为最后备用方案
        try:
            with open(filepath, 'rb') as f:
                content = f.read()
            # 尝试提取文本内容（简单的方法）
            text_content = ""
            for byte in content:
                if 32 <= byte <= 126 or byte in [9, 10, 13]:  # 可打印字符和换行符
                    text_content += chr(byte)
                elif byte == 0:  # 空字符，可能是分隔符
                    text_content += " "
            
            # 清理文本
            lines = text_content.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line and len(line) > 3:  # 过滤掉太短的行
                    cleaned_lines.append(line)
            
            if cleaned_lines:
                logger.info("使用二进制读取成功解析.doc文件")
                return "\n".join(cleaned_lines)
        except Exception as e:
            logger.warning(f"二进制读取解析失败: {e}")
        
        logger.error("所有.doc文件解析方法都失败了")
        return ""
        
    except Exception as e:
        logger.error(f"解析.doc文件失败: {e}")
        return ""

def parse_xlsx_file(filepath: str) -> str:
    """解析.xlsx文件"""
    try:
        import pandas as pd
        # 读取所有工作表
        excel_file = pd.ExcelFile(filepath)
        content_parts = []
        
        for sheet_name in excel_file.sheet_names:
            content_parts.append(f"工作表: {sheet_name}")
            df = pd.read_excel(filepath, sheet_name=sheet_name)
            
            # 将DataFrame转换为文本
            if not df.empty:
                # 处理表头
                headers = df.columns.tolist()
                content_parts.append(" | ".join(str(h) for h in headers))
                
                # 处理数据行
                for _, row in df.iterrows():
                    row_data = [str(cell) if pd.notna(cell) else "" for cell in row]
                    content_parts.append(" | ".join(row_data))
            content_parts.append("")  # 工作表间空行
        
        return "\n".join(content_parts)
    except ImportError:
        logger.warning("pandas库未安装，无法解析.xlsx文件")
        return ""
    except Exception as e:
        logger.error(f"解析.xlsx文件失败: {e}")
        return ""

def parse_xls_file(filepath: str) -> str:
    """解析.xls文件"""
    try:
        import pandas as pd
        # 读取所有工作表
        excel_file = pd.ExcelFile(filepath)
        content_parts = []
        
        for sheet_name in excel_file.sheet_names:
            content_parts.append(f"工作表: {sheet_name}")
            df = pd.read_excel(filepath, sheet_name=sheet_name)
            
            # 将DataFrame转换为文本
            if not df.empty:
                # 处理表头
                headers = df.columns.tolist()
                content_parts.append(" | ".join(str(h) for h in headers))
                
                # 处理数据行
                for _, row in df.iterrows():
                    row_data = [str(cell) if pd.notna(cell) else "" for cell in row]
                    content_parts.append(" | ".join(row_data))
            content_parts.append("")  # 工作表间空行
        
        return "\n".join(content_parts)
    except ImportError:
        logger.warning("pandas库未安装，无法解析.xls文件")
        return ""
    except Exception as e:
        logger.error(f"解析.xls文件失败: {e}")
        return ""

def clean_format_tags(content: str) -> str:
    """清理格式标签，避免重复标记"""
    import re
    
    # 清理重复的加粗标记
    content = re.sub(r'\*\*([^*]+)\*\*\*\*([^*]+)\*\*', r'**\1\2**', content)
    content = re.sub(r'\*\*([^*]+)\*\*\*\*([^*]+)\*\*\*\*([^*]+)\*\*', r'**\1\2\3**', content)
    
    # 清理重复的斜体标记
    content = re.sub(r'\*([^*]+)\*\*([^*]+)\*', r'*\1\2*', content)
    
    # 清理重复的下划线标记
    content = re.sub(r'<u>([^<]+)</u><u>([^<]+)</u>', r'<u>\1\2</u>', content)
    
    return content




def convert_format_tags_to_html(content: str) -> str:
    """将格式标签转换为HTML格式"""
    import re
    
    try:
        if not content or not content.strip():
            return '<p>内容为空</p>'
            
        lines = content.split('\n')
        html_lines = []
        in_list = False
        in_table = False
        
        for i, line in enumerate(lines):
            line = line.rstrip()  # 移除行尾空白
            
            # 处理空行
            if not line.strip():
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                elif in_table:
                    html_lines.append('</table>')
                    in_table = False
                else:
                    html_lines.append('<br>')
                continue
            
            # 处理标题 - 支持多个星号开头
            if line.startswith('*'):
                # 关闭之前的列表或表格
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                # 计算开头的星号数量
                star_count = 0
                for char in line:
                    if char == '*':
                        star_count += 1
                    else:
                        break
                
                # 如果开头有星号，提取标题文本
                if star_count > 0:
                    title_text = line[star_count:].strip()
                    # 移除格式标记
                    title_text = re.sub(r'<[^>]+>', '', title_text).strip()
                    if title_text:
                        # 根据星号数量确定标题级别
                        level = min(star_count, 6)  # 最多6级标题
                        html_lines.append(f'<h{level}>{title_text}</h{level}>')
                        continue
            
            # 处理居中对齐
            if '<align:center>' in line:
                # 关闭之前的列表或表格
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                processed_line = re.sub(r'<align:center>', '', line)
                # 提取文本内容
                text_content = re.sub(r'<[^>]+>', '', processed_line).strip()
                if text_content:
                    html_lines.append(f'<div style="text-align: center; font-weight: bold; margin: 15px 0;">{text_content}</div>')
                    continue
            
            # 处理右对齐
            if '<align:right>' in line:
                # 关闭之前的列表或表格
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                processed_line = re.sub(r'<align:right>', '', line)
                # 提取文本内容
                text_content = re.sub(r'<[^>]+>', '', processed_line).strip()
                if text_content:
                    html_lines.append(f'<div style="text-align: right; color: #666; font-style: italic; margin: 10px 0;">{text_content}</div>')
                    continue
            
            # 处理缩进
            indent_match = re.search(r'<indent:(\d+)>', line)
            if indent_match:
                # 关闭之前的列表或表格
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                indent_level = int(indent_match.group(1))
                processed_line = re.sub(r'<indent:\d+>', '', line)
                # 添加缩进
                text_content = re.sub(r'<[^>]+>', '', processed_line).strip()
                if text_content:
                    indent_px = indent_level * 20  # 每级缩进20px
                    html_lines.append(f'<div style="margin-left: {indent_px}px; border-left: 3px solid #409eff; padding-left: 10px; margin: 10px 0;">{text_content}</div>')
                    continue
            
            # 处理列表项
            if line.strip().startswith('- '):
                if not in_list:
                    html_lines.append('<ul>')
                    in_list = True
                # 关闭表格
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                list_text = line.strip()[2:].strip()
                # 处理列表项内的格式
                list_text = process_inline_formatting(list_text)
                html_lines.append(f'<li>{list_text}</li>')
                continue
            
            # 处理表格行
            if line.strip().startswith('|'):
                if not in_table:
                    html_lines.append('<table style="width: 100%; border-collapse: collapse; margin: 15px 0; box-shadow: 0 1px 3px rgba(0,0,0,0.1);">')
                    in_table = True
                # 关闭列表
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                
                # 简单的表格处理
                cells = line.strip().split('|')[1:-1]  # 去掉首尾空元素
                cell_html = ''.join(f'<td style="border: 1px solid #ddd; padding: 8px 12px; text-align: left;">{cell.strip()}</td>' for cell in cells)
                html_lines.append(f'<tr>{cell_html}</tr>')
                continue
            
            # 处理普通段落
            if line.strip():
                # 关闭之前的列表或表格
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                # 处理段落内的格式
                processed_line = process_inline_formatting(line)
                html_lines.append(f'<p style="margin: 12px 0; text-indent: 2em; line-height: 1.8;">{processed_line}</p>')
            else:
                html_lines.append('<br>')
        
        # 关闭未关闭的列表或表格
        if in_list:
            html_lines.append('</ul>')
        if in_table:
            html_lines.append('</table>')
        
        return '\n'.join(html_lines)
        
    except Exception as e:
        logger.error(f"转换格式标签为HTML失败: {e}")
        return f'<p>格式转换失败: {str(e)}</p><pre>{content}</pre>'

def process_inline_formatting(text: str) -> str:
    """处理行内格式"""
    import re
    
    # 处理加粗文本 - 黑体转换为加粗
    if '<font:黑体>' in text or '黑体' in text:
        text = re.sub(r'<font:黑体>', '<strong>', text)
        text = re.sub(r'<font:[^>]+>', '', text)
        # 确保加粗标记成对
        if text.count('<strong>') % 2 == 1:
            text += '</strong>'
    
    # 处理下划线
    text = re.sub(r'<u>(.*?)</u>', r'<u>\1</u>', text)
    
    # 处理斜体
    text = re.sub(r'<i>(.*?)</i>', r'<em>\1</em>', text)
    
    # 移除其他格式标记
    text = re.sub(r'<size:\d+>', '', text)
    text = re.sub(r'<font:[^>]+>', '', text)
    text = re.sub(r'<space_before:\d+>', '', text)
    text = re.sub(r'<space_after:\d+>', '', text)
    text = re.sub(r'<line_spacing:[\d.]+>', '', text)
    
    return text

def convert_format_tags_to_markdown(content: str) -> str:
    """将格式标签转换为Markdown格式"""
    import re
    
    try:
        lines = content.split('\n')
        markdown_lines = []
        
        for line in lines:
            if not line.strip():
                markdown_lines.append('')
                continue
            
            # 保存原始行用于处理
            original_line = line
            processed_line = line
            
            # 处理标题 - 支持多个星号开头
            if processed_line.startswith('*'):
                # 计算开头的星号数量
                star_count = 0
                for char in processed_line:
                    if char == '*':
                        star_count += 1
                    else:
                        break
                
                # 如果开头有星号，提取标题文本
                if star_count > 0:
                    title_text = processed_line[star_count:].strip()
                    # 移除格式标记
                    title_text = re.sub(r'<[^>]+>', '', title_text).strip()
                    if title_text:
                        markdown_lines.append(f"# {title_text}")
                        continue
            
            # 处理居中对齐
            if '<align:center>' in processed_line:
                processed_line = re.sub(r'<align:center>', '', processed_line)
                # 提取文本内容
                text_content = re.sub(r'<[^>]+>', '', processed_line).strip()
                if text_content:
                    markdown_lines.append(f"<div style='text-align: center;'>{text_content}</div>")
                    continue
            
            # 处理右对齐
            if '<align:right>' in processed_line:
                processed_line = re.sub(r'<align:right>', '', processed_line)
                # 提取文本内容
                text_content = re.sub(r'<[^>]+>', '', processed_line).strip()
                if text_content:
                    markdown_lines.append(f"<div style='text-align: right;'>{text_content}</div>")
                    continue
            
            # 处理缩进
            indent_match = re.search(r'<indent:(\d+)>', processed_line)
            if indent_match:
                indent_level = int(indent_match.group(1))
                processed_line = re.sub(r'<indent:\d+>', '', processed_line)
                # 添加缩进
                text_content = re.sub(r'<[^>]+>', '', processed_line).strip()
                if text_content:
                    indent_spaces = '&nbsp;' * (indent_level // 4)  # 每4个单位一个缩进
                    markdown_lines.append(f"{indent_spaces}{text_content}")
                    continue
            
            # 处理加粗文本 - 黑体转换为加粗
            if '<font:黑体>' in processed_line or '黑体' in processed_line:
                # 提取黑体文本
                processed_line = re.sub(r'<font:黑体>', '**', processed_line)
                processed_line = re.sub(r'<font:[^>]+>', '', processed_line)
                # 确保加粗标记成对
                if processed_line.count('**') % 2 == 1:
                    processed_line += '**'
            
            # 处理下划线
            processed_line = re.sub(r'<u>(.*?)</u>', r'<u>\1</u>', processed_line)
            
            # 移除其他格式标记
            processed_line = re.sub(r'<size:\d+>', '', processed_line)
            processed_line = re.sub(r'<font:[^>]+>', '', processed_line)
            processed_line = re.sub(r'<space_before:\d+>', '', processed_line)
            processed_line = re.sub(r'<space_after:\d+>', '', processed_line)
            processed_line = re.sub(r'<line_spacing:[\d.]+>', '', processed_line)
            
            # 处理列表项
            if processed_line.strip().startswith('- '):
                markdown_lines.append(processed_line)
                continue
            
            # 处理表格行
            if processed_line.strip().startswith('|'):
                markdown_lines.append(processed_line)
                continue
            
            # 处理普通段落
            markdown_lines.append(processed_line)
        
        return '\n'.join(markdown_lines)
        
    except Exception as e:
        logger.error(f"转换格式标签失败: {e}")
        return content




def get_file_type(filename: str) -> str:
    """根据文件名获取文件类型"""
    if not filename:
        return 'unknown'
    ext = os.path.splitext(filename)[1].lower()
    type_map = {
        '.txt': 'text',
        '.doc': 'word_old',
        '.docx': 'word_new',
        '.xls': 'excel_old',
        '.xlsx': 'excel_new'
    }
    return type_map.get(ext, 'unknown')


def load_template_files(use_unstructured: bool = False) -> List[Dict[str, str]]:
    """加载模板文件 - 支持多种文件格式"""
    if use_unstructured:
        # 使用Unstructured库解析
        try:
            from .unstructured_document_service import unstructured_service
            return unstructured_service.load_template_files()
        except Exception as e:
            logger.error(f"使用Unstructured加载模板失败: {e}")
            # 回退到原有方法
            pass
    
    templates = []
    
    # 从数据库加载模板记录
    try:
        from .models import DocumentTemplate
        template_records = DocumentTemplate.objects.filter(is_active=True, is_deleted=False)
        
        if not template_records.exists():
            logger.warning("数据库中没有找到启用的模板记录，请先在模版管理中上传模板文件")
            return templates
        
        for template_record in template_records:
            filepath = template_record.full_file_path
            filename = os.path.basename(filepath)
            
            if not os.path.exists(filepath):
                logger.warning(f"模板文件不存在: {filepath}，跳过该模板")
                continue
                
            # 根据文件扩展名确定文件类型
            file_type = get_file_type(filename)
            doc_type = extract_document_type_from_filename(filename)
            
            try:
                if file_type == 'text':
                    # 文本文件
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                    logger.info(f"成功加载文本模板: {filename}")
                    
                elif file_type == 'word_new':
                    # 新版Word文档 (.docx)
                    content = parse_docx_file(filepath)
                    logger.info(f"成功加载Word模板: {filename}")
                    
                elif file_type == 'word_old':
                    # 旧版Word文档 (.doc) - 使用python-docx2txt
                    content = parse_doc_file(filepath)
                    logger.info(f"成功加载Word模板: {filename}")
                    
                elif file_type == 'excel_new':
                    # 新版Excel文档 (.xlsx)
                    content = parse_xlsx_file(filepath)
                    logger.info(f"成功加载Excel模板: {filename}")
                    
                elif file_type == 'excel_old':
                    # 旧版Excel文档 (.xls)
                    content = parse_xls_file(filepath)
                    logger.info(f"成功加载Excel模板: {filename}")
                    
                else:
                    logger.warning(f"不支持的文件类型: {filename}")
                    continue
                
                if not content or content.strip() == "":
                    logger.error(f"文件 {filename} 解析内容为空，跳过该模板")
                    continue
                
                # 保留原始格式标签，不进行任何转换
                templates.append({
                    "name": filename,
                    "content": content,  # 直接使用原始内容，保留所有格式标签
                    "is_binary": False,
                    "file_type": file_type,
                    "template_type": template_record.template_type,
                    "template_id": template_record.id
                })
                
            except Exception as e:
                logger.error(f"解析模板文件 {filename} 时出错: {e}")
                continue
                
        logger.info(f"从数据库成功加载 {len(templates)} 个模板文件")
        return templates
        
    except Exception as e:
        logger.error(f"从模版管理数据库加载模板失败: {e}")
        return templates


def load_template_files_from_filesystem():
    """从文件系统加载模板文件（回退方案）"""
    templates = []
    # 使用 backend/template 文件夹
    template_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'template')
    
    if not os.path.exists(template_dir):
        logger.warning(f"模板文件夹不存在: {template_dir}")
        return templates

    # 支持的文件类型
    supported_extensions = {
        '.txt': 'text',
        '.doc': 'word_old',
        '.docx': 'word_new',
        '.xls': 'excel_old',
        '.xlsx': 'excel_new'
    }

    for filename in os.listdir(template_dir):
        if filename.startswith('~$'):  # 跳过临时文件
            continue
            
        filepath = os.path.join(template_dir, filename)
        if not os.path.isfile(filepath):
            continue
            
        # 获取文件扩展名
        if not filename:
            continue
        file_ext = os.path.splitext(filename)[1].lower()
        file_type = supported_extensions.get(file_ext, 'unknown')
        
        if file_type == 'unknown':
            logger.warning(f"不支持的文件类型: {filename}")
            continue
            
        try:
            content = ""
            doc_type = extract_document_type_from_filename(filename)
            
            if file_type == 'text':
                # 文本文件
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                logger.info(f"成功加载文本模板: {filename}")
                
            elif file_type == 'word_new':
                # 新版Word文档 (.docx)
                content = parse_docx_file(filepath)
                logger.info(f"成功加载Word模板: {filename}")
                
            elif file_type == 'word_old':
                # 旧版Word文档 (.doc) - 使用python-docx2txt
                content = parse_doc_file(filepath)
                logger.info(f"成功加载Word模板: {filename}")
                
            elif file_type == 'excel_new':
                # 新版Excel文档 (.xlsx)
                content = parse_xlsx_file(filepath)
                logger.info(f"成功加载Excel模板: {filename}")
                
            elif file_type == 'excel_old':
                # 旧版Excel文档 (.xls)
                content = parse_xls_file(filepath)
                logger.info(f"成功加载Excel模板: {filename}")
            
            # 如果解析出的内容为空，跳过该文件
            if not content.strip():
                logger.error(f"文件 {filename} 解析内容为空，跳过该模板")
                continue
            
            # 保留原始格式标签，不进行任何转换
            templates.append({
                "name": filename,
                "content": content,  # 直接使用原始内容，保留所有格式标签
                "is_binary": False,
                "file_type": file_type
            })
            
        except Exception as e:
            logger.error(f"读取模板文件 {filename} 失败: {e}")
            # 如果解析失败，跳过该文件
            continue
    
    logger.info(f"从文件系统成功加载 {len(templates)} 个模板文件")
    return templates

def generate_document_with_langchain(case_data: Dict[str, Any], document_type: str, template_info: Dict[str, Any] = None) -> Dict[str, Any]:
    """使用LangChain生成单个法律文书（优化版 - 保持模板格式）"""
    try:
        model = get_chat_model()
        
        # 启用调试模式
        import logging
        logging.getLogger("langchain").setLevel(logging.DEBUG)
        
        # 打印原始传入的case_data
        print(f"\n🔍 原始传入的case_data:")
        print(f"类型: {type(case_data)}")
        print(f"内容: {case_data}")
        print(f"键值: {list(case_data.keys()) if isinstance(case_data, dict) else 'Not a dict'}")
        
        # 准备案例数据，确保所有字段都有值
        # 首先尝试从case_data直接获取，如果没有则尝试从嵌套对象获取
        def get_nested_value(data, key, default=''):
            """递归获取嵌套字典中的值"""
            if isinstance(data, dict):
                if key in data:
                    return data[key]
                # 尝试从嵌套对象中获取
                for k, v in data.items():
                    if isinstance(v, dict) and key in v:
                        return v[key]
            return default
        
        def get_plaintiff_info(data, field, default='待填写'):
            """获取原告信息"""
            # 尝试多种可能的路径
            paths = [
                f'plaintiff_{field}',
                f'plaintiff.{field}',
                f'plaintiff_info.{field}',
                f'plaintiff_info.{field}',
                f'申请人{field}',
                f'申请人_{field}'
            ]
            
            for path in paths:
                if '.' in path:
                    # 处理嵌套路径
                    parts = path.split('.')
                    value = data
                    for part in parts:
                        if isinstance(value, dict) and part in value:
                            value = value[part]
                        else:
                            value = None
                            break
                    if value is not None:
                        return value
                else:
                    # 直接路径
                    if path in data:
                        return data[path]
            
            # 如果找不到，检查是否有简化的字段名
            if field == 'name' and 'plaintiff_name' in data:
                return data['plaintiff_name']
            elif field == 'address' and 'plaintiff_address' in data:
                return data['plaintiff_address']
            elif field == 'credit_code' and 'plaintiff_credit_code' in data:
                return data['plaintiff_credit_code']
            elif field == 'legal_representative' and 'plaintiff_legal_representative' in data:
                return data['plaintiff_legal_representative']
            
            return default
        
        def get_defendant_info(data, field, default='待填写'):
            """获取被告信息"""
            # 尝试多种可能的路径
            paths = [
                f'defendant_{field}',
                f'defendant.{field}',
                f'defendant_info.{field}',
                f'defendant_info.{field}',
                f'被申请人{field}',
                f'被申请人_{field}'
            ]
            
            for path in paths:
                if '.' in path:
                    # 处理嵌套路径
                    parts = path.split('.')
                    value = data
                    for part in parts:
                        if isinstance(value, dict) and part in value:
                            value = value[part]
                        else:
                            value = None
                            break
                    if value is not None:
                        return value
                else:
                    # 直接路径
                    if path in data:
                        return data[path]
            
            # 如果找不到，检查是否有简化的字段名
            if field == 'name' and 'defendant_name' in data:
                return data['defendant_name']
            elif field == 'address' and 'defendant_address' in data:
                return data['defendant_address']
            elif field == 'credit_code' and 'defendant_credit_code' in data:
                return data['defendant_credit_code']
            elif field == 'legal_representative' and 'defendant_legal_representative' in data:
                return data['defendant_legal_representative']
            
            return default
        
        def get_amount_info(data, field, default=0):
            """获取金额信息"""
            # 尝试多种可能的路径
            paths = [
                field,
                f'amount_info.{field}',
                f'amount.{field}',
                f'金额{field}',
                f'金额_{field}'
            ]
            
            for path in paths:
                if '.' in path:
                    # 处理嵌套路径
                    parts = path.split('.')
                    value = data
                    for part in parts:
                        if isinstance(value, dict) and part in value:
                            value = value[part]
                        else:
                            value = None
                            break
                    if value is not None:
                        return value
                else:
                    # 直接路径
                    if path in data:
                        return data[path]
            
            # 如果找不到，检查是否有简化的字段名
            if field == 'contract_amount' and 'contract_amount' in data:
                return data['contract_amount']
            elif field == 'lawyer_fee' and 'lawyer_fee' in data:
                return data['lawyer_fee']
            elif field == 'total_amount' and 'total_amount' in data:
                return data['total_amount']
            
            return default
        
        processed_case_data = {
            # 基本信息
            'case_number': get_nested_value(case_data, 'case_number', ''),
            'case_name': get_nested_value(case_data, 'case_name', ''),
            'case_type': get_nested_value(case_data, 'case_type', ''),
            'jurisdiction': get_nested_value(case_data, 'jurisdiction', ''),
            'draft_person': get_nested_value(case_data, 'draft_person', ''),
            'case_description': get_nested_value(case_data, 'case_description', ''),
            'status': get_nested_value(case_data, 'status', 'draft'),
            
            # 被告信息 - 使用专门的函数
            'defendant_name': get_defendant_info(case_data, 'name'),
            'defendant_credit_code': get_defendant_info(case_data, 'credit_code'),
            'defendant_address': get_defendant_info(case_data, 'address'),
            'defendant_legal_representative': get_defendant_info(case_data, 'legal_representative'),
            
            # 原告信息 - 使用专门的函数
            'plaintiff_name': get_plaintiff_info(case_data, 'name'),
            'plaintiff_credit_code': get_plaintiff_info(case_data, 'credit_code'),
            'plaintiff_address': get_plaintiff_info(case_data, 'address'),
            'plaintiff_legal_representative': get_plaintiff_info(case_data, 'legal_representative'),
            
            # 金额信息 - 使用专门的函数
            'contract_amount': float(get_amount_info(case_data, 'contract_amount') or 0),
            'lawyer_fee': float(get_amount_info(case_data, 'lawyer_fee') or 0),
            
            # 其他信息
            'create_date': '',
            'defendant_count': '1',
            'total_amount': float(get_amount_info(case_data, 'contract_amount') or 0) + float(get_amount_info(case_data, 'lawyer_fee') or 0)
        }
        
        # 构建系统提示词
        if template_info:
            # 让大模型进行智能填充，保持格式
            template_content = template_info.get('content', '')
            
            system_prompt = f"""你是专业的法律文书生成助手。请根据模板和案例信息生成{document_type}，严格保持模板格式。

**模板内容**：
{template_content}

**案例信息**：
案件编号：{processed_case_data['case_number']}
案件名称：{processed_case_data['case_name']}
案件类型：{processed_case_data['case_type']}
管辖法院：{processed_case_data['jurisdiction']}
拟稿人：{processed_case_data['draft_person']}
案件状态：{processed_case_data['status']}
案件描述：{processed_case_data['case_description']}

原告名称：{processed_case_data['plaintiff_name']}
原告所住地：{processed_case_data['plaintiff_address']}
原告统一社会信用代码：{processed_case_data['plaintiff_credit_code']}
原告法定代表人：{processed_case_data['plaintiff_legal_representative']}

被告名称：{processed_case_data['defendant_name']}
被告所住地：{processed_case_data['defendant_address']}
被告统一社会信用代码：{processed_case_data['defendant_credit_code']}
被告法定代表人：{processed_case_data['defendant_legal_representative']}

合同金额：{processed_case_data['contract_amount']}元
律师费：{processed_case_data['lawyer_fee']}元
总金额：{processed_case_data['total_amount']}元

**重要要求**：
1. 完全按照模板格式生成，保留所有格式标记（<size:>、<font:>、<align:>、<indent:>等）
2. 用案例信息替换模板中的示例数据，保持格式标记不变
3. 保持法律文书的专业性和规范性
4. 如果案例信息为"待填写"，请保持原样
5. 特别注意：要正确填充原告、被告的所住地、信用代码、法定代表人信息
6. 金额信息要准确填充合同金额、律师费、总金额

**格式说明**：
- <size:N>：字体大小
- <font:NAME>：字体名称
- <align:center>：居中对齐
- <align:right>：右对齐
- <indent:N>：段落缩进
- **文本**：加粗文本
- <u>文本</u>：下划线文本

请生成文档："""
            
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=f"请生成{document_type}：")
            ]
            
            # 打印传给大模型的所有数据
            print("\n" + "="*80)
            print("🔍 传给大模型的数据调试信息")
            print("="*80)
            
            print(f"\n📋 案例数据 (processed_case_data):")
            for key, value in processed_case_data.items():
                print(f"  {key}: {value}")
            
            print(f"\n📄 模板内容 (template_content):")
            template_lines = template_content.split('\n')
            for i, line in enumerate(template_lines[:20]):  # 只显示前20行
                print(f"  {i+1:2d}: {line}")
            if len(template_lines) > 20:
                print(f"  ... (还有 {len(template_lines) - 20} 行)")
            
            print(f"\n🤖 系统提示词 (system_prompt):")
            print(system_prompt)
            
            print(f"\n💬 用户消息 (human_message):")
            print(f"请生成{document_type}：")
            
            print(f"\n🚀 开始调用大模型...")
            print("="*80)
            
            response = model.invoke(messages)
            
            print(f"\n📤 大模型响应:")
            print(f"响应内容长度: {len(response.content)} 字符")
            print(f"响应内容预览:")
            response_lines = response.content.split('\n')
            for i, line in enumerate(response_lines[:20]):  # 只显示前20行
                print(f"  {i+1:2d}: {line}")
            if len(response_lines) > 20:
                print(f"  ... (还有 {len(response_lines) - 20} 行)")
            print("="*80)
            
            # 转换格式标签为实际展示效果
            formatted_content = convert_format_tags_to_display(response.content)
            
            return {
                'success': True,
                'content': formatted_content,  # 使用转换后的内容
                'document_name': document_type
            }
        else:
            # 如果没有模板，使用标准格式生成
            system_prompt = f"""你是一个专业的法律文书生成助手。请根据案例信息生成一份完整的{document_type}，使用标准法律文书格式。

完整案例信息：
【基本信息】
- 案件编号：{processed_case_data['case_number']}
- 案件名称：{processed_case_data['case_name']}
- 案件类型：{processed_case_data['case_type']}
- 管辖法院：{processed_case_data['jurisdiction']}
- 拟稿人：{processed_case_data['draft_person']}
- 案件状态：{processed_case_data['status']}
- 案件描述：{processed_case_data['case_description']}

【原告信息】
- 原告名称：{processed_case_data['plaintiff_name']}
- 原告所住地：{processed_case_data['plaintiff_address']}
- 原告统一社会信用代码：{processed_case_data['plaintiff_credit_code']}
- 原告法定代表人：{processed_case_data['plaintiff_legal_representative']}

【被告信息】
- 被告名称：{processed_case_data['defendant_name']}
- 被告所住地：{processed_case_data['defendant_address']}
- 被告统一社会信用代码：{processed_case_data['defendant_credit_code']}
- 被告法定代表人：{processed_case_data['defendant_legal_representative']}

【金额信息】
- 合同金额：{processed_case_data['contract_amount']}元
- 律师费：{processed_case_data['lawyer_fee']}元
- 总金额：{processed_case_data['total_amount']}元

【其他信息】
- 创建日期：{processed_case_data['create_date']}
- 被告数量：{processed_case_data['defendant_count']}

请按照以下要求生成文档：
1. 生成标准的{document_type}格式
2. 包含所有必要的法律文书要素
3. 根据文档类型和案例信息智能选择合适的案例信息进行填充
4. 保持法律文书的专业性和规范性
5. 使用标准法律文书格式（仿宋字体、适当缩进等）
6. 确保填充的信息准确、完整且符合法律文书规范

请生成完整的{document_type}："""
        
            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=f"请生成{document_type}：")
            ]
            
            # 打印传给大模型的所有数据
            print("\n" + "="*80)
            print("🔍 传给大模型的数据调试信息 (无模板)")
            print("="*80)
            
            print(f"\n📋 案例数据 (processed_case_data):")
            for key, value in processed_case_data.items():
                print(f"  {key}: {value}")
            
            print(f"\n🤖 系统提示词 (system_prompt):")
            print(system_prompt)
            
            print(f"\n💬 用户消息 (human_message):")
            print(f"请生成{document_type}：")
            
            print(f"\n🚀 开始调用大模型...")
            print("="*80)
            
            response = model.invoke(messages)
            
            print(f"\n📤 大模型响应:")
            print(f"响应内容长度: {len(response.content)} 字符")
            print(f"响应内容预览:")
            response_lines = response.content.split('\n')
            for i, line in enumerate(response_lines[:20]):  # 只显示前20行
                print(f"  {i+1:2d}: {line}")
            if len(response_lines) > 20:
                print(f"  ... (还有 {len(response_lines) - 20} 行)")
            print("="*80)
            
            # 转换格式标签为实际展示效果
            formatted_content = convert_format_tags_to_display(response.content)
            
            return {
                'success': True,
                'content': formatted_content,  # 使用转换后的内容
                'document_name': document_type
            }
    except Exception as e:
        logger.error(f"生成文档失败: {e}")
        import traceback
        logger.error(f"详细错误信息: {traceback.format_exc()}")
        return {
            'success': False,
            'error': str(e),
            'content': f"生成文档失败: {str(e)}"
        }


def generate_all_documents_with_langchain(case_data: Dict[str, Any], use_unstructured: bool = False) -> Dict[str, Any]:
    """使用LangChain根据所有模板生成法律文书"""
    if use_unstructured:
        # 使用Unstructured库生成
        try:
            from .unstructured_document_service import unstructured_service
            return unstructured_service.generate_all_documents(case_data)
        except Exception as e:
            logger.error(f"使用Unstructured生成文档失败: {e}")
            # 回退到原有方法
            pass
    
    # 加载所有模板文件
    templates = load_template_files(use_unstructured=False)
    
    if not templates:
        logger.error("没有找到任何模板文件")
        return {
            'success': False,
            'documents': [],
            'total_count': 0,
            'success_count': 0,
            'error_count': 0,
            'error': '没有找到任何模板文件'
        }
    
    logger.info(f"找到 {len(templates)} 个模板文件，开始生成文档...")
    
    generated_documents = []
    success_count = 0
    generated_types = set()  # 用于防止重复生成相同类型的文档
    
    for template in templates:
        # 从文件名提取文档类型
        template_name = template['name']
        document_type = extract_document_type_from_filename(template_name)
        
        # 防止重复生成相同类型的文档
        if document_type in generated_types:
            logger.warning(f"跳过重复的文档类型: {document_type}")
            continue
        
        generated_types.add(document_type)
        
        logger.info(f"正在生成文档: {document_type} (模板: {template_name})")
        result = generate_document_with_langchain(case_data, document_type, template)
        
        if result.get('success', False):
            generated_documents.append({
                "document_name": template_name,  # 使用模板文件名作为文档名称
                "template_name": template_name,
                "content": result['content'],
                "success": True
            })
            success_count += 1
            logger.info(f"成功生成文档: {document_type} (模板: {template_name})")
        else:
            generated_documents.append({
                "document_name": template_name,  # 使用模板文件名作为文档名称
                "template_name": template_name,
                "content": result.get('content', ''),
                "success": False,
                "error": result.get('error', '未知错误')
            })
            logger.error(f"生成文档失败: {document_type} (模板: {template_name}) - {result.get('error', '未知错误')}")

    return {
        'success': success_count > 0,
        'documents': generated_documents,
        'total_count': len(templates),
        'success_count': success_count,
        'error_count': len(templates) - success_count
    }

def extract_document_type_from_filename(filename: str) -> str:
    """从文件名动态提取文档类型"""
    # 移除文件扩展名
    name_without_ext = os.path.splitext(filename)[0]
    
    # 移除数字前缀和括号内容
    import re
    
    # 匹配 "数字、文档类型（其他内容）" 格式
    match = re.match(r'^\d+、(.+?)(?:（.*?）)?$', name_without_ext)
    if match:
        doc_type = match.group(1).strip()
        # 截断到20个字符以内，避免数据库字段长度限制
        return doc_type[:20] if len(doc_type) > 20 else doc_type
    
    # 匹配 "数字.文档类型" 格式
    match = re.match(r'^\d+\.(.+?)$', name_without_ext)
    if match:
        doc_type = match.group(1).strip()
        return doc_type[:20] if len(doc_type) > 20 else doc_type
    
    # 匹配 "数字-文档类型" 格式
    match = re.match(r'^\d+-(.+?)$', name_without_ext)
    if match:
        doc_type = match.group(1).strip()
        return doc_type[:20] if len(doc_type) > 20 else doc_type
    
    # 匹配 "数字_文档类型" 格式
    match = re.match(r'^\d+_(.+?)$', name_without_ext)
    if match:
        doc_type = match.group(1).strip()
        return doc_type[:20] if len(doc_type) > 20 else doc_type
    
    # 如果没有匹配到特定格式，直接返回原文件名（去掉扩展名）
    return name_without_ext[:15] if len(name_without_ext) > 15 else name_without_ext

def ai_chat_with_langchain(message: str, uploaded_files: List[str] = None) -> Dict[str, Any]:
    """使用LangChain进行AI对话"""
    try:
        model = get_chat_model()
        
        messages = [
            SystemMessage(content="你是一个专业的法律助手，请根据用户的问题提供专业的法律咨询和文书生成服务。"),
            HumanMessage(content=message)
        ]
        
        if uploaded_files:
            for file_content in uploaded_files:
                messages.append(HumanMessage(content=f"用户上传了文件，内容如下：\n{file_content}\n请根据文件内容进行分析和回复。"))
        
        chain = ChatPromptTemplate.from_messages(messages) | model | StrOutputParser()
        response_content = chain.invoke({})
        
        return {
            'success': True,
            'content': response_content
        }
    except Exception as e:
        logger.error(f"AI对话失败: {e}")
        return {
            'success': False,
            'error': str(e),
            'content': f"抱歉，AI服务暂时不可用：{str(e)}"
        }

def convert_docx_format_to_html(content: str) -> str:
    """
    专门用于Word文档预览的格式转换函数
    更好地保留Word文档的原始格式
    """
    try:
        import re
        
        if not content or not content.strip():
            return '<p>内容为空</p>'
        
        lines = content.split('\n')
        html_lines = []
        in_list = False
        in_table = False
        
        for line in lines:
            line = line.rstrip()
            
            # 处理空行
            if not line.strip():
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                elif in_table:
                    html_lines.append('</table>')
                    in_table = False
                else:
                    html_lines.append('<br>')
                continue
            
            # 处理标题 - 支持 # 开头的标题
            if line.strip().startswith('#'):
                # 关闭之前的列表或表格
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                # 计算 # 的数量
                hash_count = 0
                for char in line:
                    if char == '#':
                        hash_count += 1
                    else:
                        break
                
                title_text = line[hash_count:].strip()
                if title_text:
                    level = min(hash_count, 6)
                    html_lines.append(f'<h{level} style="margin: 20px 0 10px 0; font-weight: bold; color: #2c3e50;">{title_text}</h{level}>')
                    continue
            
            # 处理列表项
            if line.strip().startswith('- '):
                if not in_list:
                    html_lines.append('<ul style="margin: 15px 0; padding-left: 25px;">')
                    in_list = True
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                list_text = line.strip()[2:].strip()
                list_text = process_word_inline_formatting(list_text)
                html_lines.append(f'<li style="margin: 8px 0; line-height: 1.6;">{list_text}</li>')
                continue
            
            # 处理表格行
            if line.strip().startswith('|'):
                if not in_table:
                    html_lines.append('<table style="width: 100%; border-collapse: collapse; margin: 15px 0; box-shadow: 0 1px 3px rgba(0,0,0,0.1);">')
                    in_table = True
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                
                cells = line.strip().split('|')[1:-1]
                cell_html = ''.join(f'<td style="border: 1px solid #ddd; padding: 8px 12px; text-align: left;">{cell.strip()}</td>' for cell in cells)
                html_lines.append(f'<tr>{cell_html}</tr>')
                continue
            
            # 处理普通段落
            if line.strip():
                # 关闭之前的列表或表格
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if in_table:
                    html_lines.append('</table>')
                    in_table = False
                
                # 处理段落内的格式
                processed_line = process_word_inline_formatting(line)
                
                # 检查是否有格式标记
                has_format_tags = any(tag in line for tag in ['<align:', '<indent:', '<size:', '<font:', '<space_', '<line_spacing:'])
                
                if has_format_tags:
                    # 处理格式标记
                    processed_line = process_word_format_tags(processed_line)
                else:
                    # 普通段落
                    processed_line = f'<p style="margin: 12px 0; text-indent: 2em; line-height: 1.8;">{processed_line}</p>'
                
                html_lines.append(processed_line)
            else:
                html_lines.append('<br>')
        
        # 关闭未关闭的列表或表格
        if in_list:
            html_lines.append('</ul>')
        if in_table:
            html_lines.append('</table>')
        
        return '\n'.join(html_lines)
        
    except Exception as e:
        logger.error(f"转换Word格式失败: {e}")
        return f'<p>格式转换失败: {str(e)}</p><pre>{content}</pre>'

def process_word_inline_formatting(text: str) -> str:
    """处理Word文档的行内格式"""
    import re
    
    # 处理加粗文本
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
    
    # 处理斜体
    text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
    
    # 处理下划线
    text = re.sub(r'<u>(.*?)</u>', r'<u>\1</u>', text)
    
    return text

def process_word_format_tags(text: str) -> str:
    """处理Word文档的格式标记"""
    import re
    
    # 处理居中对齐
    if '<align:center>' in text:
        text = re.sub(r'<align:center>', '', text)
        text = re.sub(r'<[^>]+>', '', text).strip()
        return f'<div style="text-align: center; font-weight: bold; margin: 15px 0; font-size: 16px;">{text}</div>'
    
    # 处理右对齐
    if '<align:right>' in text:
        text = re.sub(r'<align:right>', '', text)
        text = re.sub(r'<[^>]+>', '', text).strip()
        return f'<div style="text-align: right; color: #666; font-style: italic; margin: 10px 0;">{text}</div>'
    
    # 处理缩进
    indent_match = re.search(r'<indent:(\d+)>', text)
    if indent_match:
        indent_level = int(indent_match.group(1))
        text = re.sub(r'<indent:\d+>', '', text)
        text = re.sub(r'<[^>]+>', '', text).strip()
        indent_px = indent_level * 0.75  # 转换为像素
        return f'<div style="margin-left: {indent_px}px; border-left: 3px solid #409eff; padding-left: 10px; margin: 10px 0;">{text}</div>'
    
    # 处理字体大小
    size_match = re.search(r'<size:(\d+)>', text)
    if size_match:
        size = int(size_match.group(1))
        text = re.sub(r'<size:\d+>', '', text)
        text = re.sub(r'<[^>]+>', '', text).strip()
        return f'<p style="font-size: {size}px; margin: 12px 0; text-indent: 2em; line-height: 1.8;">{text}</p>'
    
    # 处理字体类型
    font_match = re.search(r'<font:([^>]+)>', text)
    if font_match:
        font = font_match.group(1)
        text = re.sub(r'<font:[^>]+>', '', text)
        text = re.sub(r'<[^>]+>', '', text).strip()
        return f'<p style="font-family: {font}; margin: 12px 0; text-indent: 2em; line-height: 1.8;">{text}</p>'
    
    # 处理段前间距
    space_before_match = re.search(r'<space_before:(\d+)>', text)
    if space_before_match:
        space = int(space_before_match.group(1))
        text = re.sub(r'<space_before:\d+>', '', text)
        text = re.sub(r'<[^>]+>', '', text).strip()
        return f'<p style="margin-top: {space * 0.75}px; margin-bottom: 12px; text-indent: 2em; line-height: 1.8;">{text}</p>'
    
    # 处理段后间距
    space_after_match = re.search(r'<space_after:(\d+)>', text)
    if space_after_match:
        space = int(space_after_match.group(1))
        text = re.sub(r'<space_after:\d+>', '', text)
        text = re.sub(r'<[^>]+>', '', text).strip()
        return f'<p style="margin-top: 12px; margin-bottom: {space * 0.75}px; text-indent: 2em; line-height: 1.8;">{text}</p>'
    
    # 处理行间距
    line_spacing_match = re.search(r'<line_spacing:([\d.]+)>', text)
    if line_spacing_match:
        spacing = float(line_spacing_match.group(1))
        text = re.sub(r'<line_spacing:[\d.]+>', '', text)
        text = re.sub(r'<[^>]+>', '', text).strip()
        return f'<p style="margin: 12px 0; text-indent: 2em; line-height: {spacing};">{text}</p>'
    
    # 移除其他格式标记
    text = re.sub(r'<[^>]+>', '', text)
    return f'<p style="margin: 12px 0; text-indent: 2em; line-height: 1.8;">{text}</p>'

def convert_format_tags_to_display(content: str) -> str:
    """
    将格式标签转换为实际展示效果
    """
    try:
        import re
        
        # 先处理加粗文本，避免与其他标签冲突
        content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
        
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            if not line.strip():
                formatted_lines.append('<br>')  # 空行用<br>表示
                continue
            
            # 处理标题 - 支持多个星号开头
            if line.startswith('*'):
                # 计算开头的星号数量
                star_count = 0
                for char in line:
                    if char == '*':
                        star_count += 1
                    else:
                        break
                
                # 如果开头有星号，提取标题文本
                if star_count > 0:
                    title_text = line[star_count:].strip()
                    # 移除格式标记
                    title_text = re.sub(r'<[^>]+>', '', title_text).strip()
                    if title_text:
                        # 根据星号数量确定标题级别
                        level = min(star_count, 6)  # 最多6级标题
                        formatted_lines.append(f'<h{level} style="margin: 20px 0 10px 0; font-weight: bold; color: #2c3e50;">{title_text}</h{level}>')
                        continue
            
            # 处理居中对齐
            if '<align:center>' in line:
                line = re.sub(r'<align:center>', '<div style="text-align: center;">', line)
                line = line + '</div>'
            
            # 处理右对齐
            if '<align:right>' in line:
                line = re.sub(r'<align:right>', '<div style="text-align: right;">', line)
                line = line + '</div>'
            
            # 处理缩进
            indent_match = re.search(r'<indent:(\d+)>', line)
            if indent_match:
                indent_level = int(indent_match.group(1))
                line = re.sub(r'<indent:\d+>', '', line)
                # 使用CSS padding-left而不是&nbsp;
                line = f'<div style="padding-left: {indent_level}px">{line}</div>'
            
            # 处理字体大小
            size_match = re.search(r'<size:(\d+)>', line)
            if size_match:
                size = size_match.group(1)
                line = re.sub(r'<size:\d+>', f'<span style="font-size: {size}px;">', line)
                line = line + '</span>'
            
            # 处理字体类型
            font_match = re.search(r'<font:([^>]+)>', line)
            if font_match:
                font = font_match.group(1)
                line = re.sub(r'<font:[^>]+>', f'<span style="font-family: {font};">', line)
                line = line + '</span>'
            
            # 处理段前间距
            if '<space_before:' in line:
                space_match = re.search(r'<space_before:(\d+)>', line)
                if space_match:
                    space = space_match.group(1)
                    line = re.sub(r'<space_before:\d+>', f'<div style="margin-top: {space}px;">', line)
                    line = line + '</div>'
            
            # 处理段后间距
            if '<space_after:' in line:
                space_match = re.search(r'<space_after:(\d+)>', line)
                if space_match:
                    space = space_match.group(1)
                    line = re.sub(r'<space_after:\d+>', f'<div style="margin-bottom: {space}px;">', line)
                    line = line + '</div>'
            
            # 处理行间距
            if '<line_spacing:' in line:
                spacing_match = re.search(r'<line_spacing:([\d.]+)>', line)
                if spacing_match:
                    spacing = spacing_match.group(1)
                    line = re.sub(r'<line_spacing:[\d.]+>', f'<div style="line-height: {spacing};">', line)
                    line = line + '</div>'
            
            # 处理加粗文本
            if '**' in line:
                line = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', line)
            
            # 处理下划线
            line = re.sub(r'<u>(.*?)</u>', r'<u>\1</u>', line)
            
            # 移除其他未处理的格式标签
            line = re.sub(r'<size:\d+>', '', line)
            line = re.sub(r'<font:[^>]+>', '', line)
            line = re.sub(r'<space_before:\d+>', '', line)
            line = re.sub(r'<space_after:\d+>', '', line)
            line = re.sub(r'<line_spacing:[\d.]+>', '', line)
            
            # 处理普通段落
            if line.strip():
                formatted_lines.append(f'<p style="margin: 12px 0; text-indent: 2em; line-height: 1.8;">{line}</p>')
            else:
                formatted_lines.append('<br>')
        
        return '\n'.join(formatted_lines)
        
    except Exception as e:
        logger.error(f"转换格式标签失败: {e}")
        return content
