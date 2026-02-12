"""
案例管理视图
"""
import os
import json
import logging
from django.conf import settings
from django.db import models
from django.shortcuts import get_object_or_404
from django.http import HttpResponse, FileResponse, JsonResponse
from django.views.decorators.clickjacking import xframe_options_exempt
from rest_framework import status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.filters import SearchFilter, OrderingFilter

# 设置日志
logger = logging.getLogger(__name__)

from dvadmin.utils.json_response import DetailResponse, ErrorResponse
from dvadmin.utils.request_util import get_request_user
from dvadmin.utils.viewset import CustomModelViewSet
from .models import CaseManagement, CaseDocument, DocumentTemplate, CaseFolder
from .serializers import (
    CaseManagementSerializer, 
    CaseDocumentSerializer, 
    DocumentTemplateSerializer,
    ManualGenerateSerializer,
    AIChatSerializer,
    PlaceholderEditSerializer,
    CaseFolderSerializer,
    CaseDocumentDetailSerializer,
    DocumentUploadSerializer,
    DocumentMoveSerializer,
    DocumentRenameSerializer,
    CheckGeneratedDocumentsSerializer,
    GenerateDocumentsSerializer
)
from .ai_service import generate_document_with_ai, generate_all_documents_with_ai, ai_chat_with_documents
from .smart_document_filler import smart_fill_document, smart_fill_all_templates, generate_smart_documents
from .xpert_integration import XpertAIClient
from .utils.document_converter import DocumentConverter
from .utils.document_file_manager import DocumentFileManager
from .utils.image_handler import ImageHandler


class CaseManagementViewSet(CustomModelViewSet):
    """案例管理视图集"""
    
    queryset = CaseManagement.objects.filter(is_deleted=False)
    serializer_class = CaseManagementSerializer
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['case_type', 'case_status', 'draft_person']
    search_fields = ['case_number', 'case_name', 'case_description', 'defendant_name', 'plaintiff_name']
    ordering_fields = ['id', 'case_number']
    ordering = ['-id']
    
    # 临时禁用权限检查，确保功能正常
    permission_classes = []
    # 临时禁用数据权限过滤器，因为案例管理不需要数据权限过滤
    extra_filter_class = []
    
    # 只返回JSON，不渲染HTML模板
    from rest_framework.renderers import JSONRenderer
    renderer_classes = [JSONRenderer]
    
    def get_queryset(self):
        """Get queryset with data scope"""
        queryset = super().get_queryset()
        user = get_request_user(self.request) or getattr(self.request, 'user', None)
        if not user or not getattr(user, 'is_authenticated', False):
            return queryset.none()
        role_level = getattr(user, 'role_level', None) or getattr(user, 'org_scope', None)
        if role_level == 'HQ':
            return queryset
        return queryset.filter(handlers=user)
    
    @action(detail=False, methods=['get'], url_path='home_dashboard')
    def home_dashboard(self, request):
        """首页仪表板数据
        
        返回：
        - stats: 统计数据（固定值）
        - upcoming_deadlines: 即将到期的案件（从案件中取8条）
        - ongoing_cases: 进行中的案件（从案件中取8条）
        """
        try:
            # 获取参数
            try:
                limit = int(request.query_params.get('limit', 8))
            except Exception:
                limit = 8
            if limit <= 0:
                limit = 8
            
            # 1. 即将到期的案件（根据立案日期或案例日期排序，取最近的）
            try:
                upcoming_cases = CaseManagement.objects.filter(
                    is_deleted=False
                ).filter(
                    models.Q(filing_date__isnull=False) | models.Q(case_date__isnull=False)
                ).order_by('-filing_date', '-case_date')[:limit]
            except Exception:
                upcoming_cases = CaseManagement.objects.filter(
                    is_deleted=False
                ).order_by('-id')[:limit]
            
            upcoming_deadlines = [
                {
                    'id': case.id,
                    'case_number': case.case_number,
                    'case_name': case.case_name,
                    'case_type': case.case_type,
                    'deadline': case.filing_date.strftime('%Y-%m-%d') if case.filing_date else (
                        case.case_date.strftime('%Y-%m-%d') if case.case_date else None
                    ),
                    'status': case.case_status,
                }
                for case in upcoming_cases
            ]
            
            # 2. 进行中的案件（case_status 包含"进行"或"处理"等关键词）
            try:
                ongoing_cases_queryset = CaseManagement.objects.filter(
                    is_deleted=False,
                    case_status__in=['进行中', '待处理', '审理中', '处理中']
                ).order_by('-id')[:limit]
            except Exception:
                ongoing_cases_queryset = CaseManagement.objects.filter(
                    is_deleted=False
                ).order_by('-id')[:limit]
            
            # 如果上面的状态筛选结果为空，则取最新的案件
            if not ongoing_cases_queryset.exists():
                ongoing_cases_queryset = CaseManagement.objects.filter(
                    is_deleted=False
                ).order_by('-id')[:limit]
            
            ongoing_cases = [
                {
                    'id': case.id,
                    'case_number': case.case_number,
                    'case_name': case.case_name,
                    'case_type': case.case_type,
                    'case_status': case.case_status,
                    'plaintiff_name': case.plaintiff_name,
                    'defendant_name': case.defendant_name,
                    'filing_date': case.filing_date.strftime('%Y-%m-%d') if case.filing_date else None,
                }
                for case in ongoing_cases_queryset
            ]
            
            # 3. 统计数据（固定返回）
            stats = {
                'ongoing_cases_count': 5,
                'upcoming_deadlines_count': 8,
                'documents_today_count': 234,
                'win_rate': 87,
                'ongoing_cases_change': '+12%',
                'documents_change': 23,
                'win_rate_change': '+5%'
            }
            
            return DetailResponse(
                data={
                    'stats': stats,
                    'upcoming_deadlines': upcoming_deadlines,
                    'ongoing_cases': ongoing_cases
                },
                msg='success'
            )
            
        except Exception as e:
            logger.error(f"获取首页仪表板数据失败: {str(e)}")
            return ErrorResponse(msg=f"获取数据失败: {str(e)}")
    
    def create(self, request, *args, **kwargs):
        """创建案例"""
        data = request.data.copy()
        # 支持多被告：将 defendants 数组拍扁为逗号分隔的四个字段，兼容现有模型字段
        try:
            defendants = data.get('defendants')
            if defendants:
                # 允许传 JSON 字符串或数组
                if isinstance(defendants, str):
                    try:
                        defendants = json.loads(defendants)
                    except Exception:
                        defendants = []
                if isinstance(defendants, list):
                    names = []
                    codes = []
                    addrs = []
                    reps = []
                    for d in defendants:
                        if not isinstance(d, dict):
                            continue
                        name = (d.get('name') or '').strip()
                        code = (d.get('credit_code') or '').strip()
                        addr = (d.get('address') or '').strip()
                        rep = (d.get('legal_representative') or '').strip()
                        if name or code or addr or rep:
                            names.append(name)
                            codes.append(code)
                            addrs.append(addr)
                            reps.append(rep)
                    if names or codes or addrs or reps:
                        data['defendant_name'] = ','.join([n for n in names if n])
                        data['defendant_credit_code'] = ','.join([c for c in codes if c])
                        data['defendant_address'] = ','.join([a for a in addrs if a])
                        data['defendant_legal_representative'] = ','.join([r for r in reps if r])
        except Exception as _e:
            logger.warning(f"创建案例时处理多被告数据出错: {_e}")

        serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        # 不传递created_by，让模型自动处理
        serializer.save()
        return DetailResponse(data=serializer.data, msg="案例创建成功")
    
    def update(self, request, *args, **kwargs):
        """更新案例"""
        instance = self.get_object()
        data = request.data.copy()
        # 支持多被告：将 defendants 数组拍扁为逗号分隔的四个字段，兼容现有模型字段
        try:
            defendants = data.get('defendants')
            if defendants:
                if isinstance(defendants, str):
                    try:
                        defendants = json.loads(defendants)
                    except Exception:
                        defendants = []
                if isinstance(defendants, list):
                    names = []
                    codes = []
                    addrs = []
                    reps = []
                    for d in defendants:
                        if not isinstance(d, dict):
                            continue
                        name = (d.get('name') or '').strip()
                        code = (d.get('credit_code') or '').strip()
                        addr = (d.get('address') or '').strip()
                        rep = (d.get('legal_representative') or '').strip()
                        if name or code or addr or rep:
                            names.append(name)
                            codes.append(code)
                            addrs.append(addr)
                            reps.append(rep)
                    if names or codes or addrs or reps:
                        data['defendant_name'] = ','.join([n for n in names if n])
                        data['defendant_credit_code'] = ','.join([c for c in codes if c])
                        data['defendant_address'] = ','.join([a for a in addrs if a])
                        data['defendant_legal_representative'] = ','.join([r for r in reps if r])
        except Exception as _e:
            logger.warning(f"更新案例时处理多被告数据出错: {_e}")

        serializer = self.get_serializer(instance, data=data, partial=True)
        serializer.is_valid(raise_exception=True)
        # 不传递updated_by，让模型自动处理
        serializer.save()
        return DetailResponse(data=serializer.data, msg="案例更新成功")
    
    def destroy(self, request, *args, **kwargs):
        """永久删除案件及其关联数据"""
        instance = self.get_object()
        
        try:
            # 删除关联的文档
            CaseDocument.objects.filter(case=instance).delete()
            
            # 删除关联的目录
            CaseFolder.objects.filter(case=instance).delete()
            
            # 永久删除案件
            instance.delete()
            
            return DetailResponse(msg="案件删除成功")
        except Exception as e:
            logger.error(f"删除案件失败: {str(e)}")
            return ErrorResponse(msg=f"删除案件失败: {str(e)}")
    
    @action(detail=True, methods=['post'], url_path='apply_filing')
    def apply_filing(self, request, pk=None):
        """申请立案 - 将草稿状态的案件申请立案，状态变更为已立案"""
        try:
            case = self.get_object()
            
            # 检查当前状态是否为草稿或待受理
            if case.case_status not in ["草稿", "待受理","待处理"]:
                return ErrorResponse(msg=f"当前案件状态为【{case.case_status}】，只有草稿/待受理/待处理状态的案件才能申请立案")
            
            # 更新状态为已立案
            case.case_status = "已立案"
            case.save()
            
            # 返回更新后的案件信息
            serializer = CaseManagementSerializer(case)
            return DetailResponse(
                data=serializer.data, 
                msg="申请立案成功，案件状态已变更为【已立案】"
            )
            
        except Exception as e:
            logger.error(f"申请立案失败: {str(e)}")
            return ErrorResponse(msg=f"申请立案失败: {str(e)}")
    
    @action(detail=True, methods=['get'])
    def documents(self, request, pk=None):
        """获取案例的文档列表"""
        case = self.get_object()
        documents = case.casedocument_set.filter(is_deleted=False)
        serializer = CaseDocumentSerializer(documents, many=True)
        return DetailResponse(data=serializer.data)
    
    @action(detail=True, methods=['post'])
    def save_case_info(self, request, pk=None):
        """保存案例信息"""
        try:
            case = self.get_object()
            
            # 获取请求数据
            data = request.data
            
            # 更新案例信息
            case.defendant_name = data.get('defendant_name', case.defendant_name)
            case.defendant_credit_code = data.get('defendant_credit_code', case.defendant_credit_code)
            case.defendant_address = data.get('defendant_address', case.defendant_address)
            case.defendant_legal_representative = data.get('defendant_legal_representative', case.defendant_legal_representative)
            case.plaintiff_name = data.get('plaintiff_name', case.plaintiff_name)
            case.plaintiff_credit_code = data.get('plaintiff_credit_code', case.plaintiff_credit_code)
            case.plaintiff_address = data.get('plaintiff_address', case.plaintiff_address)
            case.plaintiff_legal_representative = data.get('plaintiff_legal_representative', case.plaintiff_legal_representative)
            case.contract_amount = data.get('contract_amount', case.contract_amount)
            case.lawyer_fee = data.get('lawyer_fee', case.lawyer_fee)
            case.litigation_request = data.get('litigation_request', case.litigation_request)
            case.facts_and_reasons = data.get('facts_and_reasons', case.facts_and_reasons)
            case.jurisdiction = data.get('jurisdiction', case.jurisdiction)
            case.petitioner = data.get('petitioner', case.petitioner)
            
            # 处理日期字段
            if data.get('filing_date'):
                from datetime import datetime
                logger.info(f"处理日期字段: {data['filing_date']}, 类型: {type(data['filing_date'])}")
                try:
                    # 如果前端传递的是字符串，需要解析
                    if isinstance(data['filing_date'], str):
                        case.filing_date = datetime.strptime(data['filing_date'], '%Y-%m-%d').date()
                        logger.info(f"解析后的日期: {case.filing_date}")
                    else:
                        case.filing_date = data['filing_date']
                        logger.info(f"直接使用日期: {case.filing_date}")
                except ValueError as e:
                    logger.error(f"日期解析失败: {data['filing_date']}, 错误: {e}")
            else:
                logger.info("没有提供日期字段")
            
            # 保存案例
            case.save()
            
            # 返回更新后的案例信息
            serializer = CaseManagementSerializer(case)
            return DetailResponse(data=serializer.data, msg="案例信息保存成功")
            
        except Exception as e:
            logger.error(f"保存案例信息失败: {str(e)}")
            return ErrorResponse(msg=f"保存失败: {str(e)}")
    
    @action(detail=True, methods=['post'])
    def generate_manual_document(self, request, pk=None):
        """人工录入生成文档 - 直接使用字段名匹配占位符"""
        case = self.get_object()
        serializer = ManualGenerateSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        try:
            # 获取所有可用的模板
            from .models import DocumentTemplate
            templates = DocumentTemplate.objects.filter(is_active=True, is_deleted=False)
            if not templates.exists():
                return DetailResponse(
                    data={'documents': [], 'summary': {'success': 0, 'failed': 0, 'total': 0}},
                    msg="没有启用的模板可用于生成文书"
                )
            
            # 准备表单数据
            form_data = serializer.validated_data
            logger.info(f"序列化器验证后的数据: {form_data}")
            
            # 将Decimal字段转换为float，避免JSON序列化错误
            for key, value in form_data.items():
                if hasattr(value, 'as_tuple'):  # 检查是否为Decimal类型
                    form_data[key] = float(value)
            
            logger.info(f"开始生成文档，表单数据: {form_data}")
            logger.info(f"表单数据中的关键字段 - litigation_request: {form_data.get('litigation_request')}, facts_and_reasons: {form_data.get('facts_and_reasons')}, jurisdiction: {form_data.get('jurisdiction')}, petitioner: {form_data.get('petitioner')}, filing_date: {form_data.get('filing_date')}")
            
            # 生成文档
            created_documents = []
            success_count = 0
            error_count = 0
            total_count = templates.count()
            
            for template in templates:
                try:
                    # 直接使用字段名作为占位符进行匹配
                    template_data = {}
                    
                    if template.placeholder_info and 'placeholders' in template.placeholder_info:
                        logger.info(f"模板 {template.template_name} 的占位符: {template.placeholder_info['placeholders']}")
                        for placeholder in template.placeholder_info['placeholders']:
                            placeholder_key = placeholder['key']
                            logger.info(f"处理占位符: {placeholder_key}")
                            
                            # 直接使用占位符key作为字段名进行匹配
                            if placeholder_key in form_data:
                                value = form_data[placeholder_key]
                                logger.info(f"找到匹配字段 {placeholder_key}: {value}")
                                
                                # 特殊处理：对contract_amount进行法律语言格式化
                                if placeholder_key == 'contract_amount' and value:
                                    try:
                                        amount = float(value)
                                        value = f"请求判令被告支付合同款项人民币{amount:,.2f}元"
                                    except (ValueError, TypeError):
                                        value = f"请求判令被告支付合同款项人民币{value}元"
                                
                                template_data[placeholder_key] = value
                            else:
                                # 如果字段不存在，使用默认值
                                logger.warning(f"字段 {placeholder_key} 在表单数据中不存在，使用默认值")
                                template_data[placeholder_key] = '待填写'
                    
                    # 使用占位符服务生成文档
                    from .placeholder_template_service import placeholder_service
                    doc = placeholder_service.fill_and_save_by_record(case.id, template, template_data, request=request)
                    created_documents.append(CaseDocumentSerializer(doc).data)
                    success_count += 1
                    logger.info(f"成功生成文档: {template.template_name}")
                    
                except Exception as e:
                    error_count += 1
                    logger.error(f"生成文档失败 - 模板ID: {template.id}, 模板名称: {template.template_name}, 错误: {e}")
                    import traceback
                    logger.error(f"详细错误信息: {traceback.format_exc()}")
            
            return DetailResponse(
                data={
                    'documents': created_documents,
                    'summary': {
                        'success': success_count,
                        'failed': error_count,
                        'total': total_count
                    }
                },
                msg=f"人工录入文书生成完成！成功: {success_count}个，失败: {error_count}个"
            )
            
        except Exception as e:
            logger.error(f"人工录入生成文档失败: {e}")
            import traceback
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            return DetailResponse(
                data={'documents': [], 'summary': {'success': 0, 'failed': 0, 'total': 0}},
                msg=f"人工录入生成文档失败: {str(e)}"
            )
    
    
    @action(detail=True, methods=['post'])
    def generate_ai_document(self, request, pk=None):
        """AI生成文档"""
        case = self.get_object()
        serializer = AIChatSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        try:
            # 调用AI对话生成文档
            result = ai_chat_with_documents(
                serializer.validated_data['message'],
                serializer.validated_data.get('files', [])
            )
            
            if result['success']:
                # 保存生成的文档
                document = CaseDocument.objects.create(
                    case=case,
                    document_name=f"AI生成文档_{case.case_number}_{case.case_name}",
                    document_content=result['content'],
                    document_type='AI生成',
                    file_path='',  # 提供空字符串作为默认值
                    generation_method='ai',
                    template_used='AI生成'
                )
                
                return DetailResponse(data=CaseDocumentSerializer(document).data, msg="AI文档生成成功")
            else:
                return Response({"error": result['error']}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['post'])
    def generate_all_documents(self, request, pk=None):
        """生成所有文档（根据11个模板）"""
        try:
            case = self.get_object()
            
            # 准备案例数据
            case_data = {
                'case_number': case.case_number,
                'case_name': case.case_name,
                'case_type': case.case_type,
                'jurisdiction': case.jurisdiction,
                'draft_person': case.draft_person,
                'defendant_name': case.defendant_name,
                'defendant_credit_code': case.defendant_credit_code,
                'defendant_address': case.defendant_address,
                'plaintiff_name': case.plaintiff_name,
                'plaintiff_credit_code': case.plaintiff_credit_code,
                'plaintiff_address': case.plaintiff_address,
                'contract_amount': float(case.contract_amount) if case.contract_amount else 0.0,
                'lawyer_fee': float(case.lawyer_fee) if case.lawyer_fee else 0.0
            }
            
            # 调用AI生成所有文档
            result = generate_all_documents_with_ai(case_data)
            
            if result['success']:
                # 保存所有生成的文档
                created_documents = []
                for doc_data in result['documents']:
                    if doc_data['success']:
                        # 检查是否已存在相同类型的文档
                        existing_doc = CaseDocument.objects.filter(
                            case=case,
                            document_type=doc_data['document_name']
                        ).first()
                        
                        if existing_doc:
                            # 更新现有文档
                            existing_doc.document_content = doc_data['content']
                            existing_doc.template_used = doc_data['template_name']
                            existing_doc.save()
                            created_documents.append(CaseDocumentSerializer(existing_doc).data)
                        else:
                            # 创建新文档
                            document = CaseDocument.objects.create(
                                case=case,
                                document_name=doc_data['template_name'],  # 使用模板名称作为文档名称
                                document_content=doc_data['content'],
                                document_type=doc_data['document_name'],
                                file_path='',  # 提供空字符串作为默认值
                                generation_method='manual',
                                template_used=doc_data['template_name']
                            )
                            created_documents.append(CaseDocumentSerializer(document).data)
                
                return DetailResponse(
                    data={
                        'documents': created_documents,
                        'summary': {
                            'success': result['success_count'],
                            'failed': result['error_count'],
                            'total': result['total_count']
                        }
                    }, 
                    msg=f"成功生成 {result['success_count']} 个文档，失败 {result['error_count']} 个"
                )
            else:
                error_msg = result.get('error', '未知错误')
                print(f"文档生成失败: {error_msg}")
                return Response({"error": f"文档生成失败: {error_msg}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            print(f"生成所有文档异常: {e}")
            import traceback
            traceback.print_exc()
            return Response({"error": f"生成所有文档异常: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=False, methods=['post'])
    def ai_chat(self, request):
        """AI对话并生成文书"""
        serializer = AIChatSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        try:
            # 获取案例ID（从请求中获取）
            case_id = serializer.validated_data.get('case_id')
            case = None
            
            if case_id:
                # 获取案例信息
                try:
                    case = CaseManagement.objects.get(id=case_id)
                except CaseManagement.DoesNotExist:
                    return Response({"error": "案例不存在"}, status=status.HTTP_404_NOT_FOUND)
            
            # 准备案例数据
            if case:
                case_data = {
                    'case_number': case.case_number,
                    'case_name': case.case_name,
                    'case_type': case.case_type,
                    'jurisdiction': case.jurisdiction,
                    'draft_person': case.draft_person,
                    'defendant_name': case.defendant_name,
                    'defendant_credit_code': case.defendant_credit_code,
                    'defendant_address': case.defendant_address,
                    'plaintiff_name': case.plaintiff_name,
                    'plaintiff_credit_code': case.plaintiff_credit_code,
                    'plaintiff_address': case.plaintiff_address,
                    'contract_amount': float(case.contract_amount) if case.contract_amount else 0.0,
                    'lawyer_fee': float(case.lawyer_fee) if case.lawyer_fee else 0.0
                }
            else:
                # 如果没有案例，使用默认数据
                case_data = {
                    'case_number': '临时案例',
                    'case_name': '新建案例',
                    'case_type': '民事纠纷',
                    'jurisdiction': '北京市朝阳区人民法院',
                    'draft_person': '系统',
                    'defendant_name': '待填写',
                    'defendant_credit_code': '待填写',
                    'defendant_address': '待填写',
                    'plaintiff_name': '待填写',
                    'plaintiff_credit_code': '待填写',
                    'plaintiff_address': '待填写',
                    'contract_amount': 0,
                    'lawyer_fee': 0
                }
            
            # 调用AI对话
            result = ai_chat_with_documents(
                serializer.validated_data['message'],
                serializer.validated_data.get('files', [])
            )
            
            if result['success']:
                # 如果AI建议生成文书，则生成所有文书
                # 检查用户消息是否包含文书相关的关键词
                user_message_lower = serializer.validated_data['message'].lower()
                if any(keyword in user_message_lower for keyword in ['生成文书', '文书', '起诉状', '委托书', '申请书', '合同', '法律文书', '文档', '帮我写', '需要', '要']):
                    # 生成所有文书
                    documents_result = generate_all_documents_with_ai(case_data)
                    
                    if documents_result.get('success'):
                        if case:
                            # 如果有案例，保存生成的文档到数据库
                            created_documents = []
                            for doc_data in documents_result['documents']:
                                if doc_data['success']:
                                    # 检查是否已存在相同类型的文档
                                    existing_doc = CaseDocument.objects.filter(
                                        case=case,
                                        document_type=doc_data['document_name']
                                    ).first()
                                    
                                    if existing_doc:
                                        # 更新现有文档
                                        existing_doc.document_content = doc_data['content']
                                        existing_doc.template_used = doc_data['template_name']
                                        existing_doc.generation_method = 'ai'
                                        existing_doc.save()
                                        created_documents.append(CaseDocumentSerializer(existing_doc).data)
                                    else:
                                        # 创建新文档
                                        document = CaseDocument.objects.create(
                                            case=case,
                                            document_name=doc_data['template_name'],
                                            document_content=doc_data['content'],
                                            document_type=doc_data['document_name'],
                                            file_path='',
                                            generation_method='ai',
                                            template_used=doc_data['template_name']
                                        )
                                        created_documents.append(CaseDocumentSerializer(document).data)
                            
                            return DetailResponse(
                                data={
                                    'response': result['content'],
                                    'documents': created_documents,
                                    'generated_count': len(created_documents)
                                }, 
                                msg="AI回复成功，已生成文书"
                            )
                        else:
                            # 如果没有案例，只返回生成的文档内容
                            generated_documents = []
                            for doc_data in documents_result['documents']:
                                if doc_data['success']:
                                    generated_documents.append({
                                        'template_name': doc_data['template_name'],
                                        'document_name': doc_data['document_name'],
                                        'content': doc_data['content']
                                    })
                            
                            return DetailResponse(
                                data={
                                    'response': result['content'],
                                    'documents': generated_documents,
                                    'generated_count': len(generated_documents)
                                }, 
                                msg="AI回复成功，已生成文书预览"
                            )
                
                return DetailResponse(
                    data={
                        'response': result['content'],
                        'documents': [],
                        'generated_count': 0
                    }, 
                    msg="AI回复成功"
                )
            else:
                return Response({"error": result['error']}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['post'])
    def smart_fill_document(self, request, pk=None):
        """智能填充文档（基于模板）"""
        case = self.get_object()
        
        try:
            # 获取请求参数
            template_path = request.data.get('template_path')
            output_path = request.data.get('output_path')
            use_xml = request.data.get('use_xml', False)
            
            if not template_path or not output_path:
                return Response({"error": "缺少必要参数：template_path 和 output_path"}, 
                              status=status.HTTP_400_BAD_REQUEST)
            
            # 准备案例数据
            case_data = {
                'case_number': case.case_number,
                'case_name': case.case_name,
                'case_type': case.case_type,
                'jurisdiction': case.jurisdiction,
                'draft_person': case.draft_person,
                'defendant_name': case.defendant_name,
                'defendant_credit_code': case.defendant_credit_code,
                'defendant_address': case.defendant_address,
                'defendant_legal_representative': case.defendant_legal_representative,
                'plaintiff_name': case.plaintiff_name,
                'plaintiff_credit_code': case.plaintiff_credit_code,
                'plaintiff_address': case.plaintiff_address,
                'plaintiff_legal_representative': case.plaintiff_legal_representative,
                'contract_amount': float(case.contract_amount) if case.contract_amount else 0.0,
                'lawyer_fee': float(case.lawyer_fee) if case.lawyer_fee else 0.0
            }
            
            # 调用智能填充功能
            result = smart_fill_document(template_path, case_data, output_path, use_xml)
            
            if result['success']:
                return DetailResponse(data=result, msg="智能填充成功")
            else:
                return Response({"error": result.get('error', '智能填充失败')}, 
                              status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['post'])
    def smart_fill_all_templates(self, request, pk=None):
        """智能填充所有模板"""
        case = self.get_object()
        
        try:
            # 获取输出目录
            output_dir = request.data.get('output_dir', 'smart_output')
            
            # 准备案例数据
            case_data = {
                'case_number': case.case_number,
                'case_name': case.case_name,
                'case_type': case.case_type,
                'jurisdiction': case.jurisdiction,
                'draft_person': case.draft_person,
                'defendant_name': case.defendant_name,
                'defendant_credit_code': case.defendant_credit_code,
                'defendant_address': case.defendant_address,
                'defendant_legal_representative': case.defendant_legal_representative,
                'plaintiff_name': case.plaintiff_name,
                'plaintiff_credit_code': case.plaintiff_credit_code,
                'plaintiff_address': case.plaintiff_address,
                'plaintiff_legal_representative': case.plaintiff_legal_representative,
                'contract_amount': float(case.contract_amount) if case.contract_amount else 0.0,
                'lawyer_fee': float(case.lawyer_fee) if case.lawyer_fee else 0.0
            }
            
            # 调用智能填充所有模板功能
            result = smart_fill_all_templates(case_data, output_dir)
            
            if result['success']:
                return DetailResponse(data=result, msg="智能填充所有模板成功")
            else:
                return Response({"error": result.get('error', '智能填充所有模板失败')}, 
                              status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=True, methods=['post'])
    def generate_smart_documents(self, request, pk=None):
        """智能生成文档（结合AI生成和模板填充）"""
        case = self.get_object()
        
        try:
            # 获取输出目录
            output_dir = request.data.get('output_dir', 'smart_output')
            
            # 准备案例数据
            case_data = {
                'case_number': case.case_number,
                'case_name': case.case_name,
                'case_type': case.case_type,
                'jurisdiction': case.jurisdiction,
                'draft_person': case.draft_person,
                'defendant_name': case.defendant_name,
                'defendant_credit_code': case.defendant_credit_code,
                'defendant_address': case.defendant_address,
                'defendant_legal_representative': case.defendant_legal_representative,
                'plaintiff_name': case.plaintiff_name,
                'plaintiff_credit_code': case.plaintiff_credit_code,
                'plaintiff_address': case.plaintiff_address,
                'plaintiff_legal_representative': case.plaintiff_legal_representative,
                'contract_amount': float(case.contract_amount) if case.contract_amount else 0.0,
                'lawyer_fee': float(case.lawyer_fee) if case.lawyer_fee else 0.0
            }
            
            # 调用智能生成文档功能
            result = generate_smart_documents(case_data, output_dir)
            
            if result['success']:
                # 保存AI生成的文档到数据库
                created_documents = []
                for doc in result.get('documents', []):
                    if doc.get('type') == 'ai_generated' and doc.get('success'):
                        # 检查是否已存在相同类型的文档
                        existing_doc = CaseDocument.objects.filter(
                            case=case,
                            document_type=doc.get('name', '')
                        ).first()
                        
                        if existing_doc:
                            # 更新现有文档
                            existing_doc.document_content = doc.get('content', '')
                            existing_doc.template_used = doc.get('template_used', '')
                            existing_doc.generation_method = 'smart_ai'
                            existing_doc.save()
                            created_documents.append(CaseDocumentSerializer(existing_doc).data)
                        else:
                            # 创建新文档
                            document = CaseDocument.objects.create(
                                case=case,
                                document_name=doc.get('name', ''),
                                document_content=doc.get('content', ''),
                                document_type=doc.get('name', ''),
                                file_path='',
                                generation_method='smart_ai',
                                template_used=doc.get('template_used', '')
                            )
                            created_documents.append(CaseDocumentSerializer(document).data)
                
                return DetailResponse(
                    data={
                        'result': result,
                        'created_documents': created_documents,
                        'ai_generated_count': result.get('ai_generated_count', 0),
                        'template_filled_count': result.get('template_filled_count', 0)
                    }, 
                    msg="智能生成文档成功"
                )
            else:
                return Response({"error": result.get('error', '智能生成文档失败')}, 
                              status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=False, methods=['post'])
    def extract_content(self, request):
        """提取上传文件的内容"""
        try:
            files = request.FILES.getlist('files')
            if not files:
                return DetailResponse(
                    data={},
                    msg="请上传文件"
                )
            
            # 使用Unstructured服务提取内容
            from .unstructured_document_service import unstructured_service
            
            all_content = []
            for file in files:
                try:
                    # 保存临时文件
                    import tempfile
                    import os
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.name)[1]) as tmp_file:
                        for chunk in file.chunks():
                            tmp_file.write(chunk)
                        tmp_file_path = tmp_file.name
                    
                    # 提取内容
                    content = unstructured_service.parse_template_file(tmp_file_path)
                    all_content.append(f"文件: {file.name}\n内容:\n{content}\n\n")
                    
                    # 清理临时文件
                    os.unlink(tmp_file_path)
                    
                except Exception as e:
                    logger.error(f"提取文件 {file.name} 内容失败: {e}")
                    all_content.append(f"文件: {file.name}\n提取失败: {str(e)}\n\n")
            
            return DetailResponse(
                data={
                    'content': ''.join(all_content)
                },
                msg="内容提取成功"
            )
            
        except Exception as e:
            logger.error(f"提取内容失败: {e}")
            return DetailResponse(
                data={},
                msg=f"提取内容失败: {str(e)}"
            )
    
    @action(detail=True, methods=['post'])
    def generate_from_files(self, request, pk=None):
        """基于上传文件生成文档"""
        try:
            case = self.get_object()
            template_id = request.data.get('template_id')
            files = request.FILES.getlist('files')
            
            if not template_id:
                return DetailResponse(
                    data={},
                    msg="请选择模板"
                )
            
            if not files:
                return DetailResponse(
                    data={},
                    msg="请上传文件"
                )
            
            # 获取模板
            try:
                template = DocumentTemplate.objects.get(id=template_id, is_deleted=False)
            except DocumentTemplate.DoesNotExist:
                return DetailResponse(
                    data={},
                    msg="模板不存在"
                )
            
            # 提取文件内容
            from .unstructured_document_service import unstructured_service
            from .smart_content_extractor import smart_extractor
            
            all_content = []
            for file in files:
                try:
                    # 保存临时文件
                    import tempfile
                    import os
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.name)[1]) as tmp_file:
                        for chunk in file.chunks():
                            tmp_file.write(chunk)
                        tmp_file_path = tmp_file.name
                    
                    # 提取内容
                    content = unstructured_service.parse_template_file(tmp_file_path)
                    all_content.append(content)
                    
                    # 清理临时文件
                    os.unlink(tmp_file_path)
                    
                except Exception as e:
                    logger.error(f"提取文件 {file.name} 内容失败: {e}")
                    all_content.append(f"文件 {file.name} 提取失败: {str(e)}")
            
            # 合并所有提取的内容
            combined_content = '\n\n'.join(all_content)
            
            # 使用智能提取器从内容中提取字段信息
            extracted_data = smart_extractor.extract_from_content(combined_content)
            
            # 准备案例数据
            case_data = {
                'case_id': case.id,
                'case_number': case.case_number,
                'case_name': case.case_name,
                'case_type': case.case_type,
                'jurisdiction': case.jurisdiction,
                'draft_person': case.draft_person,
                'defendant_name': case.defendant_name,
                'defendant_credit_code': case.defendant_credit_code,
                'defendant_address': case.defendant_address,
                'defendant_legal_representative': case.defendant_legal_representative,
                'plaintiff_name': case.plaintiff_name,
                'plaintiff_credit_code': case.plaintiff_credit_code,
                'plaintiff_address': case.plaintiff_address,
                'plaintiff_legal_representative': case.plaintiff_legal_representative,
                'contract_amount': float(case.contract_amount) if case.contract_amount else 0.0,
                'lawyer_fee': float(case.lawyer_fee) if case.lawyer_fee else 0.0,
                'extracted_content': combined_content
            }
            
            # 将智能提取的数据与案例数据合并
            case_data = smart_extractor.merge_with_case_data(case_data, extracted_data)
            
            # 使用AI进一步增强数据填充
            from .ai_template_filler import ai_template_filler
            case_data = ai_template_filler.enhance_with_ai(combined_content, case_data)
            
            # 使用占位符服务生成文档
            from .placeholder_template_service import placeholder_service
            
            doc = placeholder_service.fill_and_save_by_record(case.id, template, case_data, request=request)
            
            return DetailResponse(
                data=CaseDocumentSerializer(doc).data,
                msg="文档生成成功"
            )
            
        except Exception as e:
            logger.error(f"基于文件生成文档失败: {e}")
            return DetailResponse(
                data={},
                msg=f"生成文档失败: {str(e)}"
            )
    
    @action(detail=True, methods=['get'])
    def documents(self, request, pk=None):
        """获取案例的所有文书"""
        try:
            case = self.get_object()
            documents = CaseDocument.objects.filter(case=case, is_deleted=False)
            
            serializer = CaseDocumentSerializer(documents, many=True)
            return DetailResponse(
                data=serializer.data,
                msg="获取文书列表成功"
            )
        except Exception as e:
            return DetailResponse(
                data=[],
                msg=f"获取文书列表失败: {str(e)}"
            )
    
    @action(detail=False, methods=['post'], url_path='expert_analyze')
    def expert_analyze(self, request):
        """AI专家解析文档"""
        try:
            case_id = request.data.get('case_id')
            files = request.data.get('files', [])
            templates = request.data.get('templates', [])
            
            # 添加调试日志
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"接收到的文件数量: {len(files)}")
            logger.info(f"接收到的模板数量: {len(templates)}")
            if files:
                logger.info(f"第一个文件信息: name={files[0].get('name', '未知')}, type={files[0].get('type', '未知')}, has_base64={bool(files[0].get('base64'))}")
            
            if not files:
                return DetailResponse(msg="请上传数据源文档")
            
            if not templates:
                return DetailResponse(msg="没有可用的模板")
            
            # 初始化XpertAI客户端
            xpert_client = XpertAIClient(
                api_url=os.getenv("XPERTAI_API_URL", "https://api.mtda.cloud/api/ai/"),
                api_key=os.getenv("XPERTAI_API_KEY", "")
            )
            
            # 准备发送给专家的数据
            expert_data = {
                "files": files,
                "templates": templates,
                "case_id": case_id
            }
            
            # 调用XpertAI专家进行解析
            import asyncio
            analysis_result = asyncio.run(xpert_client.analyze_documents_with_expert(expert_data))
            
            # 解析专家返回的JSON数据
            parsed_data = {}
            if isinstance(analysis_result, dict) and 'parsed_data' in analysis_result:
                parsed_data = analysis_result['parsed_data']
            elif isinstance(analysis_result, str):
                try:
                    import json
                    parsed_data = json.loads(analysis_result)
                except:
                    parsed_data = {"raw_result": analysis_result}
            else:
                parsed_data = analysis_result
            
            # 返回解析结果
            result = {
                "original_result": analysis_result,
                "parsed_data": parsed_data,
                "matched_placeholders": analysis_result.get('matched_placeholders', []) if isinstance(analysis_result, dict) else [],
                "matched_templates": analysis_result.get('matched_templates', []) if isinstance(analysis_result, dict) else []
            }
            
            return DetailResponse(data=result, msg="AI专家解析完成")
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"AI专家解析失败: {e}")
            return DetailResponse(msg=f"AI专家解析失败: {str(e)}")
    
    @action(detail=True, methods=['post'], url_path='expert_generate')
    def expert_generate(self, request, pk=None):
        """基于占位符信息生成文书"""
        try:
            case_id = request.data.get('case_id')
            analysis_result = request.data.get('analysis_result')
            templates = request.data.get('templates', [])
            
            if not analysis_result:
                return DetailResponse(msg="请先进行AI专家解析")
            
            if not templates:
                return DetailResponse(msg="没有可用的模板")
            
            # 直接调用本地文书生成方法，不调用XpertAI平台
            generation_result = self._generate_documents_from_placeholders(
                case_id=case_id,
                analysis_result=analysis_result,
                templates=templates
            )
            
            # 将生成的文书保存到数据库（参考人工录入逻辑）
            if generation_result.get('documents'):
                saved_documents = self._save_generated_documents_with_templates(pk, generation_result['documents'], analysis_result)
                generation_result['saved_documents'] = saved_documents
            
            return DetailResponse(data=generation_result, msg="文书生成完成")
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"文书生成失败: {e}")
            return DetailResponse(msg=f"文书生成失败: {str(e)}")
    
    def _generate_documents_from_placeholders(self, case_id: str, analysis_result: dict, templates: list) -> dict:
        """基于占位符信息直接生成文书（本地方法，不调用XpertAI）"""
        try:
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"开始为案例 {case_id} 生成文书")
            
            generated_documents = []
            
            # 从分析结果中提取占位符信息
            matched_placeholders = analysis_result.get('matched_placeholders', [])
            matched_templates = analysis_result.get('matched_templates', [])
            
            if not matched_placeholders:
                logger.warning("没有找到匹配的占位符信息")
                return {
                    'documents': [],
                    'total_documents': 0,
                    'message': '没有找到匹配的占位符信息'
                }
            
            # 按模板分组占位符
            placeholders_by_template = {}
            for placeholder in matched_placeholders:
                template_id = placeholder.get('template_id', 'unknown')
                if template_id not in placeholders_by_template:
                    placeholders_by_template[template_id] = []
                placeholders_by_template[template_id].append(placeholder)
            
            # 打印调试信息
            logger.info(f"可用的占位符模板ID: {list(placeholders_by_template.keys())}")
            logger.info(f"传入的模板: {templates}")
            
            # 为每个模板生成文档
            for template in templates:
                # 处理模板数据，支持字典和整数两种格式
                if isinstance(template, dict):
                    template_id = str(template.get('template_id', ''))
                    template_name = template.get('template_name', '未知模板')
                elif isinstance(template, (int, str)):
                    # 如果传入的是模板ID，需要从分析结果中获取模板信息
                    template_id = str(template)
                    template_name = f"模板{template_id}"
                else:
                    logger.warning(f"不支持的模板格式: {type(template)}")
                    continue
                
                logger.info(f"处理模板ID: {template_id}")
                template_placeholders = placeholders_by_template.get(template_id, [])
                
                if not template_placeholders:
                    logger.warning(f"模板 {template_name} 没有匹配的占位符")
                    continue
                
                # 生成文档内容
                document_content = self._generate_document_content(template, template_placeholders)
                
                generated_documents.append({
                    'document_name': f"{template_name}_{case_id}",
                    'document_type': 'word',
                    'content': document_content,
                    'template_id': template_id,
                    'template_name': template_name,
                    'generation_method': 'placeholder_fill',
                    'placeholders_count': len(template_placeholders)
                })
                
                logger.info(f"为模板 {template_name} 生成了文档，包含 {len(template_placeholders)} 个占位符")
            
            result = {
                'documents': generated_documents,
                'total_documents': len(generated_documents),
                'message': f'成功生成 {len(generated_documents)} 个文档'
            }
            
            logger.info(f"文书生成完成，生成了 {len(generated_documents)} 个文档")
            return result
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"文书生成失败: {e}")
            return {
                'documents': [],
                'total_documents': 0,
                'error': str(e),
                'message': f'文书生成失败: {str(e)}'
            }
    
    def _generate_document_content(self, template, placeholders: list) -> str:
        """生成文档内容"""
        try:
            # 处理模板名称
            if isinstance(template, dict):
                template_name = template.get('template_name', '未知模板')
            else:
                template_name = f"模板{template}"
            
            content_lines = [
                f"# {template_name}",
                "",
                "基于以下占位符信息生成：",
                ""
            ]
            
            for placeholder in placeholders:
                key = placeholder.get('key', '')
                value = placeholder.get('value', '')
                content_lines.append(f"**{key}**: {value}")
                content_lines.append("")
            
            return "\n".join(content_lines)
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"生成文档内容失败: {e}")
            return f"文档生成失败: {str(e)}"
    
    def _save_generated_documents(self, case_id: int, documents: list) -> list:
        """保存生成的文书到数据库"""
        try:
            import logging
            logger = logging.getLogger(__name__)
            
            # 获取案例对象
            case = CaseManagement.objects.get(pk=int(case_id))
            saved_documents = []
            
            for doc_data in documents:
                # 创建文书记录
                case_document = CaseDocument.objects.create(
                    case=case,
                    document_name=doc_data.get('document_name', '未知文书'),
                    document_type=doc_data.get('document_type', 'word'),
                    document_content=doc_data.get('content', ''),
                    generation_method='ai',
                    template_used=doc_data.get('template_name', ''),
                    generation_params={
                        'template_id': doc_data.get('template_id'),
                        'placeholders_count': doc_data.get('placeholders_count', 0),
                        'generation_method': doc_data.get('generation_method', 'placeholder_fill')
                    }
                )
                
                saved_documents.append({
                    'id': case_document.id,
                    'document_name': case_document.document_name,
                    'document_type': case_document.document_type,
                    'template_used': case_document.template_used,
                    'generation_method': case_document.generation_method,
                    'create_datetime': case_document.create_datetime.strftime('%Y-%m-%d %H:%M:%S')
                })
                
                logger.info(f"保存文书成功: {case_document.document_name} (ID: {case_document.id})")
            
            logger.info(f"成功保存 {len(saved_documents)} 个文书到数据库")
            return saved_documents
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"保存文书失败: {e}")
            return []
    
    def _save_generated_documents_with_templates(self, case_id: int, documents: list, analysis_result: dict) -> list:
        """参考人工录入逻辑保存生成的文书到数据库"""
        try:
            import logging
            logger = logging.getLogger(__name__)
            
            # 获取案例对象
            case = CaseManagement.objects.get(pk=int(case_id))
            saved_documents = []
            
            # 导入占位符服务
            from .placeholder_template_service import placeholder_service
            from .models import DocumentTemplate
            
            # 获取所有启用的模板
            templates = DocumentTemplate.objects.filter(is_active=True, is_deleted=False)
            if not templates.exists():
                logger.warning("没有启用的模板可用于生成文书")
                return []
            
            # 构建符合占位符模板要求的数据结构
            case_data = self._build_case_data_for_templates(case, analysis_result)
            
            # 为每个生成的文档找到对应的模板并生成实际文档
            for doc_data in documents:
                template_id = doc_data.get('template_id')
                template_name = doc_data.get('template_name', '')
                
                # 查找对应的模板
                template = None
                if template_id:
                    try:
                        template = templates.get(id=int(template_id))
                    except (DocumentTemplate.DoesNotExist, ValueError):
                        logger.warning(f"未找到模板ID {template_id}，尝试按名称查找")
                
                if not template and template_name:
                    try:
                        template = templates.get(template_name__icontains=template_name)
                    except DocumentTemplate.DoesNotExist:
                        logger.warning(f"未找到模板名称 {template_name}")
                
                if not template:
                    # 使用第一个可用模板作为默认
                    template = templates.first()
                    logger.info(f"使用默认模板: {template.template_name}")
                
                if template:
                    try:
                        # 使用占位符服务生成实际文档
                        case_document = placeholder_service.fill_and_save_by_record(
                            case_id=int(case_id),
                            template_record=template,
                            data=case_data,
                            request=request
                        )
                        
                        saved_documents.append({
                            'id': case_document.id,
                            'document_name': case_document.document_name,
                            'document_type': case_document.document_type,
                            'template_used': template.template_name,
                            'generation_method': 'ai',
                            'file_path': case_document.file_path,
                            'file_size': case_document.file_size,
                            'create_datetime': case_document.create_datetime.strftime('%Y-%m-%d %H:%M:%S')
                        })
                        
                        logger.info(f"成功生成文书: {case_document.document_name} (ID: {case_document.id})")
                        
                    except Exception as e:
                        logger.error(f"生成文书失败 - 模板: {template.template_name}, 错误: {e}")
                        continue
                else:
                    logger.warning(f"没有可用的模板用于生成文书: {template_name}")
            
            logger.info(f"成功生成 {len(saved_documents)} 个文书")
            return saved_documents
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"保存文书失败: {e}")
            return []
    
    def _build_case_data_for_templates(self, case: 'CaseManagement', analysis_result: dict) -> dict:
        """构建符合占位符模板要求的数据结构"""
        try:
            # 从分析结果中提取占位符信息
            matched_placeholders = analysis_result.get('matched_placeholders', [])
            
            # 构建扁平化的数据字典
            case_data = {
                'case_id': case.id,
                'case_number': case.case_number,
                'case_name': case.case_name,
                'case_type': case.case_type,
                'jurisdiction': case.jurisdiction,
                'draft_person': case.draft_person,
                
                # 原告信息
                'plaintiff_name': case.plaintiff_name or '待填写',
                'plaintiff_address': case.plaintiff_address or '待填写',
                'plaintiff_credit_code': case.plaintiff_credit_code or '待填写',
                'plaintiff_legal_representative': case.plaintiff_legal_representative or '待填写',
                
                # 被告信息
                'defendant_name': case.defendant_name or '待填写',
                'defendant_address': case.defendant_address or '待填写',
                'defendant_credit_code': case.defendant_credit_code or '待填写',
                'defendant_legal_representative': case.defendant_legal_representative or '待填写',
                
                # 金额信息
                'contract_amount': float(case.contract_amount or 0.0),
                'lawyer_fee': float(case.lawyer_fee or 0.0),
                
                # 其他信息
                'case_description': case.case_description or '待填写',
            }
            
            # 添加嵌套结构以支持点号路径
            case_data.update({
                'plaintiff': {
                    'name': case.plaintiff_name or '待填写',
                    'address': case.plaintiff_address or '待填写',
                    'credit_code': case.plaintiff_credit_code or '待填写',
                    'legal_representative': case.plaintiff_legal_representative or '待填写'
                },
                'defendant': {
                    'name': case.defendant_name or '待填写',
                    'address': case.defendant_address or '待填写',
                    'credit_code': case.defendant_credit_code or '待填写',
                    'legal_representative': case.defendant_legal_representative or '待填写'
                }
            })
            
            # 将AI解析的占位符信息添加到数据中
            for placeholder in matched_placeholders:
                key = placeholder.get('key', '')
                value = placeholder.get('value', '')
                if key and value:
                    case_data[key] = value
            
            return case_data
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"构建案例数据失败: {e}")
            return {}
    
    @action(detail=True, methods=['post'], url_path='check_generated_documents')
    def check_generated_documents(self, request, pk=None):
        """检查指定模板是否已生成文书
        
        ✅ 简化逻辑：只根据 case_id 和 template_id 判断
        """
        try:
            case = self.get_object()
            serializer = CheckGeneratedDocumentsSerializer(data=request.data)
            serializer.is_valid(raise_exception=True)
            
            template_ids = serializer.validated_data['template_ids']
            
            # ✅ 简化查询：只检查 case_id 和 template_id
            existing_documents = CaseDocument.objects.filter(
                case=case,
                template_id__in=template_ids,
                is_deleted=False
            ).select_related('template')
            
            # ✅ 构建返回数据
            exists = existing_documents.exists()
            existing_doc_names = []
            existing_doc_list = []
            
            # 记录已生成的模板ID
            generated_template_ids = set()
            
            for doc in existing_documents:
                doc_name = f"{doc.document_name}{doc.file_ext or ''}"
                existing_doc_names.append(doc_name)
                existing_doc_list.append({
                    'id': doc.id,
                    'document_name': doc_name,
                    'template_id': doc.template_id,
                    'template_name': doc.template.template_name if doc.template else '',
                    'file_path': doc.file_path or '',
                    'file_size': doc.file_size or 0,
                    'created_at': doc.create_datetime.strftime('%Y-%m-%d %H:%M:%S') if doc.create_datetime else None
                })
                if doc.template_id:
                    generated_template_ids.add(doc.template_id)
            
            # ✅ 找出未生成的模板ID（在请求列表中但数据库没有的）
            missing_template_ids = [
                tid for tid in template_ids 
                if tid not in generated_template_ids
            ]
            
            return DetailResponse(
                data={
                    'exists': exists,
                    'existing_documents': existing_doc_names,
                    'existing_documents_detail': existing_doc_list,
                    'total_count': len(existing_doc_list),
                    'missing_template_ids': missing_template_ids,  # ✅ 新增：未生成的模板ID列表
                    'generated_count': len(existing_doc_list),      # ✅ 新增：已生成的数量
                    'missing_count': len(missing_template_ids)      # ✅ 新增：未生成的数量
                },
                msg="检查完成"
            )
            
        except Exception as e:
            logger.error(f"检查已生成文书失败: {str(e)}")
            return ErrorResponse(msg=f"检查失败: {str(e)}")
    
    @action(detail=True, methods=['get'], url_path='generated_template_ids')
    def generated_template_ids(self, request, pk=None):
        """获取案件已生成的模板ID列表"""
        try:
            case = self.get_object()
            
            # 查询该案件所有已生成的文档（未删除的）
            generated_documents = CaseDocument.objects.filter(
                case=case,
                is_deleted=False,
                template_id__isnull=False
            ).values_list('template_id', flat=True).distinct()
            
            # 转换为列表并去重
            template_ids = list(set(generated_documents))
            
            return DetailResponse(
                data={
                    'template_ids': template_ids,
                    'count': len(template_ids)
                },
                msg="获取已生成模板ID列表成功"
            )
            
        except Exception as e:
            logger.error(f"获取已生成模板ID列表失败: {str(e)}")
            return ErrorResponse(msg=f"获取失败: {str(e)}")
    
    @action(detail=True, methods=['get'], url_path='document_tree')
    def document_tree(self, request, pk=None):
        """获取案件的完整目录树结构，如果为空则自动创建预设目录"""
        try:
            from .utils.folder_helper import get_case_document_tree, create_case_folders
            
            case = self.get_object()
            
            # 检查该案件是否已有目录
            folder_count = CaseFolder.objects.filter(case=case, is_deleted=False).count()
            
            # 如果目录为空，自动创建预设目录
            if folder_count == 0:
                logger.info(f"案件 {case.id} 目录为空，自动创建预设目录")
                create_case_folders(case.id)
            
            # 获取并返回目录树
            tree = get_case_document_tree(case.id)
            
            message = "success" if folder_count > 0 else "目录已自动创建"
            return DetailResponse(data=tree, msg=message)
            
        except Exception as e:
            logger.error(f"获取文档树失败: {str(e)}")
            return ErrorResponse(msg=f"获取文档树失败: {str(e)}")
    
    @action(detail=True, methods=['get'])
    def folders(self, request, pk=None):
        """获取案件的目录结构"""
        case = self.get_object()
        folders = CaseFolder.objects.filter(case=case, is_deleted=False).order_by('sort_order')
        serializer = CaseFolderSerializer(folders, many=True)
        return DetailResponse(data=serializer.data)
    
    @action(detail=True, methods=['get'])
    def folder_documents(self, request, pk=None):
        """获取指定目录下的文档列表"""
        case = self.get_object()
        folder_path = request.query_params.get('folder_path', '/')
        
        documents = CaseDocument.objects.filter(
            case=case,
            folder_path=folder_path,
            is_deleted=False
        ).select_related('template', 'folder')
        
        serializer = CaseDocumentDetailSerializer(documents, many=True)
        return DetailResponse(data=serializer.data)
    
    @action(detail=True, methods=['post'])
    def upload_document(self, request, pk=None):
        """上传文档到指定目录"""
        case = self.get_object()
        serializer = DocumentUploadSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        try:
            from .utils.folder_helper import save_uploaded_file
            
            folder_path = serializer.validated_data['folder_path']
            file = serializer.validated_data['file']
            document_name = serializer.validated_data.get('document_name', file.name)
            
            # 保存文件并创建文档记录
            document = save_uploaded_file(
                case_id=case.id,
                folder_path=folder_path,
                file=file,
                document_name=document_name
            )
            
            return DetailResponse(
                data=CaseDocumentDetailSerializer(document).data,
                msg="文档上传成功"
            )
        except Exception as e:
            logger.error(f"上传文档失败: {str(e)}")
            return ErrorResponse(msg=f"上传文档失败: {str(e)}")


class CaseDocumentViewSet(CustomModelViewSet):
    """案例文档视图集"""
    
    queryset = CaseDocument.objects.filter(is_deleted=False)
    serializer_class = CaseDocumentSerializer
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['document_type', 'generation_method', 'case']
    search_fields = ['document_name', 'template_used']
    ordering_fields = ['id']
    ordering = ['-id']
    permission_classes = []
    extra_filter_class = []  # 移除可能有问题的过滤器
    
    # 只返回JSON，不渲染HTML模板
    from rest_framework.renderers import JSONRenderer
    renderer_classes = [JSONRenderer]
    
    def destroy(self, request, *args, **kwargs):
        """删除文档 - 直接删除记录并删除物理文件"""
        try:
            document = self.get_object()
            
            # 获取文件的完整物理路径
            file_path = document.full_file_path
            
            # 删除物理文件
            if file_path and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    logger.info(f"已删除文件: {file_path}")
                except Exception as e:
                    logger.error(f"删除文件失败: {file_path}, 错误: {str(e)}")
            
            # 真实删除数据库记录（不是软删除）
            document.delete(soft_delete=False)
            
            return DetailResponse(msg="文档删除成功")
            
        except Exception as e:
            logger.error(f"删除文档失败: {str(e)}")
            return ErrorResponse(msg=f"删除文档失败: {str(e)}")
    
    @action(detail=True, methods=['get'], permission_classes=[])
    def download(self, request, pk=None):
        """下载文档"""
        try:
            logger.info(f"下载文档请求 - pk: {pk}, 类型: {type(pk)}")
            
            # 检查pk是否为有效数字
            if pk is None or pk == 'undefined' or not str(pk).isdigit():
                logger.error(f"无效的文档ID: {pk}")
                return DetailResponse(msg="无效的文档ID")
            
            document = self.get_object()
            logger.info(f"找到文档: {document.id} - {document.document_name}")
        except:
            # 如果get_object失败，直接通过pk获取
            try:
                document = CaseDocument.objects.get(pk=pk)
            except CaseDocument.DoesNotExist:
                return DetailResponse(msg="文档不存在")
            except ValueError:
                return DetailResponse(msg="无效的文档ID")
        
        # 获取文档文件的完整路径
        doc_file_path = document.full_file_path
        if not doc_file_path or not os.path.exists(doc_file_path):
            return Response(
                {
                    "error": "文件不存在",
                    "message": "文档文件不存在或已被删除",
                    "code": "FILE_NOT_FOUND"
                },
                status=status.HTTP_404_NOT_FOUND
            )
        
        # 根据文档类型确定Content-Type
        import mimetypes
        file_ext = os.path.splitext(doc_file_path)[1].lower()
        content_type, _ = mimetypes.guess_type(doc_file_path)
        
        if content_type is None:
            if file_ext == '.docx':
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif file_ext == '.doc':
                content_type = 'application/msword'
            elif file_ext == '.pdf':
                content_type = 'application/pdf'
            elif file_ext in ['.txt', '.md']:
                content_type = 'text/plain; charset=utf-8'
            else:
                content_type = 'application/octet-stream'
        
        try:
            # 使用 FileResponse 流式传输文件（更高效，支持大文件）
            from django.http import FileResponse
            response = FileResponse(
                open(doc_file_path, 'rb'),
                content_type=content_type,
                as_attachment=True
            )
            
            # 设置文件名 - 支持中文文件名
            filename = document.document_name
            if not filename.endswith(file_ext):
                filename += file_ext
            
            # 使用标准的 UTF-8 编码方式设置文件名
            import urllib.parse
            encoded_filename = urllib.parse.quote(filename)
            response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{encoded_filename}'
            
            # 添加支持 blob 下载的响应头
            response['Access-Control-Expose-Headers'] = 'Content-Disposition'
            
            return response
            
        except Exception as e:
            logger.error(f"文件下载失败: {str(e)}")
            return Response(
                {
                    "error": "下载失败",
                    "message": "文件下载失败，请稍后重试",
                    "detail": str(e),
                    "code": "DOWNLOAD_ERROR"
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'], permission_classes=[])
    @xframe_options_exempt
    def public_download(self, request, pk=None):
        """
        公开下载接口，用于WPS等第三方预览服务
        
        重要：必须直接返回文件内容（二进制流），不能返回JSON
        注意：使用 @xframe_options_exempt 装饰器，确保 Django 中间件不会添加 X-Frame-Options 响应头
        """
        try:
            # 检查pk是否为有效数字
            if pk is None or pk == 'undefined' or not str(pk).isdigit():
                # 返回JSON错误（但应该尽量避免）
                return JsonResponse({
                    "code": 40004,
                    "message": "invalid document id"
                }, status=400)
            
            document = self.get_object()
        except:
            # 如果get_object失败，直接通过pk获取
            try:
                document = CaseDocument.objects.get(pk=pk, is_deleted=False)
            except CaseDocument.DoesNotExist:
                return JsonResponse({
                    "code": 40004,
                    "message": "file not exists"
                }, status=404)
            except ValueError:
                return JsonResponse({
                    "code": 40004,
                    "message": "invalid document id"
                }, status=400)
        
        try:
            import urllib.parse
            import base64
            
            # 创建自定义响应类，确保X-Frame-Options不会被设置
            class FrameAllowedHttpResponse(HttpResponse):
                """允许iframe加载的HTTP响应类，阻止 X-Frame-Options 响应头"""
                def __init__(self, *args, **kwargs):
                    super().__init__(*args, **kwargs)
                    # 设置为 True，告诉 Django 的 XFrameOptionsMiddleware 不要添加 X-Frame-Options 头部
                    self.xframe_options_exempt = True
                
                def __setitem__(self, key, value):
                    # 完全阻止设置 X-Frame-Options 响应头
                    if key.lower() == 'x-frame-options':
                        logger.warning(f"阻止设置 X-Frame-Options 响应头: {value}")
                        return
                    super().__setitem__(key, value)
                
                def __delitem__(self, key):
                    # 允许删除 X-Frame-Options 响应头
                    super().__delitem__(key)
            
            class FrameAllowedFileResponse(FileResponse):
                """允许iframe加载的文件响应类，阻止 X-Frame-Options 响应头"""
                def __init__(self, *args, **kwargs):
                    # 移除 as_attachment 参数（如果存在），确保不会自动设置 attachment
                    kwargs.pop('as_attachment', None)
                    super().__init__(*args, **kwargs)
                    # 设置为 True，告诉 Django 的 XFrameOptionsMiddleware 不要添加 X-Frame-Options 头部
                    self.xframe_options_exempt = True
                
                def __setitem__(self, key, value):
                    # 完全阻止设置 X-Frame-Options 响应头
                    if key.lower() == 'x-frame-options':
                        logger.warning(f"阻止设置 X-Frame-Options 响应头: {value}")
                        return
                    super().__setitem__(key, value)
                
                def __delitem__(self, key):
                    # 允许删除 X-Frame-Options 响应头
                    super().__delitem__(key)
                
                def _set_content_disposition(self, value):
                    """确保 Content-Disposition 响应头被正确设置"""
                    # 直接设置响应头，确保不被覆盖
                    if hasattr(self, '_headers'):
                        self._headers['content-disposition'] = ('Content-Disposition', value)
                    else:
                        super().__setitem__('Content-Disposition', value)
            
            # 优先从文件路径读取文件
            file_path = document.full_file_path
            if file_path and os.path.exists(file_path):
                # 从文件路径读取
                try:
                    file_handle = open(file_path, 'rb')
                    filename = document.file_name or document.document_name or 'document.docx'
                    # 必须设置 Content-Disposition: inline，允许在 iframe 中内联显示文档
                    content_disposition = f'inline; filename="{filename}"'
                    
                    response = FrameAllowedFileResponse(
                        file_handle,
                        content_type=document.mime_type or 
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    # 必须设置 Content-Disposition: inline，允许在 iframe 中内联显示文档
                    # 使用直接设置方式，确保响应头不被覆盖
                    if hasattr(response, '_set_content_disposition'):
                        response._set_content_disposition(content_disposition)
                    else:
                        response['Content-Disposition'] = content_disposition
                        # 确保响应头被设置（防止被覆盖）
                        if hasattr(response, '_headers'):
                            response._headers['content-disposition'] = ('Content-Disposition', content_disposition)
                    response['Access-Control-Allow-Origin'] = '*'
                    response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
                    response['Access-Control-Allow-Headers'] = 'Range'
                    # 支持Range请求（断点续传）
                    if os.path.getsize(file_path) > 0:
                        response['Accept-Ranges'] = 'bytes'
                    
                    # 确保响应头已设置（调试用）
                    # 直接设置 _headers 字典，确保响应头不被覆盖
                    if hasattr(response, '_headers'):
                        response._headers['content-disposition'] = ('Content-Disposition', content_disposition)
                    # 双重检查：确保响应头存在
                    final_content_disposition = response.get('Content-Disposition', 'NOT SET')
                    logger.info(f"公开下载文档（文件路径）: {document.id} - {document.document_name}, Content-Disposition: {final_content_disposition}")
                    
                    # 如果响应头仍未设置，强制设置
                    if final_content_disposition == 'NOT SET' or 'inline' not in final_content_disposition:
                        if hasattr(response, '_headers'):
                            response._headers['content-disposition'] = ('Content-Disposition', content_disposition)
                        else:
                            response['Content-Disposition'] = content_disposition
                    
                    # 明确删除 X-Frame-Options 响应头（如果存在）
                    if hasattr(response, '_headers') and 'x-frame-options' in response._headers:
                        del response._headers['x-frame-options']
                    # 确保 xframe_options_exempt 属性被设置
                    response.xframe_options_exempt = True
                    
                    return response
                except IOError as e:
                    logger.error(f"读取文件失败: {str(e)}")
                    # 继续尝试从document_content读取
            
            # 如果文件路径不存在，尝试从document_content读取（base64）
            if document.document_content:
                try:
                    docx_bytes = base64.b64decode(document.document_content)
                    filename = f"{document.document_name}.docx"
                    ascii_filename = urllib.parse.quote(filename.encode('utf-8'))
                    # 必须设置 Content-Disposition: inline，允许在 iframe 中内联显示文档
                    content_disposition = f'inline; filename*=UTF-8\'\'{ascii_filename}'
                    
                    response = FrameAllowedHttpResponse(
                        docx_bytes,
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    # 必须设置 Content-Disposition: inline，允许在 iframe 中内联显示文档
                    response['Content-Disposition'] = content_disposition
                    response['Access-Control-Allow-Origin'] = '*'
                    response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
                    response['Access-Control-Allow-Headers'] = 'Range'
                    
                    # 确保响应头已设置（调试用）
                    # 直接设置 _headers 字典，确保响应头不被覆盖
                    if hasattr(response, '_headers'):
                        response._headers['content-disposition'] = ('Content-Disposition', content_disposition)
                    # 双重检查：确保响应头存在
                    final_content_disposition = response.get('Content-Disposition', 'NOT SET')
                    logger.info(f"公开下载文档（Base64）: {document.id} - {document.document_name}, Content-Disposition: {final_content_disposition}")
                    
                    # 如果响应头仍未设置，强制设置
                    if final_content_disposition == 'NOT SET' or 'inline' not in final_content_disposition:
                        if hasattr(response, '_headers'):
                            response._headers['content-disposition'] = ('Content-Disposition', content_disposition)
                        else:
                            response['Content-Disposition'] = content_disposition
                    
                    # 明确删除 X-Frame-Options 响应头（如果存在）
                    if hasattr(response, '_headers') and 'x-frame-options' in response._headers:
                        del response._headers['x-frame-options']
                    # 确保 xframe_options_exempt 属性被设置
                    response.xframe_options_exempt = True
                    
                    return response
                except Exception as e:
                    logger.error(f"Base64解码失败: {str(e)}")
                    return JsonResponse({
                        "code": 40004,
                        "message": "document content format error"
                    }, status=400)
            
            # 文件不存在且内容为空
            return JsonResponse({
                "code": 40004,
                "message": "file not exists"
            }, status=404)
            
        except Exception as e:
            logger.error(f"公开下载文档失败: {str(e)}", exc_info=True)
            # 错误响应也不能设置X-Frame-Options: deny
            error_response = JsonResponse({
                "code": 50000,
                "message": str(e)
            }, status=500)
            error_response.xframe_options_exempt = True
            return error_response
    
    @action(detail=True, methods=['get'])
    def template_placeholders(self, request, pk=None):
        """获取当前文档所使用模板的占位符信息"""
        try:
            # 获取文档
            document = self.get_object()
            
            # 获取案件信息并序列化为字典（仅基础字段）
            from django.forms.models import model_to_dict
            case_instance = getattr(document, 'case', None)
            case_data_all = {}
            if case_instance:
                # 仅取模型的本地字段，避免外键/多对多
                case_field_names = [f.name for f in case_instance._meta.fields]
                case_data_all = model_to_dict(case_instance, fields=case_field_names)
            
            # 校验是否关联模板
            template = document.template
            if not template:
                return DetailResponse(
                    data={
                        'document_id': document.id,
                        'template_id': None,
                        'placeholder_info': {},
                        'data': {}
                    },
                    msg="该文档未关联模板，无法获取占位符信息"
                )
            
            def build_filtered_case_data(placeholder_info_obj: dict) -> dict:
                """根据占位符key集合过滤案件字段，只返回有占位符的字段"""
                try:
                    placeholders = placeholder_info_obj.get('placeholders', []) if isinstance(placeholder_info_obj, dict) else []
                    placeholder_keys = {p.get('key') for p in placeholders if isinstance(p, dict) and p.get('key')}
                except Exception:
                    placeholder_keys = set()
                if not placeholder_keys:
                    return {}
                filtered = {}
                for key in placeholder_keys:
                    if key in case_data_all:
                        value = case_data_all.get(key)
                        # 统一转成基础可序列化类型
                        filtered[key] = "" if value is None else value
                return filtered
            
            # 如果模板已缓存占位符信息，基于缓存生成结果
            if template.placeholder_info:
                filtered_case_data = build_filtered_case_data(template.placeholder_info)
                return DetailResponse(
                    data={
                        'document_id': document.id,
                        'template_id': template.id,
                        'template_name': template.template_name,
                        'placeholder_info': template.placeholder_info,
                        'data': filtered_case_data
                    },
                    msg="获取占位符信息成功（缓存）"
                )
            
            # 未缓存则解析模板文件并缓存
            from .placeholder_parser import placeholder_parser
            template_file_path = template.full_file_path
            placeholder_info = placeholder_parser.parse_template_file(template_file_path)
            # 修复中文转义问题
            placeholder_info_fixed = placeholder_parser.fix_unicode_escaped_data(placeholder_info)
            # 缓存保存
            template.placeholder_info = placeholder_info_fixed
            template.save()
            # 基于解析结果过滤案件字段
            filtered_case_data = build_filtered_case_data(placeholder_info_fixed)
            
            return DetailResponse(
                data={
                    'document_id': document.id,
                    'template_id': template.id,
                    'template_name': template.template_name,
                    'placeholder_info': placeholder_info_fixed,
                    'data': filtered_case_data
                },
                msg="解析占位符信息成功"
            )
        except Exception as e:
            logger.error(f"获取模板占位符失败: {str(e)}")
            return ErrorResponse(msg=f"获取模板占位符失败: {str(e)}")
    
    @action(detail=True, methods=['get'])
    def preview(self, request, pk=None):
        """预览文档内容"""
        try:
            # 检查pk是否为有效数字
            if pk is None or pk == 'undefined' or not str(pk).isdigit():
                return DetailResponse(msg="无效的文档ID")
            
            document = self.get_object()
        except:
            # 如果get_object失败，直接通过pk获取
            try:
                document = CaseDocument.objects.get(pk=pk)
            except CaseDocument.DoesNotExist:
                return DetailResponse(msg="文档不存在")
            except ValueError:
                return DetailResponse(msg="无效的文档ID")
        
        # 获取预览类型参数
        preview_type = request.GET.get('type', 'html')  # 默认为html预览
        
        if preview_type == 'third_party':
            # 第三方预览服务
            return self._get_third_party_preview(document, request)
        else:
            # 原有的HTML预览
            return self._get_html_preview(document)
    
    def _get_html_preview(self, document):
        """获取HTML格式预览"""
        # 根据文档类型选择合适的格式转换函数
        from .direct_langchain_ai_service import convert_format_tags_to_display, convert_docx_format_to_html
        
        # 检查文档内容是否包含Word格式标记
        if any(tag in document.document_content for tag in ['<align:', '<indent:', '<size:', '<font:', '<space_', '<line_spacing:']):
            # 使用Word格式转换函数
            html_content = convert_docx_format_to_html(document.document_content)
        else:
            # 使用通用格式转换函数
            html_content = convert_format_tags_to_display(document.document_content)
        
        return DetailResponse(
            data={
                'id': document.id,  # 添加文档ID
                'document_name': document.document_name,
                'document_type': document.document_type,
                'content': html_content,  # 使用HTML格式
                'raw_content': document.document_content,  # 保留原始内容
                'create_datetime': document.create_datetime,
                'generation_method': document.generation_method,
                'template_used': document.template_used,
                'preview_type': 'html'
            },
            msg="文档预览成功"
        )
    
    def _get_third_party_preview(self, document, request):
        """获取第三方预览服务URL"""
        try:
            import urllib.parse
            from django.urls import reverse
            from django.conf import settings
            
            # 构建文档公开下载URL
            # 使用相对URL构建，避免URL名称问题
            download_url = f"{request.scheme}://{request.get_host()}/api/case/documents/{document.id}/public_download/"
            
            # URL编码，确保特殊字符正确处理
            encoded_download_url = urllib.parse.quote(download_url, safe='')
            
            # 生成第三方预览服务URL
            third_party_urls = {
                'microsoft': f"https://view.officeapps.live.com/op/embed.aspx?src={encoded_download_url}",
                'google': f"https://docs.google.com/gview?url={encoded_download_url}&embedded=true",
                'mozilla': f"https://mozilla.github.io/pdf.js/web/viewer.html?file={encoded_download_url}",
                # 添加更多可靠的预览服务
                'office365': f"https://view.officeapps.live.com/op/view.aspx?src={encoded_download_url}",
                'pdfjs': f"https://mozilla.github.io/pdf.js/web/viewer.html?file={encoded_download_url}",
                'viewerjs': f"https://viewerjs.org/{encoded_download_url}"
            }
            
            logger.info(f"生成第三方预览URL成功: {download_url}")
            
            return DetailResponse(
                data={
                    'id': document.id,
                    'document_name': document.document_name,
                    'document_type': document.document_type,
                    'download_url': download_url,
                    'preview_urls': third_party_urls,
                    'preview_type': 'third_party',
                    'create_datetime': document.create_datetime,
                    'generation_method': document.generation_method,
                    'template_used': document.template_used
                },
                msg="第三方预览URL生成成功"
            )
            
        except Exception as e:
            logger.error(f"生成第三方预览URL失败: {str(e)}")
            logger.error(f"错误详情: {type(e).__name__}: {str(e)}")
            return ErrorResponse(msg=f"生成预览URL失败: {str(e)}")
    
    @action(detail=False, methods=['post'], url_path='batch-update-print-count')
    def batch_update_print_count(self, request):
        """批量更新文档打印数量
        
        请求格式:
        {
            "documents": [
                {"id": 1, "print_count": 3},
                {"id": 2, "print_count": 2}
            ]
        }
        """
        try:
            # 验证请求数据
            from .serializers import BatchUpdatePrintCountSerializer
            serializer = BatchUpdatePrintCountSerializer(data=request.data)
            if not serializer.is_valid():
                return ErrorResponse(msg=f"数据验证失败: {serializer.errors}")
            
            documents_data = serializer.validated_data['documents']
            
            # 使用事务批量更新
            from django.db import transaction
            
            with transaction.atomic():
                # 提取所有文档ID
                document_ids = [item['id'] for item in documents_data]
                
                # 查询对应的文档对象（只查询未删除的）
                documents = CaseDocument.objects.filter(
                    id__in=document_ids,
                    is_deleted=False
                )
                
                # 检查是否所有文档都存在
                found_ids = set(d.id for d in documents)
                missing_ids = set(document_ids) - found_ids
                if missing_ids:
                    return ErrorResponse(
                        msg=f"部分文档不存在",
                        data={'not_found_ids': list(missing_ids)}
                    )
                
                # 检查请求中是否包含is_selected字段（用于向后兼容）
                has_is_selected = any('is_selected' in item for item in documents_data)
                
                # 创建ID到数据的映射
                document_data_map = {
                    item['id']: {
                        'print_count': item['print_count'],
                        'is_selected': item.get('is_selected', False) if has_is_selected else None
                        # 如果请求中没有is_selected字段，设置为None，表示不更新
                    }
                    for item in documents_data
                }
                
                # 更新每个文档的print_count和is_selected
                for document in documents:
                    data = document_data_map[document.id]
                    document.print_count = data['print_count']
                    # 只有当前端提供了is_selected字段时才更新（向后兼容）
                    if has_is_selected and data['is_selected'] is not None:
                        document.is_selected = data['is_selected']
                
                # 批量保存（包含is_selected字段）
                update_fields = ['print_count']
                # 如果请求中包含is_selected字段，也一起更新
                if has_is_selected:
                    update_fields.append('is_selected')
                CaseDocument.objects.bulk_update(documents, update_fields)
                
                # 准备返回数据
                updated_documents = [
                    {
                        'id': doc.id,
                        'document_name': doc.document_name,
                        'print_count': doc.print_count,
                        'is_selected': doc.is_selected
                    }
                    for doc in documents
                ]
            
            logger.info(
                f"成功更新 {len(documents)} 个文档的打印数量和选中状态: "
                f"has_is_selected={has_is_selected}, "
                f"updated_fields={update_fields}"
            )
            
            return DetailResponse(
                data={
                    'updated_count': len(documents),
                    'updated_documents': updated_documents
                },
                msg=f"成功更新 {len(documents)} 个文档的打印数量和选中状态"
            )
            
        except Exception as e:
            logger.error(f"批量更新打印数量失败: {str(e)}")
            return ErrorResponse(msg=f"更新打印数量失败: {str(e)}")
    
    @action(detail=True, methods=['post'], url_path='wps/edit-config', permission_classes=[])
    def wps_edit_config(self, request, pk=None):
        """获取WPS编辑配置
        
        请求参数：
        {
            "mode": "edit",  // "view" 或 "edit"
            "userId": 456,  // 可选
            "userName": "张三"  // 可选
        }
        """
        try:
            document = self.get_object()
            mode = request.data.get('mode', 'edit')  # 'view' 或 'edit'
            user_id = request.data.get('userId') or get_request_user(request).id if hasattr(get_request_user(request), 'id') else None
            user_name = request.data.get('userName') or (get_request_user(request).username if hasattr(get_request_user(request), 'username') else '用户')
            
            # 构建文档文件URL（使用public_download接口，支持跨域）
            file_url = f"{request.scheme}://{request.get_host()}/api/case/documents/{document.id}/public_download/"
            
            # 构建回调URL和保存URL
            callback_url = f"{request.scheme}://{request.get_host()}/api/case/documents/{document.id}/wps/callback/"
            save_url = f"{request.scheme}://{request.get_host()}/api/case/documents/{document.id}/wps/save/"
            download_url = f"{request.scheme}://{request.get_host()}/api/case/documents/{document.id}/wps/download/"
            
            # 生成简单的Token（实际项目中应该使用JWT或更安全的算法）
            import hashlib
            import time
            token_data = f"{document.id}_{user_id}_{int(time.time())}"
            token = hashlib.sha256(token_data.encode()).hexdigest()
            
            # WPS配置（需要从环境变量或配置中获取，这里先用占位符）
            # CDN 地址配置（支持环境变量配置，默认使用官方 CDN）
            wps_cdn_url = os.getenv('WPS_CDN_URL', 'https://wwo.wps.cn/office/v1/index.js')
            
            wps_config = {
                'fileUrl': file_url,
                'fileId': str(document.id),
                'appId': os.getenv('WPS_APP_ID', 'wps_app_id_placeholder'),  # 需要配置
                'token': token,
                'mode': mode,
                'userId': str(user_id) if user_id else '0',
                'userName': user_name or '用户',
                'callbackUrl': callback_url,
                'saveUrl': save_url,
                'downloadUrl': download_url,
                'cdnUrl': wps_cdn_url,  # ✅ 新增：CDN 地址，前端可以使用此地址加载 SDK
            }
            
            logger.info(f"WPS配置生成成功: document_id={document.id}, mode={mode}, user_id={user_id}")
            
            return DetailResponse(
                data={
                    'documentId': document.id,
                    'wpsConfig': wps_config
                },
                msg="WPS配置生成成功"
            )
            
        except Exception as e:
            logger.error(f"生成WPS配置失败: {str(e)}")
            import traceback
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            return ErrorResponse(msg=f"生成WPS配置失败: {str(e)}")
    
    @action(detail=True, methods=['post'], url_path='wps/save', permission_classes=[])
    def wps_save(self, request, pk=None):
        """保存WPS编辑后的文档
        
        请求参数：
        Content-Type: multipart/form-data
        file: [文档文件]
        """
        try:
            document = self.get_object()
            file = request.FILES.get('file')
            
            if not file:
                return ErrorResponse(msg="未上传文件")
            
            # 验证文件格式
            if not file.name.lower().endswith('.docx'):
                return ErrorResponse(msg="仅支持 .docx 格式文件")
            
            # 验证文件大小（限制50MB）
            max_size = 50 * 1024 * 1024  # 50MB
            if file.size > max_size:
                return ErrorResponse(msg=f"文件大小超过限制（最大50MB），当前文件大小：{file.size / 1024 / 1024:.2f}MB")
            
            # 备份原文件（如果存在）
            original_file_path = document.full_file_path
            if original_file_path and os.path.exists(original_file_path):
                try:
                    # 创建备份目录
                    backup_dir = os.path.join(os.path.dirname(original_file_path), 'backup')
                    os.makedirs(backup_dir, exist_ok=True)
                    
                    # 备份文件（带时间戳）
                    import datetime
                    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                    backup_filename = f"{document.id}_{timestamp}{os.path.splitext(original_file_path)[1]}"
                    backup_path = os.path.join(backup_dir, backup_filename)
                    
                    import shutil
                    shutil.copy2(original_file_path, backup_path)
                    logger.info(f"已备份原文件: {backup_path}")
                except Exception as e:
                    logger.warning(f"备份文件失败: {str(e)}")
            
            # 保存新文件到原位置（如果原文件存在）或新位置
            original_file_path = document.full_file_path
            
            if original_file_path and os.path.exists(original_file_path):
                # 保存到原位置（覆盖原文件）
                saved_full_path = original_file_path
                saved_relative_path = document.file_path  # 保持原路径
            else:
                # 原文件不存在，保存到新位置
                from django.conf import settings
                # 构建文件路径（与 save_uploaded_file 保持一致）
                base_path = os.path.join(settings.MEDIA_ROOT, 'cases', str(document.case.id))
                folder_path = document.folder_path or '/case_documents'
                folder_rel_path = folder_path.strip('/')
                folder_full_path = os.path.join(base_path, folder_rel_path)
                
                # 确保目录存在
                os.makedirs(folder_full_path, exist_ok=True)
                
                # 生成文件名（使用原文件名或生成新文件名）
                import datetime
                timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                if document.file_name:
                    filename = f"{os.path.splitext(document.file_name)[0]}_{timestamp}.docx"
                else:
                    filename = f"{document.document_name}_{timestamp}.docx"
                
                saved_full_path = os.path.join(folder_full_path, filename)
                
                # 构建相对路径（用于数据库存储）
                saved_relative_path = os.path.join('cases', str(document.case.id), folder_rel_path, filename).replace('\\', '/')
            
            # 保存文件
            with open(saved_full_path, 'wb+') as destination:
                for chunk in file.chunks():
                    destination.write(chunk)
            
            # 更新文档记录
            document.file_name = file.name
            document.file_path = saved_relative_path
            document.file_size = file.size
            document.file_ext = os.path.splitext(file.name)[1].lower()
            document.mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            document.save()
            
            logger.info(f"WPS文档保存成功: document_id={document.id}, file_size={file.size}")
            
            return DetailResponse(
                data={
                    'documentId': document.id,
                    'filePath': saved_relative_path,
                    'fileSize': file.size,
                    'updateTime': document.update_datetime.strftime('%Y-%m-%dT%H:%M:%S') if hasattr(document, 'update_datetime') else None,
                    'version': 1  # 版本号可以后续实现
                },
                msg="文档保存成功"
            )
            
        except Exception as e:
            logger.error(f"WPS文档保存失败: {str(e)}")
            import traceback
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            return ErrorResponse(msg=f"文档保存失败: {str(e)}")
    
    @action(detail=True, methods=['get'], url_path='wps/download', permission_classes=[])
    def wps_download(self, request, pk=None):
        """下载WPS文档（直接使用现有的download action）"""
        return self.download(request, pk)
    
    @action(detail=True, methods=['post'], url_path='wps/callback', permission_classes=[])
    def wps_callback(self, request, pk=None):
        """处理WPS回调
        
        回调数据格式：
        {
            "event": "file_save",  // "file_save", "file_close", "file_error"
            "fileId": "123",
            "userId": "456",
            "timestamp": "2025-11-02T10:30:00",
            "data": {
                "fileUrl": "https://wps-server.com/files/xxx.docx",
                "error": null
            }
        }
        """
        try:
            document = self.get_object()
            event = request.data.get('event', '')
            file_id = request.data.get('fileId', '')
            user_id = request.data.get('userId', '')
            timestamp = request.data.get('timestamp', '')
            event_data = request.data.get('data', {})
            
            logger.info(f"WPS回调: document_id={document.id}, event={event}, file_id={file_id}")
            
            # 处理不同的事件类型
            if event == 'file_save':
                logger.info(f"WPS文档保存事件: document_id={document.id}")
                # 可以在这里触发自动备份或其他操作
            elif event == 'file_close':
                logger.info(f"WPS文档关闭事件: document_id={document.id}")
                # 可以在这里清理临时文件
            elif event == 'file_error':
                error_msg = event_data.get('error', '未知错误')
                logger.error(f"WPS文档错误事件: document_id={document.id}, error={error_msg}")
            
            return DetailResponse(
                data={
                    'documentId': document.id,
                    'event': event,
                    'status': 'success'
                },
                msg="回调处理成功"
            )
            
        except Exception as e:
            logger.error(f"WPS回调处理失败: {str(e)}")
            import traceback
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            return ErrorResponse(msg=f"回调处理失败: {str(e)}")


class DocumentTemplateViewSet(CustomModelViewSet):
    """文档模板视图集"""
    
    queryset = DocumentTemplate.objects.filter(is_deleted=False)
    serializer_class = DocumentTemplateSerializer
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['template_type', 'is_active']
    search_fields = ['template_name', 'description']
    ordering_fields = ['id', 'sort_order', 'create_datetime']  # ✅ 添加 sort_order 排序支持
    ordering = ['sort_order', 'id']  # ✅ 默认按 sort_order 升序排列
    permission_classes = []
    extra_filter_class = []  # 禁用额外的过滤类
    
    # 只返回JSON，不渲染HTML模板
    from rest_framework.renderers import JSONRenderer
    renderer_classes = [JSONRenderer]
    
    def get_queryset(self):
        """获取查询集"""
        # 直接返回所有未删除的模板，不过滤is_active
        return DocumentTemplate.objects.filter(is_deleted=False)
    
    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """下载带占位符的模板文件"""
        try:
            template = self.get_object()
            
            # 获取模板文件的完整路径
            template_file_path = template.full_file_path
            
            # 检查模板文件是否存在
            if not template_file_path or not os.path.exists(template_file_path):
                return DetailResponse(
                    data={},
                    msg="模板文件不存在"
                )
            
            # 读取文件内容
            file_ext = os.path.splitext(template_file_path)[1].lower()
            
            if file_ext in ['.txt', '.md', '.html']:
                # 文本文件
                with open(template_file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # 创建HTTP响应
                response = HttpResponse(content, content_type='text/plain; charset=utf-8')
                response['Content-Disposition'] = f'attachment; filename="{template.template_name}_placeholder.txt"'
                return response
                
            elif file_ext in ['.docx', '.doc']:
                # Word文档，直接下载文件
                with open(template_file_path, 'rb') as f:
                    content = f.read()
                
                response = HttpResponse(content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{template.template_name}_placeholder{file_ext}"'
                return response
            else:
                # 其他格式，尝试下载
                with open(template_file_path, 'rb') as f:
                    content = f.read()
                
                response = HttpResponse(content, content_type='application/octet-stream')
                response['Content-Disposition'] = f'attachment; filename="{template.template_name}_placeholder{file_ext}"'
                return response
                
        except Exception as e:
            return DetailResponse(
                data={},
                msg=f"模板下载失败: {str(e)}"
            )
    
    @action(detail=True, methods=['get'])
    def preview(self, request, pk=None):
        """预览模板内容"""
        try:
            template = self.get_object()
            
            # 获取模板文件的完整路径
            template_file_path = template.full_file_path
            
            # 读取模板文件内容
            if not template_file_path or not os.path.exists(template_file_path):
                return DetailResponse(
                    data={'content': '模板文件不存在'},
                    msg="模板文件不存在"
                )
            
            # 根据文件类型读取内容
            file_ext = os.path.splitext(template_file_path)[1].lower()
            
            if file_ext == '.txt':
                # 文本文件
                with open(template_file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_ext in ['.doc', '.docx']:
                # Word文档，提取文本内容（不依赖langchain）
                try:
                    if file_ext == '.docx':
                        # 解析 .docx 文件
                        from docx import Document
                        doc = Document(template_file_path)
                        
                        content_parts = []
                        for paragraph in doc.paragraphs:
                            text = paragraph.text.strip()
                            if text:
                                content_parts.append(text)
                            else:
                                # 保留空行以保持格式
                                content_parts.append('')
                        
                        content = '\n'.join(content_parts) if content_parts else '(文档内容为空)'
                        
                        # 添加表格内容
                        if doc.tables:
                            content += '\n\n--- 表格内容 ---\n'
                            for i, table in enumerate(doc.tables, 1):
                                content += f'\n表格 {i}:\n'
                                for row in table.rows:
                                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                                    if row_text.strip():
                                        content += row_text + '\n'
                    else:
                        # .doc 文件（旧格式）
                        content = "⚠️ 暂不支持预览 .doc 格式文件（仅支持 .docx），请下载后使用 Word 查看"
                        
                except ImportError as e:
                    logger.error(f"导入docx模块失败: {str(e)}")
                    content = "❌ 无法解析Word文档：缺少 python-docx 库\n请安装: pip install python-docx"
                except Exception as e:
                    logger.error(f"解析Word文档失败: {str(e)}")
                    content = f"❌ 解析Word文档失败: {str(e)}\n建议下载文件后查看"
            else:
                # 其他格式，尝试读取为文本
                try:
                    with open(template_file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    content = f"无法预览此格式的文件: {file_ext}"
            
            return DetailResponse(
                data={
                    'template_name': template.template_name,
                    'template_type': template.template_type,
                    'content': content,
                    'file_size': template.file_size,
                    'create_datetime': template.create_datetime
                },
                msg="模板预览成功"
            )
        except Exception as e:
            return DetailResponse(
                data={},
                msg=f"模板预览失败: {str(e)}"
            )
    
    @action(detail=True, methods=['get'])
    def placeholders(self, request, pk=None):
        """获取模板的占位符信息"""
        try:
            template = self.get_object()
            
            # 如果模板已有占位符信息，直接返回
            if template.placeholder_info:
                return DetailResponse(
                    data={
                        'template_id': template.id,
                        'template_name': template.template_name,
                        'placeholder_info': template.placeholder_info
                    },
                    msg="获取占位符信息成功"
                )
            
            # 如果没有占位符信息，重新解析
            from .placeholder_parser import placeholder_parser
            template_file_path = template.full_file_path
            placeholder_info = placeholder_parser.parse_template_file(template_file_path)
            
            # 更新模板的占位符信息
            template.placeholder_info = placeholder_info
            template.save()
            
            return DetailResponse(
                data={
                    'template_id': template.id,
                    'template_name': template.template_name,
                    'placeholder_info': placeholder_info
                },
                msg="解析占位符信息成功"
            )
            
        except Exception as e:
            return DetailResponse(
                data={},
                msg=f"获取占位符信息失败: {str(e)}"
            )
    
    @action(detail=True, methods=['put', 'patch'])
    def update_placeholders(self, request, pk=None):
        """更新模板的占位符信息"""
        try:
            template = self.get_object()
            
            # 验证请求数据
            serializer = PlaceholderEditSerializer(data=request.data)
            if not serializer.is_valid():
                return DetailResponse(
                    data={},
                    msg=f"数据验证失败: {serializer.errors}"
                )
            
            # 更新占位符信息，确保中文不被转义
            from .placeholder_parser import placeholder_parser
            placeholder_info = serializer.validated_data['placeholder_info']
            placeholder_info_fixed = placeholder_parser.fix_unicode_escaped_data(placeholder_info)
            template.placeholder_info = placeholder_info_fixed
            template.save()
            
            # 返回更新后的信息
            template_serializer = self.get_serializer(template)
            return DetailResponse(
                data=template_serializer.data,
                msg="占位符信息更新成功"
            )
            
        except Exception as e:
            return DetailResponse(
                data={},
                msg=f"更新占位符信息失败: {str(e)}"
            )
    
    @action(detail=True, methods=['post'])
    def reparse_placeholders(self, request, pk=None):
        """重新解析模板的占位符信息"""
        try:
            template = self.get_object()
            
            # 重新解析占位符
            from .placeholder_parser import placeholder_parser
            template_file_path = template.full_file_path
            placeholder_info = placeholder_parser.parse_template_file(template_file_path)
            
            # 确保占位符信息中的中文不被转义
            placeholder_info_fixed = placeholder_parser.fix_unicode_escaped_data(placeholder_info)
            
            # 更新模板的占位符信息
            template.placeholder_info = placeholder_info_fixed
            template.save()
            
            # 返回解析结果
            template_serializer = self.get_serializer(template)
            return DetailResponse(
                data=template_serializer.data,
                msg=f"重新解析成功，发现 {placeholder_info.get('total_count', 0)} 个占位符"
            )
            
        except Exception as e:
            return DetailResponse(
                data={},
                msg=f"重新解析占位符失败: {str(e)}"
            )
    
    @action(detail=False, methods=['post'])
    def upload(self, request):
        """上传带占位符的模板文件并解析占位符信息"""
        try:
            # 获取上传的文件
            uploaded_file = request.FILES.get('file')
            if not uploaded_file:
                return DetailResponse(
                    data={},
                    msg="请选择要上传的文件"
                )
            
            # 保存原始文件到 media/case_templates 目录
            import time
            file_name, file_ext = os.path.splitext(uploaded_file.name)
            unique_filename = f"{file_name}_{int(time.time())}{file_ext}"
            
            # 生成相对路径（相对于 MEDIA_ROOT）
            relative_path = os.path.join('case_templates', unique_filename)
            
            # 创建物理目录
            template_dir = os.path.join(settings.MEDIA_ROOT, 'case_templates')
            os.makedirs(template_dir, exist_ok=True)
            
            # 完整物理路径
            template_file_path = os.path.join(settings.MEDIA_ROOT, relative_path)
            
            # 保存上传的文件
            with open(template_file_path, 'wb') as f:
                for chunk in uploaded_file.chunks():
                    f.write(chunk)
            
            # 解析模板中的占位符信息
            from .placeholder_parser import placeholder_parser
            placeholder_info = placeholder_parser.parse_template_file(template_file_path)
            
            # 确保占位符信息中的中文不被转义
            placeholder_info_fixed = placeholder_parser.fix_unicode_escaped_data(placeholder_info)
            
            # 创建模板记录，包含占位符信息
            template_data = {
                'template_name': request.data.get('template_name', uploaded_file.name),
                'template_type': request.data.get('template_type', 'other'),
                'file_path': relative_path,  # 保存相对路径而非绝对路径
                'file_size': uploaded_file.size,
                'description': request.data.get('description', ''),
                'placeholder_info': placeholder_info_fixed,  # 存储修复后的占位符信息
                'is_active': True
            }
            
            serializer = self.get_serializer(data=template_data)
            if serializer.is_valid():
                template = serializer.save()
                
                # 返回详细的占位符信息
                response_data = serializer.data.copy()
                response_data['placeholder_summary'] = {
                    'total_placeholders': placeholder_info.get('total_count', 0),
                    'unique_placeholders': placeholder_info.get('unique_count', 0),
                    'categories': placeholder_info.get('analysis', {}).get('by_category', {}),
                    'with_defaults': len(placeholder_info.get('analysis', {}).get('with_defaults', [])),
                    'without_defaults': len(placeholder_info.get('analysis', {}).get('without_defaults', []))
                }
                
                # 添加文件访问 URL
                response_data['file_url'] = template.file_url
                
                return DetailResponse(
                    data=response_data,
                    msg=f"模板上传成功，解析到 {placeholder_info.get('total_count', 0)} 个占位符"
                )
            else:
                return DetailResponse(
                    data={},
                    msg=f"模板上传失败: {serializer.errors}"
                )
                
        except Exception as e:
            return DetailResponse(
                data={},
                msg=f"模板上传失败: {str(e)}"
            )
    
    def _convert_to_placeholder_template(self, original_file_path: str, file_name: str, file_ext: str) -> str:
        """将原始模板转换为占位符模板"""
        try:
            from .placeholder_template_service import placeholder_service
            
            # 定义样例值到占位符的映射规则
            mapping = {
                # 基本主体信息 - 公司名称
                '深圳市振惠建混凝土有限公司': 'plaintiff.name',
                '中建三局第一建设工程有限责任公司': 'defendant.name',
                '中建三局集团有限公司': 'defendant2.name',
                '深圳市某某有限公司': 'plaintiff.name',
                '中建某某有限公司': 'defendant.name',
                '某某建筑公司': 'defendant.name',
                '某某混凝土公司': 'plaintiff.name',
                
                # 统一社会信用代码
                '91440300743220274C': 'plaintiff.credit_code',
                '914201007483157744': 'defendant.credit_code',
                '914201007483157745': 'defendant2.credit_code',
                '91440300XXXXXXXXXX': 'plaintiff.credit_code',
                '91420100XXXXXXXXXX': 'defendant.credit_code',
                
                # 地址信息 - 更全面的地址匹配
                '深圳市南山区': 'plaintiff.address',
                '武汉市江夏区': 'defendant.address',
                '武汉市洪山区': 'defendant2.address',
                '深圳市': 'plaintiff.address',
                '武汉市': 'defendant.address',
                '广东省深圳市': 'plaintiff.address',
                '湖北省武汉市': 'defendant.address',
                '所住地：深圳市南山区': '所住地：{{ plaintiff.address }}',
                '所住地：武汉市江夏区': '所住地：{{ defendant.address }}',
                '住所地：深圳市南山区': '住所地：{{ plaintiff.address }}',
                '住所地：武汉市江夏区': '住所地：{{ defendant.address }}',
                '地址：深圳市南山区': '地址：{{ plaintiff.address }}',
                '地址：武汉市江夏区': '地址：{{ defendant.address }}',
                
                # 法定代表人 - 更全面的匹配
                '张三': 'plaintiff.legal_representative',
                '李四': 'defendant.legal_representative',
                '王五': 'defendant2.legal_representative',
                '法定代表人：张三': '法定代表人：{{ plaintiff.legal_representative }}',
                '法定代表人：李四': '法定代表人：{{ defendant.legal_representative }}',
                '法定代表人：王五': '法定代表人：{{ defendant2.legal_representative }}',
                '法人代表：张三': '法人代表：{{ plaintiff.legal_representative }}',
                '法人代表：李四': '法人代表：{{ defendant.legal_representative }}',
                
                # 金额信息 - 各种格式
                '1000000.00': 'contract_amount',
                '50000.00': 'lawyer_fee',
                '1000000': 'contract_amount',
                '50000': 'lawyer_fee',
                '1000000元': '{{ contract_amount }}元',
                '50000元': '{{ lawyer_fee }}元',
                '合同款1000000.00元': '合同款{{ contract_amount }}元',
                '律师费50000.00元': '律师费{{ lawyer_fee }}元',
                '合同金额：1000000.00元': '合同金额：{{ contract_amount }}元',
                '律师费用：50000.00元': '律师费用：{{ lawyer_fee }}元',
                
                # 案件信息
                '2024年1月1日': 'contract_date',
                '2024年2月1日': 'dispute_date',
                '2024年3月1日': 'filing_date',
                '深圳市南山区人民法院': 'court_name',
                '民事': 'case_type',
                '合同纠纷': 'case_category',
                '案件类型：民事': '案件类型：{{ case_type }}',
                '纠纷类型：合同纠纷': '纠纷类型：{{ case_category }}',
                
                # 其他常见样例值和占位符
                '待填写': 'plaintiff.name',
                '待补充': 'defendant.name',
                '待定': 'plaintiff.name',
                'XXX': 'plaintiff.name',
                'YYY': 'defendant.name',
                'ZZZ': 'defendant2.name',
                '原告姓名': '{{ plaintiff.name }}',
                '被告姓名': '{{ defendant.name }}',
                '原告地址': '{{ plaintiff.address }}',
                '被告地址': '{{ defendant.address }}',
                '原告法定代表人': '{{ plaintiff.legal_representative }}',
                '被告法定代表人': '{{ defendant.legal_representative }}',
            }
            
            # 生成占位符模板文件名和路径
            placeholder_filename = f"{file_name}_placeholder{file_ext}"
            placeholder_dir = os.path.join(settings.BASE_DIR, 'templates', 'case_templates_placeholders')
            os.makedirs(placeholder_dir, exist_ok=True)
            placeholder_file_path = os.path.join(placeholder_dir, placeholder_filename)
            
            # 根据文件类型进行转换
            if file_ext.lower() in ['.txt', '.md', '.html']:
                # 文本文件转换
                with open(original_file_path, 'r', encoding='utf-8') as f:
                    template_text = f.read()
                
                converted_text = placeholder_service.convert_text_template_to_placeholders(
                    template_text, mapping
                )
                
                with open(placeholder_file_path, 'w', encoding='utf-8') as f:
                    f.write(converted_text)
                    
            elif file_ext.lower() in ['.docx']:
                # DOCX文件转换
                placeholder_service.convert_docx_template_to_placeholders(
                    original_file_path, placeholder_file_path, mapping
                )
                
            elif file_ext.lower() in ['.doc']:
                # DOC文件需要先转换为文本，然后保存为占位符模板
                # 这里可以调用现有的解析方法
                from .direct_langchain_ai_service import parse_doc_file
                template_text = parse_doc_file(original_file_path)
                
                converted_text = placeholder_service.convert_text_template_to_placeholders(
                    template_text, mapping
                )
                
                # 保存为文本格式的占位符模板
                placeholder_file_path = placeholder_file_path.replace('.doc', '.txt')
                with open(placeholder_file_path, 'w', encoding='utf-8') as f:
                    f.write(converted_text)
            else:
                # 其他格式，直接复制
                import shutil
                shutil.copy2(original_file_path, placeholder_file_path)
            
            return placeholder_file_path
            
        except Exception as e:
            # 如果转换失败，返回原始文件路径
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"转换占位符模板失败: {e}")
            return original_file_path
    
    def destroy(self, request, *args, **kwargs):
        """删除模板 - 直接删除记录并删除物理文件"""
        try:
            template = self.get_object()
            
            # 获取模板文件的完整物理路径
            template_file_path = template.full_file_path
            
            # 删除模板文件（占位符模板文件）
            if template_file_path and os.path.exists(template_file_path):
                try:
                    os.remove(template_file_path)
                    logger.info(f"已删除模板文件: {template_file_path}")
                except Exception as e:
                    logger.error(f"删除模板文件失败: {template_file_path}, 错误: {str(e)}")
            
            # 删除原始模板文件（如果存在）
            # 从占位符文件路径推断原始文件路径
            if template_file_path:
                placeholder_path = template_file_path
                if 'case_templates_placeholders' in placeholder_path:
                    # 将占位符路径转换为原始路径
                    original_path = placeholder_path.replace('case_templates_placeholders', 'case_templates')
                    # 移除_placeholder后缀
                    if '_placeholder' in original_path:
                        original_path = original_path.replace('_placeholder', '')
                    
                    if os.path.exists(original_path):
                        try:
                            os.remove(original_path)
                            logger.info(f"已删除原始模板文件: {original_path}")
                        except Exception as e:
                            logger.warning(f"删除原始模板文件失败: {original_path}, 错误: {str(e)}")
            
            # ✅ 直接删除数据库记录（不是软删除）
            template.delete(soft_delete=False)
            
            logger.info(f"模板删除成功: id={template.id}, name={template.template_name}")
            return DetailResponse(msg="模板删除成功")
            
        except Exception as e:
            logger.error(f"删除模板失败: {str(e)}")
            return ErrorResponse(msg=f"删除模板失败: {str(e)}")

    @action(detail=False, methods=['post'], url_path='generate_documents')
    def generate_documents(self, request):
        """根据选择的模板生成文书"""
        try:
            case_id = request.data.get('case_id')
            templates = request.data.get('templates', [])
            
            if not case_id:
                return ErrorResponse(msg="案例ID不能为空")
            
            if not templates:
                return ErrorResponse(msg="请至少选择一个模板")
            
            # 获取案例
            try:
                case = CaseManagement.objects.get(id=case_id, is_deleted=False)
            except CaseManagement.DoesNotExist:
                return ErrorResponse(msg="案例不存在")
            
            # 获取选中的模板
            selected_templates = DocumentTemplate.objects.filter(
                id__in=templates, 
                is_deleted=False,
                is_active=True
            )
            
            if not selected_templates.exists():
                return ErrorResponse(msg="选择的模板不存在或未启用")
            
            # 准备案例数据（使用与generate_manual_document相同的逻辑）
            case_data = {
                'case_number': case.case_number,
                'case_name': case.case_name,
                'case_type': case.case_type,
                'case_status': case.case_status,
                'case_description': case.case_description,
                'case_date': case.case_date,
                'draft_person': case.draft_person,
                'defendant_name': case.defendant_name,
                'defendant_credit_code': case.defendant_credit_code,
                'defendant_address': case.defendant_address,
                'defendant_legal_representative': case.defendant_legal_representative,
                'plaintiff_name': case.plaintiff_name,
                'plaintiff_credit_code': case.plaintiff_credit_code,
                'plaintiff_address': case.plaintiff_address,
                'plaintiff_legal_representative': case.plaintiff_legal_representative,
                'contract_amount': float(case.contract_amount) if case.contract_amount else 0.0,
                'lawyer_fee': float(case.lawyer_fee) if case.lawyer_fee else 0.0,
                'litigation_request': case.litigation_request,
                'facts_and_reasons': case.facts_and_reasons,
                'case_result': case.case_result,
                'case_notes': case.case_notes,
                'jurisdiction': case.jurisdiction,
                'petitioner': case.petitioner,
                'filing_date': case.filing_date
            }

            # 构建被告人数组（供 docxtpl 循环使用）
            def _split_multi(value):
                if not value:
                    return []
                if isinstance(value, str):
                    # 支持中文顿号、逗号
                    parts = [p.strip() for p in re.split(r'[，,、]', value) if p and p.strip()]
                    return parts if parts else [value]
                return [str(value)]

            try:
                import re  # 局部导入，避免顶部过多依赖
                names = _split_multi(case.defendant_name)
                codes = _split_multi(case.defendant_credit_code)
                addrs = _split_multi(case.defendant_address)
                reps  = _split_multi(case.defendant_legal_representative)

                max_len = max(len(names), len(codes), len(addrs), len(reps), 1)
                defendants_list = []
                for i in range(max_len):
                    defendants_list.append({
                        # 注意：模板中使用 defendant.defendant_name 等键名
                        'defendant_name': names[i] if i < len(names) else '',
                        'defendant_credit_code': codes[i] if i < len(codes) else '',
                        'defendant_address': addrs[i] if i < len(addrs) else '',
                        'defendant_legal_representative': reps[i] if i < len(reps) else '',
                    })
                case_data['defendants'] = defendants_list
                case_data['defendants_count'] = len(defendants_list)
            except Exception as e:
                logger.warning(f"构建被告人数组失败: {e}")
                case_data['defendants'] = []
                case_data['defendants_count'] = 0
            
            # 将Decimal字段转换为float，避免JSON序列化错误
            for key, value in case_data.items():
                if hasattr(value, 'as_tuple'):  # 检查是否为Decimal类型
                    case_data[key] = float(value)
            
            success_count = 0
            failed_count = 0
            generated_documents = []
            
            # 为每个选中的模板生成文书（使用与generate_manual_document相同的逻辑）
            for template in selected_templates:
                try:
                    # 直接使用字段名作为占位符进行匹配
                    template_data = {}
                    
                    if template.placeholder_info and 'placeholders' in template.placeholder_info:
                        logger.info(f"模板 {template.template_name} 的占位符: {template.placeholder_info['placeholders']}")
                        for placeholder in template.placeholder_info['placeholders']:
                            placeholder_key = placeholder['key']
                            logger.info(f"处理占位符: {placeholder_key}")
                            
                            # 直接使用占位符key作为字段名进行匹配
                            if placeholder_key in case_data:
                                value = case_data[placeholder_key]
                                logger.info(f"找到匹配字段 {placeholder_key}: {value}")
                                
                                # 特殊处理：对contract_amount进行法律语言格式化
                                if placeholder_key == 'contract_amount' and value:
                                    try:
                                        amount = float(value)
                                        value = f"请求判令被告支付合同款项人民币{amount:,.2f}元"
                                    except (ValueError, TypeError):
                                        value = f"请求判令被告支付合同款项人民币{value}元"
                                
                                template_data[placeholder_key] = value
                            else:
                                # 如果字段不存在，使用默认值
                                logger.warning(f"字段 {placeholder_key} 在案例数据中不存在，使用默认值")
                                template_data[placeholder_key] = '待填写'

                    # 无论占位符里是否声明，确保将可循环的数据结构放入上下文，供 docxtpl 使用
                    if 'defendants' in case_data:
                        template_data['defendants'] = case_data['defendants']
                        template_data['defendants_count'] = case_data.get('defendants_count', len(case_data['defendants']))
                    
                    # 使用占位符服务生成文档
                    from .placeholder_template_service import placeholder_service
                    doc = placeholder_service.fill_and_save_by_record(case.id, template, template_data, request=request)
                    generated_documents.append({
                        'id': doc.id,
                        'name': doc.document_name,
                        'type': doc.document_type
                    })
                    success_count += 1
                    logger.info(f"成功生成文档: {template.template_name}")
                    
                except Exception as e:
                    failed_count += 1
                    logger.error(f"生成模板 {template.template_name} 的文书失败: {str(e)}")
            
            return DetailResponse(
                data={
                    'success_count': success_count,
                    'failed_count': failed_count,
                    'generated_documents': generated_documents
                },
                msg=f"文书生成完成！成功: {success_count}个，失败: {failed_count}个"
            )
            
        except Exception as e:
            logger.error(f"批量生成文书失败: {str(e)}")
            return ErrorResponse(msg=f"生成文书失败: {str(e)}")
    
    @action(detail=False, methods=['post'], url_path='batch-sort')
    def batch_sort(self, request):
        """批量更新模板排序
        
        请求格式:
        {
            "templates": [
                {"id": 1, "sort_order": 1000},
                {"id": 2, "sort_order": 2000}
            ]
        }
        """
        try:
            # 验证请求数据
            from .serializers import BatchSortSerializer
            serializer = BatchSortSerializer(data=request.data)
            if not serializer.is_valid():
                return ErrorResponse(msg=f"数据验证失败: {serializer.errors}")
            
            templates_data = serializer.validated_data['templates']
            
            # 使用事务批量更新
            from django.db import transaction
            
            with transaction.atomic():
                # 提取所有模板ID
                template_ids = [item['id'] for item in templates_data]
                
                # 查询对应的模板对象
                templates = DocumentTemplate.objects.filter(id__in=template_ids)
                
                # 检查是否所有模板都存在
                found_ids = set(t.id for t in templates)
                missing_ids = set(template_ids) - found_ids
                if missing_ids:
                    return ErrorResponse(msg=f"模板不存在: {list(missing_ids)}")
                
                # 创建ID到sort_order的映射
                sort_map = {item['id']: item['sort_order'] for item in templates_data}
                
                # 更新sort_order
                for template in templates:
                    template.sort_order = sort_map[template.id]
                
                # 批量保存
                DocumentTemplate.objects.bulk_update(templates, ['sort_order'])
            
            logger.info(f"成功更新 {len(templates)} 个模板的排序")
            
            return DetailResponse(
                data={
                    'updated_count': len(templates),
                    'updated_ids': template_ids
                },
                msg=f"成功更新 {len(templates)} 个模板的排序"
            )
            
        except Exception as e:
            logger.error(f"批量更新排序失败: {str(e)}")
            return ErrorResponse(msg=f"更新排序失败: {str(e)}")


class CaseFolderViewSet(CustomModelViewSet):
    """案件目录管理ViewSet"""
    
    queryset = CaseFolder.objects.filter(is_deleted=False)
    serializer_class = CaseFolderSerializer
    permission_classes = []
    
    filter_backends = [DjangoFilterBackend, OrderingFilter]
    filterset_fields = ['case', 'parent', 'folder_type']
    ordering_fields = ['sort_order', 'id']
    ordering = ['sort_order', 'id']
    
    # 只返回JSON，不渲染HTML模板
    from rest_framework.renderers import JSONRenderer
    renderer_classes = [JSONRenderer]


# ==================== 文档转换 API ====================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def convert_docx_to_html(request):
    """
    DOCX转HTML接口
    
    请求参数:
        - documentId: 文档ID（必需）
    """
    try:
        data = request.data
        document_id = data.get('documentId')
        
        if not document_id:
            return ErrorResponse(msg='请提供 documentId')
        
        # 获取文档对象
        document = get_object_or_404(CaseDocument, id=document_id)
        
        # 检查权限
        user = get_request_user(request)
        if document.case_id and hasattr(document, 'case'):
            # 可以添加更多权限检查
            pass
        
        # 初始化转换器
        converter = DocumentConverter()
        file_manager = DocumentFileManager()
        
        # 获取文件路径
        media_root = getattr(settings, 'MEDIA_ROOT', 'media')
        if document.file_path:
            full_path = os.path.join(media_root, document.file_path)
        elif document.file_name:
            # 尝试从file_name构建路径
            full_path = os.path.join(media_root, 'cases', str(document.case_id), document.file_name)
        else:
            return ErrorResponse(msg='文档文件不存在')
        
        if not os.path.exists(full_path):
            return ErrorResponse(msg='文档文件不存在')
        
        docx_path = full_path
        title = document.document_name or os.path.splitext(document.file_name or 'document')[0]
        
        # 创建图片输出目录
        from datetime import datetime
        now = datetime.now()
        year_month = now.strftime('%Y/%m')
        image_output_dir = os.path.join(settings.MEDIA_ROOT, 'images', 'documents', year_month)
        os.makedirs(image_output_dir, exist_ok=True)
        
        # 转换为HTML
        result = converter.docx_to_html(docx_path, image_output_dir)
        
        return DetailResponse(
            data={
                'documentId': document_id,
                'html': result['html'],
                'title': result.get('title', title),
                'images': result['images']
            },
            msg='转换成功'
        )
        
    except CaseDocument.DoesNotExist:
        return ErrorResponse(msg='文档不存在')
    except Exception as e:
        logger.error(f"DOCX转HTML失败: {str(e)}", exc_info=True)
        return ErrorResponse(msg=f'转换失败: {str(e)}')


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def convert_html_to_docx(request):
    """
    HTML转DOCX接口
    
    请求参数:
        - documentId: 文档ID（必需）
        - html: HTML内容（必需）
        - title: 文档标题（可选）
        - savePath: 保存路径（可选）
    """
    try:
        data = request.data
        document_id = data.get('documentId')
        html_content = data.get('html', '')
        title = data.get('title', '')
        save_path = data.get('savePath')
        
        if not document_id:
            return ErrorResponse(msg='请提供 documentId')
        
        if not html_content:
            return ErrorResponse(msg='HTML内容不能为空')
        
        # 获取文档对象
        document = get_object_or_404(CaseDocument, id=document_id)
        
        # 检查权限
        user = get_request_user(request)
        
        # 初始化转换器和文件管理器
        converter = DocumentConverter()
        file_manager = DocumentFileManager()
        
        # 创建临时文件
        temp_file_path = file_manager.create_temp_file(suffix='.docx')
        
        try:
            # 获取图片基础目录
            image_base_dir = None
            if document.case_id:
                media_root = getattr(settings, 'MEDIA_ROOT', 'media')
                image_base_dir = os.path.join(media_root, 'images', 'documents')
            
            # 转换为DOCX
            converter.html_to_docx(html_content, temp_file_path, image_base_dir)
            
            # 读取生成的DOCX文件
            with open(temp_file_path, 'rb') as f:
                docx_content = f.read()
            
            # 保存文档
            filename = save_path or f"{document.document_name or 'document'}.docx"
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            saved_path = file_manager.save_document(docx_content, document_id, filename)
            
            # 更新文档记录
            document.file_path = saved_path
            document.file_size = len(docx_content)
            document.file_ext = '.docx'
            document.mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            document.save()
            
            return DetailResponse(
                data={
                    'documentId': document_id,
                    'documentPath': saved_path,
                    'fileSize': len(docx_content),
                    'updateTime': document.update_datetime.isoformat() if hasattr(document, 'update_datetime') else None
                },
                msg='转换并保存成功'
            )
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_file_path):
                try:
                    os.remove(temp_file_path)
                except Exception as e:
                    logger.warning(f"删除临时文件失败: {e}")
        
    except Exception as e:
        logger.error(f"HTML转DOCX失败: {str(e)}", exc_info=True)
        return ErrorResponse(msg=f'转换失败: {str(e)}')


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def upload_document_image(request):
    """
    上传文档图片接口
    
    请求参数:
        - file: 图片文件（必需，multipart/form-data）
        - documentId: 文档ID（可选，用于关联）
    """
    try:
        if 'file' not in request.FILES:
            return ErrorResponse(msg='请上传图片文件')
        
        uploaded_file = request.FILES['file']
        document_id = request.data.get('documentId')
        
        # 初始化图片处理器
        image_handler = ImageHandler()
        
        # 上传图片
        result = image_handler.upload_image(uploaded_file, document_id)
        
        return DetailResponse(
            data={
                'url': result['url'],
                'alt': result.get('filename', ''),
                'href': result['url']
            },
            msg='图片上传成功'
        )
        
    except ValueError as e:
        return ErrorResponse(msg=str(e))
    except Exception as e:
        logger.error(f"图片上传失败: {str(e)}", exc_info=True)
        return ErrorResponse(msg=f'上传失败: {str(e)}')
