#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
占位符模版填充与保存服务

功能目标：
- 提供统一的占位符渲染能力（支持文本/HTML与DOCX），占位符形如 {{ key }} 或 {{ key | 默认值 }}，支持点号路径（如 {{ plaintiff.name }}）
- 将渲染后的结果保存为文书，登记到数据库 `CaseDocument`
- 提供将现有模版转换为占位符模版的工具方法（文本与DOCX）

注意：
- DOCX 优先使用 docxtpl 保留格式；如未安装 docxtpl，则使用 python-docx 执行逐段落/逐run替换（可能损失部分复杂格式）
- 本文件不依赖外部命令行工具，兼容无 LibreOffice 环境
"""

from __future__ import annotations

import os
import re
import json
import logging
from typing import Any, Dict, Optional

logger = logging.getLogger(__name__)


def number_to_chinese(num: int) -> str:
    """将数字转换为中文数字（支持 1-99）。

    例如：1 -> 一, 2 -> 二, 10 -> 十, 11 -> 十一
    """
    chinese_nums = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']

    if num <= 0:
        return ''
    if num <= 10:
        return chinese_nums[num]
    if num < 20:
        return '十' + chinese_nums[num % 10] if num % 10 != 0 else '十'
    if num < 100:
        tens = num // 10
        ones = num % 10
        tens_part = chinese_nums[tens] + '十'
        return tens_part if ones == 0 else tens_part + chinese_nums[ones]
    # 超过 99 时直接返回数字字符串，确保不会报错
    return str(num)


def _get_by_path(data: Dict[str, Any], path: str, default: str = "") -> str:
    """根据点号路径从字典中取值，取不到返回默认值

    例：path = 'plaintiff.name' -> data['plaintiff']['name']
    """
    try:
        current: Any = data
        for key in path.split('.'):  # 逐层下钻
            if not isinstance(current, dict) or key not in current:
                return default
            current = current[key]
        if current is None or current == "":
            return default
        return str(current)
    except Exception:
        return default


def _flatten_dict(data: Dict[str, Any], parent_key: str = "", sep: str = ".") -> Dict[str, Any]:
    """将嵌套字典拍平成扁平字典，便于 docxtpl 渲染

    例：{"plaintiff": {"name": "张三"}} -> {"plaintiff.name": "张三"}
    """
    items: Dict[str, Any] = {}
    for k, v in data.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.update(_flatten_dict(v, new_key, sep=sep))
        else:
            items[new_key] = v
    return items


class PlaceholderTemplateService:
    """占位符模版渲染与保存服务类"""

    # ==================== 占位符渲染（文本/HTML） ====================
    def render_text(self, template_text: str, data: Dict[str, Any], alias: Optional[Dict[str, str]] = None) -> str:
        """渲染文本/HTML模版中的占位符

        占位符语法：
        - {{ key }}
        - {{ key | 默认值 }}
        - 支持点号路径：{{ a.b.c }}
        - alias 用于将旧字段名映射到点号路径，如 {'plaintiff_name': 'plaintiff.name'}
        """
        alias = alias or {}
        pattern = re.compile(r"\{\{\s*([a-zA-Z0-9_.]+)\s*(?:\|\s*([^}]+))?\s*\}\}")

        def _replace(m: re.Match) -> str:
            key = (m.group(1) or "").strip()
            default = (m.group(2) or "").strip()
            real_key = alias.get(key, key)
            return _get_by_path(data, real_key, default)

        # 逐处替换
        return pattern.sub(_replace, template_text)

    # ==================== 占位符渲染（DOCX） ====================
    def render_docx(self, template_path: str, data: Dict[str, Any], out_path: str) -> str:
        """渲染 DOCX 模版：优先使用 docxtpl，其次回退到 python-docx 逐 run 替换

        返回：生成文件的绝对路径
        """
        # 确保输出目录存在
        output_dir = os.path.dirname(out_path)
        if output_dir:  # 只有当目录路径不为空时才创建
            os.makedirs(output_dir, exist_ok=True)

        # 优先 docxtpl（保留版式最完整）
        try:
            from docxtpl import DocxTemplate  # type: ignore
            from docx import Document  # type: ignore
            
            flat = _flatten_dict(data)
            # 同时将扁平键的最末段也放一份，提升容错
            simple = {k.split('.')[-1]: v for k, v in flat.items()}
            ctx = {**flat, **simple}
            
            # ✅ 重要：确保列表类型的数据（如 defendants）被保留在上下文中
            # 因为 _flatten_dict 不会处理列表，所以需要手动添加
            if 'defendants' in data:
                ctx['defendants'] = data['defendants']
            if 'defendants_count' in data:
                ctx['defendants_count'] = data['defendants_count']

            doc = DocxTemplate(template_path)

            # 注入工具函数，供 Jinja 模板中使用
            ctx['number_to_chinese'] = number_to_chinese

            doc.render(ctx)
            doc.save(out_path)
            
            # ✅ 渲染后处理：删除空段落（去除多余换行）
            try:
                docx_doc = Document(out_path)
                paragraphs_to_remove = []
                for para in docx_doc.paragraphs:
                    # 检查段落是否为空（只有空白字符或完全为空）
                    text = para.text.strip()
                    if not text:
                        paragraphs_to_remove.append(para)
                
                # 删除空段落
                for para in paragraphs_to_remove:
                    p = para._element
                    p.getparent().remove(p)
                
                # 保存处理后的文档
                docx_doc.save(out_path)
                logger.info("已清理文档中的空段落")
            except Exception as e:
                logger.warning(f"清理空段落时出错，但不影响文档生成: {e}")
            
            logger.info("使用 docxtpl 成功渲染 DOCX 模版")
            return out_path
        except Exception as e:
            logger.warning(f"docxtpl 渲染失败，尝试 python-docx：{e}")

        # 回退 python-docx：对每个 run 做简单的字符串替换（复杂版式可能有损）
        try:
            from docx import Document  # type: ignore

            # 准备一个替换表：{{ key }} 与 {{ key | 默认 }} 都替换
            flat = _flatten_dict(data)

            def _build_replacements() -> Dict[str, str]:
                repl: Dict[str, str] = {}
                for k, v in flat.items():
                    repl[f"{{{{ {k} }}}}"] = str(v)
                    repl[f"{{{{{k}}}}}"] = str(v)
                # 同时补一份末级键，便于兼容 {{ name }}
                for k, v in flat.items():
                    end = k.split('.')[-1]
                    repl[f"{{{{ {end} }}}}"] = str(v)
                    repl[f"{{{{{end}}}}}"] = str(v)
                return repl

            repl_map = _build_replacements()

            def _replace_text(text: str) -> str:
                if not text:
                    return text
                # 简单多次替换（不解析默认值语法）
                for k, v in repl_map.items():
                    if k in text:
                        text = text.replace(k, v)
                # 处理带默认值 {{ key | default }}：正则匹配后降级为直接取值，不存在则替换成 default
                def _repl_default(m: re.Match) -> str:
                    key = (m.group(1) or "").strip()
                    default = (m.group(2) or "").strip()
                    real = flat.get(key, flat.get(key.split('.')[-1], default))
                    return str(real if real not in (None, '') else default)
                text = re.sub(r"\{\{\s*([a-zA-Z0-9_.]+)\s*\|\s*([^}]+)\s*\}\}", _repl_default, text)
                return text

            doc = Document(template_path)
            # 段落替换
            for p in doc.paragraphs:
                for r in p.runs:
                    r.text = _replace_text(r.text)
            # 表格中的文本
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = _replace_text(r.text)

            doc.save(out_path)
            logger.info("使用 python-docx 成功渲染 DOCX 模版（简化替换）")
            return out_path
        except Exception as e:
            logger.error(f"python-docx 渲染 DOCX 失败：{e}")
            # 如果所有方法都失败，尝试创建一个简单的DOCX文件
            try:
                from docx import Document
                doc = Document()
                doc.add_paragraph("文档生成失败，请检查模板文件格式。")
                doc.save(out_path)
                logger.warning(f"创建了简单的DOCX文件作为回退: {out_path}")
                return out_path
            except Exception as e2:
                logger.error(f"创建回退DOCX文件也失败: {e2}")
                raise

    # ==================== 保存为文书（登记数据库） ====================
    def save_text_document(self, case_id: int, document_name: str, content: str, file_ext: str = '.txt', folder_path: str = '/case_documents', template_id: int = None, creator=None, dept_belong_id: str = None, template_print_count: int = 1, template_sort_order: int = 0) -> "CaseDocument":
        """保存文本/HTML渲染结果为文书（存储内容到 document_content 字段）
        
        Args:
            case_id: 案件ID
            document_name: 文档名称
            content: 文档内容
            file_ext: 文件扩展名
            folder_path: 目录路径，默认为 '/case_documents'（案件文书目录）
            template_id: 模板ID，用于检查是否已存在相同文档
            creator: 创建人对象（User实例）
            dept_belong_id: 所属部门ID
            template_print_count: 模板配置的默认打印份数（生成时记录到文档）
            template_sort_order: 模板的排序序号（生成时记录到文档）
        
        逻辑：
            - 如果存在相同的文档（同案件、同模板、同目录），则更新该文档
            - 如果不存在，则创建新文档
        """
        from .models import CaseDocument, CaseManagement, CaseFolder  # 延迟导入避免循环依赖
        from django.db.models import Q
        case = CaseManagement.objects.get(id=case_id)
        
        # 获取目录对象
        folder = None
        try:
            folder = CaseFolder.objects.get(
                case=case,
                folder_path=folder_path,
                is_deleted=False
            )
        except CaseFolder.DoesNotExist:
            logger.warning(f"目录不存在: {folder_path}")
        
        # 根据文件扩展名确定文档类型
        doc_type_map = {
            '.txt': 'text',
            '.md': 'text', 
            '.html': 'text',
            '.doc': 'word',
            '.docx': 'word',
            '.pdf': 'pdf'
        }
        doc_type = doc_type_map.get(file_ext.lower(), 'text')
        
        # ✅ 简化检查：只根据 case_id 和 template_id 判断
        existing_doc = None
        if template_id:
            existing_doc = CaseDocument.objects.filter(
                case=case,
                template_id=template_id,
                is_deleted=False
            ).first()
            
            if existing_doc:
                logger.info(f"[DEBUG] 找到已存在的文本文档: id={existing_doc.id}, case_id={case.id}, template_id={template_id}, 将更新而非创建")
        
        if existing_doc:
            # ✅ 更新现有文档（不修改 print_count，保持原值）
            existing_doc.document_name = document_name
            existing_doc.document_type = doc_type
            existing_doc.document_content = content
            existing_doc.file_size = len(content.encode('utf-8'))
            existing_doc.folder = folder
            existing_doc.folder_path = folder_path
            existing_doc.template_used = document_name
            # ❌ 不修改 print_count，保持文档首次创建时的值
            existing_doc.save()
            logger.info(f"[DEBUG] 文本文档记录已更新: id={existing_doc.id}, print_count保持不变={existing_doc.print_count}")
            return existing_doc
        else:
            # 创建新文档
            doc = CaseDocument.objects.create(
                case=case,
                folder=folder,                    # ✅ 设置目录对象
                folder_path=folder_path,          # ✅ 设置目录路径
                document_name=document_name,
                document_type=doc_type,
                document_content=content,
                generation_method='manual',
                template_used=document_name,
                template_id=template_id,          # ✅ 设置模板ID
                creator=creator,                  # ✅ 设置创建人
                dept_belong_id=dept_belong_id,    # ✅ 设置所属部门
                print_count=template_print_count, # ✅ 记录模板配置的默认打印份数
                sort_order=template_sort_order,   # ✅ 记录模板的排序序号
                file_path='',
                file_size=len(content.encode('utf-8')),
            )
            logger.info(f"[DEBUG] 文本文档记录已创建: id={doc.id}, creator={creator}, dept_belong_id={dept_belong_id}, print_count={template_print_count}, sort_order={template_sort_order}")
            return doc

    def save_docx_document(self, case_id: int, document_name: str, file_path: str, folder_path: str = '/case_documents', template_id: int = None, creator=None, dept_belong_id: str = None, template_print_count: int = 1, template_sort_order: int = 0) -> "CaseDocument":
        """保存 DOCX 渲染结果为文书（登记文件路径与大小）
        
        Args:
            case_id: 案件ID
            document_name: 文档名称
            file_path: 文件完整路径
            folder_path: 目录路径，默认为 '/case_documents'（案件文书目录）
            template_id: 模板ID，用于检查是否已存在相同文档
            creator: 创建人对象（User实例）
            dept_belong_id: 所属部门ID
            template_print_count: 模板配置的默认打印份数（生成时记录到文档）
            template_sort_order: 模板的排序序号（生成时记录到文档）
        
        逻辑：
            - 如果存在相同的文档（同案件、同模板、同目录），则更新该文档
            - 如果不存在，则创建新文档
        """
        from .models import CaseDocument, CaseManagement, CaseFolder
        from django.conf import settings
        from django.db.models import Q
        logger.info(f"[DEBUG] 保存DOCX文档: case_id={case_id}, document_name={document_name}, file_path={file_path}, folder_path={folder_path}, template_id={template_id}, creator={creator}, dept_belong_id={dept_belong_id}, template_print_count={template_print_count}")
        
        case = CaseManagement.objects.get(id=case_id)
        size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        logger.info(f"[DEBUG] 文件信息: 存在={os.path.exists(file_path)}, 大小={size}")
        
        # 转换为相对路径（相对于 MEDIA_ROOT）
        relative_path = file_path
        if file_path.startswith(settings.MEDIA_ROOT):
            # 如果是绝对路径，转换为相对路径
            relative_path = os.path.relpath(file_path, settings.MEDIA_ROOT)
        elif file_path.startswith('media/') or file_path.startswith('media\\'):
            # 如果以 'media/' 开头，去掉这个前缀
            relative_path = file_path[6:]  # 去掉 'media/' 或 'media\'
        
        # 标准化路径分隔符为正斜杠
        relative_path = relative_path.replace('\\', '/')
        logger.info(f"[DEBUG] 相对路径: {relative_path}")
        
        # 获取目录对象
        folder = None
        try:
            folder = CaseFolder.objects.get(
                case=case,
                folder_path=folder_path,
                is_deleted=False
            )
            logger.info(f"[DEBUG] 找到目录: {folder.folder_name} (ID: {folder.id})")
        except CaseFolder.DoesNotExist:
            logger.warning(f"[DEBUG] 目录不存在: {folder_path}，尝试创建")
            # 如果目录不存在，创建案件文书目录
            if folder_path == '/case_documents':
                from .utils.folder_helper import create_case_folders
                create_case_folders(case_id)
                try:
                    folder = CaseFolder.objects.get(
                        case=case,
                        folder_path=folder_path,
                        is_deleted=False
                    )
                    logger.info(f"[DEBUG] 目录创建成功: {folder.folder_name}")
                except:
                    logger.warning(f"[DEBUG] 无法创建目录，folder设置为None")
        
        # 读取渲染后的DOCX文件内容作为预览
        content_preview = ""
        try:
            from docx import Document
            doc = Document(file_path)
            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text.strip())
            content_preview = '\n'.join(content_parts)
            logger.info(f"[DEBUG] 读取DOCX内容成功，段落数: {len(content_parts)}")
        except Exception as e:
            logger.warning(f"读取DOCX内容失败: {e}")
            content_preview = f"文档已生成，文件路径: {file_path}"
        
        # ✅ 简化检查：只根据 case_id 和 template_id 判断
        existing_doc = None
        if template_id:
            existing_doc = CaseDocument.objects.filter(
                case=case,
                template_id=template_id,
                is_deleted=False
            ).first()
            
            if existing_doc:
                logger.info(f"[DEBUG] 找到已存在的文档: id={existing_doc.id}, case_id={case.id}, template_id={template_id}, 将删除旧文件并更新记录")
        
        if existing_doc:
            # ✅ 删除旧文件
            old_file_path = existing_doc.full_file_path
            if old_file_path and os.path.exists(old_file_path):
                try:
                    os.remove(old_file_path)
                    logger.info(f"[DEBUG] 已删除旧文件: {old_file_path}")
                except Exception as e:
                    logger.warning(f"[DEBUG] 删除旧文件失败: {old_file_path}, 错误: {e}")
            
            # ✅ 更新现有文档记录（不修改 print_count，保持原值）
            existing_doc.document_name = document_name
            existing_doc.document_content = content_preview
            existing_doc.file_path = relative_path
            existing_doc.file_size = size
            existing_doc.folder = folder
            existing_doc.folder_path = folder_path
            existing_doc.template_used = document_name
            # ❌ 不修改 print_count，保持文档首次创建时的值
            existing_doc.save()
            logger.info(f"[DEBUG] 文档记录已更新: id={existing_doc.id}, 新文件路径={relative_path}, print_count保持不变={existing_doc.print_count}")
            return existing_doc
        else:
            # 创建新文档
            doc = CaseDocument.objects.create(
                case=case,
                folder=folder,                    # ✅ 设置目录对象
                folder_path=folder_path,          # ✅ 设置目录路径
                document_name=document_name,
                document_type='word',
                document_content=content_preview,
                generation_method='manual',
                template_used=document_name,
                template_id=template_id,          # ✅ 设置模板ID
                creator=creator,                  # ✅ 设置创建人
                dept_belong_id=dept_belong_id,    # ✅ 设置所属部门
                print_count=template_print_count, # ✅ 记录模板配置的默认打印份数
                sort_order=template_sort_order,   # ✅ 记录模板的排序序号
                file_path=relative_path,          # ✅ 保存相对路径
                file_size=size,
            )
            logger.info(f"[DEBUG] 文档记录已创建: id={doc.id}, document_type={doc.document_type}, folder_path={doc.folder_path}, file_path={relative_path}, creator={creator}, dept_belong_id={dept_belong_id}, print_count={template_print_count}, sort_order={template_sort_order}")
            return doc

    def _convert_doc_to_docx_and_render(self, doc_path: str, data: Dict[str, Any], out_path: str) -> str:
        """将DOC文件转换为DOCX并渲染占位符"""
        try:
            # 先读取DOC文件内容
            from .direct_langchain_ai_service import parse_doc_file
            doc_content = parse_doc_file(doc_path)
            
            # 创建新的DOCX文件
            from docx import Document
            doc = Document()
            
            # 将文本内容按行分割并添加到段落中
            lines = doc_content.split('\n')
            for line in lines:
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()  # 空行
            
            # 保存为临时DOCX文件
            temp_docx = out_path.replace('.docx', '_temp.docx')
            doc.save(temp_docx)
            
            # 使用占位符渲染
            self.render_docx(temp_docx, data, out_path)
            
            # 删除临时文件
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
                
            return out_path
            
        except Exception as e:
            logger.error(f"转换DOC到DOCX并渲染失败: {e}")
            # 如果转换失败，尝试直接文本渲染
            with open(doc_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            rendered_content = self.render_text(content, data)
            
            # 创建简单的DOCX文件
            from docx import Document
            doc = Document()
            for line in rendered_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.save(out_path)
            return out_path

    # ==================== 一键：根据模板记录填充并保存 ====================
    def fill_and_save_by_record(self, case_id: int, template_record: "DocumentTemplate", data: Dict[str, Any], request=None) -> "CaseDocument":
        """根据模板记录（数据库中的 DocumentTemplate）进行占位符渲染并保存为文书
        
        Args:
            case_id: 案件ID
            template_record: 模板记录对象
            data: 用于填充的数据字典
            request: 请求对象（用于获取创建人信息）
        """
        filepath = template_record.full_file_path
        logger.info(f"[DEBUG] 开始处理模板: template_id={template_record.id}, template_name={template_record.template_name}, file_path={template_record.file_path}, filepath={filepath}")
        
        # 获取用户信息
        creator = None
        dept_belong_id = None
        template_print_count = getattr(template_record, 'print_count', 1)  # ✅ 获取模板配置的默认打印份数
        template_sort_order = getattr(template_record, 'sort_order', 0)  # ✅ 获取模板的排序序号
        
        if request and hasattr(request, 'user') and str(request.user) != 'AnonymousUser':
            creator = request.user
            dept_belong_id = getattr(request.user, 'dept_id', None)
            logger.info(f"[DEBUG] 从请求获取用户信息: creator={creator}, dept_belong_id={dept_belong_id}, template_print_count={template_print_count}, template_sort_order={template_sort_order}")
        
        if not filepath:
            logger.error(f"[DEBUG] 模板文件路径为空: template_id={template_record.id}, template_name={template_record.template_name}, file_path={template_record.file_path}")
            raise FileNotFoundError(f"模板文件路径为空：模板ID={template_record.id}, 模板名称={template_record.template_name}")
        
        if not os.path.exists(filepath):
            logger.error(f"[DEBUG] 模板文件不存在: filepath={filepath}, template_id={template_record.id}, template_name={template_record.template_name}, file_path={template_record.file_path}")
            # 尝试列出 case_templates 目录下的所有文件，帮助调试
            from django.conf import settings
            case_templates_dir = os.path.join(settings.MEDIA_ROOT, 'case_templates')
            if os.path.exists(case_templates_dir):
                try:
                    files = os.listdir(case_templates_dir)
                    logger.info(f"[DEBUG] case_templates 目录下的文件列表: {files[:10]}")  # 只显示前10个
                except Exception as e:
                    logger.warning(f"[DEBUG] 无法列出 case_templates 目录: {e}")
            raise FileNotFoundError(f"模板文件不存在：{filepath} (模板ID={template_record.id}, 模板名称={template_record.template_name})")

        # 获取模板名称，优先使用template_name，但确保不包含扩展名
        if template_record.template_name:
            # 如果template_name不包含扩展名，直接使用
            if '.' not in template_record.template_name:
                name = template_record.template_name
            else:
                name = os.path.splitext(template_record.template_name)[0]
        else:
            # 从文件路径获取名称
            name = os.path.splitext(os.path.basename(filepath))[0]
        
        ext = os.path.splitext(filepath)[1].lower()
        logger.info(f"[DEBUG] 模板信息: name={name}, ext={ext}")

        # 根据扩展名决定处理方式，保持与模板相同的格式
        # 保存到案件目录下的 case_documents 文件夹
        out_dir = os.path.join('media', 'cases', str(case_id), 'case_documents')
        os.makedirs(out_dir, exist_ok=True)
        
        if ext == '.docx':
            # DOCX模板，生成DOCX文档
            out_name = f"{name}.docx"
            out_path = os.path.join(out_dir, out_name)
            logger.info(f"[DEBUG] DOCX处理: out_name={out_name}, out_path={out_path}")
            logger.info(f"开始渲染DOCX模板: {filepath} -> {out_path}")
            try:
                self.render_docx(filepath, data, out_path)
                logger.info(f"DOCX渲染成功: {out_path}")
                
                # 验证生成的文件
                if os.path.exists(out_path):
                    file_size = os.path.getsize(out_path)
                    actual_ext = os.path.splitext(out_path)[1]
                    logger.info(f"[DEBUG] 生成的文件验证: 存在={True}, 大小={file_size}, 扩展名={actual_ext}")
                else:
                    logger.error(f"[DEBUG] 生成的文件不存在: {out_path}")
                
                # ✅ 传递 template_id、creator、dept_belong_id、template_print_count、template_sort_order 参数
                return self.save_docx_document(
                    case_id, 
                    out_name, 
                    out_path, 
                    template_id=template_record.id,
                    creator=creator,
                    dept_belong_id=dept_belong_id,
                    template_print_count=template_print_count,
                    template_sort_order=template_sort_order
                )
            except Exception as e:
                logger.error(f"[DEBUG] DOCX渲染异常: {e}")
                logger.error(f"DOCX渲染失败: {e}")
                raise
        elif ext == '.doc':
            # DOC模板，先转换为DOCX再渲染
            out_name = f"{name}.docx"
            out_path = os.path.join(out_dir, out_name)
            # 对于.doc文件，先读取内容，然后创建新的DOCX文件
            self._convert_doc_to_docx_and_render(filepath, data, out_path)
            # ✅ 传递 template_id、creator、dept_belong_id、template_print_count、template_sort_order 参数
            return self.save_docx_document(
                case_id, 
                out_name, 
                out_path, 
                template_id=template_record.id,
                creator=creator,
                dept_belong_id=dept_belong_id,
                template_print_count=template_print_count,
                template_sort_order=template_sort_order
            )
        else:
            # 其他文本类（.txt/.md/.html），保持原格式
            out_name = f"{name}{ext}"
            out_path = os.path.join(out_dir, out_name)
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                tpl = f.read()
            content = self.render_text(tpl, data)
            with open(out_path, 'w', encoding='utf-8') as f:
                f.write(content)
            # ✅ 传递 template_id、creator、dept_belong_id、template_print_count、template_sort_order 参数
            return self.save_text_document(
                case_id, 
                out_name, 
                content, 
                ext, 
                template_id=template_record.id,
                creator=creator,
                dept_belong_id=dept_belong_id,
                template_print_count=template_print_count,
                template_sort_order=template_sort_order
            )

    # ==================== 将现有模板转换为占位符模板（文本） ====================
    def convert_text_template_to_placeholders(self, template_text: str, mapping: Dict[str, str]) -> str:
        """将已有文本模板中的样例值批量替换为占位符

        参数：
        - template_text: 原始模板内容
        - mapping: { '样例值或正则': '占位符键名' }，正则以 r"^...$" 样式传入时使用正则匹配
        """
        result = template_text
        for sample, key in mapping.items():
            # 优先当作精确子串替换，否则尝试正则
            placeholder = f"{{{{ {key} }}}}"
            if sample.startswith('re:'):
                # 正则：re:后面是表达式
                pattern = sample[3:]
                try:
                    result = re.sub(pattern, placeholder, result)
                except re.error:
                    logger.warning(f"无效正则：{pattern}")
            else:
                if sample in result:
                    result = result.replace(sample, placeholder)
        return result

    # ==================== 将现有 DOCX 模板转换为占位符模板 ====================
    def convert_docx_template_to_placeholders(self, template_path: str, out_path: str, mapping: Dict[str, str]) -> str:
        """将 DOCX 模板中的样例值批量替换为占位符，并保存为新文件

        - mapping: { '样例值': '占位符键名' }，若需正则请在上层先转换文本后写回
        返回：输出文件路径
        """
        from docx import Document  # type: ignore
        doc = Document(template_path)

        def _replace_in_text(text: str) -> str:
            if not text:
                return text
            for sample, key in mapping.items():
                placeholder = f"{{{{ {key} }}}}"
                if sample in text:
                    text = text.replace(sample, placeholder)
            return text

        for p in doc.paragraphs:
            for r in p.runs:
                r.text = _replace_in_text(r.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.text = _replace_in_text(r.text)

        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        doc.save(out_path)
        return out_path


# 创建一个可复用的全局实例
placeholder_service = PlaceholderTemplateService()


