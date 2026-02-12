"""
法规文档解析工具
解析Word文档(.docx)格式的法规文件，提取章节、条款等结构化信息
"""

import os
import re
from typing import Dict, List, Any
from docx import Document
import logging

logger = logging.getLogger(__name__)


class RegulationDocumentParser:
    """法规文档解析器"""
    
    def __init__(self, file_path: str):
        """
        初始化解析器
        
        Args:
            file_path: Word文档路径
        """
        self.file_path = file_path
        self.document = None
        
    def parse(self) -> Dict[str, Any]:
        """
        解析文档
        
        Returns:
            包含法规结构化信息的字典
        """
        try:
            self.document = Document(self.file_path)
            
            # 提取文档内容
            title = self._extract_title()
            metadata = self._extract_metadata()
            chapters = self._extract_chapters()
            
            return {
                'success': True,
                'title': title,
                'metadata': metadata,
                'chapters': chapters,
                'total_articles': self._count_articles(chapters)
            }
            
        except Exception as e:
            logger.error(f"解析文档失败: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_title(self) -> str:
        """提取文档标题"""
        if not self.document or len(self.document.paragraphs) == 0:
            return "未知标题"
        
        # 通常标题在第一段
        for para in self.document.paragraphs[:5]:
            text = para.text.strip()
            if text and ('中华人民共和国' in text or '法' in text):
                return text
        
        return self.document.paragraphs[0].text.strip()
    
    def _extract_metadata(self) -> Dict[str, Any]:
        """提取文档元数据"""
        metadata = {
            'publish_date': '',
            'effective_date': '',
            'authority': '',
            'document_number': ''
        }
        
        # 从文档前几段提取元数据
        for para in self.document.paragraphs[:20]:
            text = para.text.strip()
            
            # 匹配日期
            date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', text)
            if date_match:
                year, month, day = date_match.groups()
                date_str = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                if not metadata['publish_date']:
                    metadata['publish_date'] = date_str
                elif not metadata['effective_date'] and '施行' in text:
                    metadata['effective_date'] = date_str
            
            # 匹配发文机关
            if '全国人' in text or '人民代表大会' in text:
                metadata['authority'] = '全国人民代表大会' if '全国人民代表大会' in text else '全国人大常委会'
            
            # 匹配发文字号
            if '主席令' in text or '第' in text and '号' in text:
                metadata['document_number'] = text
        
        return metadata
    
    def _extract_chapters(self) -> List[Dict[str, Any]]:
        """提取章节和条款"""
        chapters = []
        current_chapter = None
        current_section = None
        
        # 章节模式（如：第一编、第二编）
        chapter_pattern = re.compile(r'^第[一二三四五六七八九十零百]+编\s+(.+)$')
        # 章模式（如：第一章、第二章）
        section_pattern = re.compile(r'^第[一二三四五六七八九十零百]+章\s+(.+)$')
        # 特殊章节模式（如：附则、总则、说明等）
        special_section_pattern = re.compile(r'^(附\s*则|总\s*则|说\s*明|补充规定)$')
        # 条款模式（如：第一条、第二条）- 添加"零"字支持（如第一百零六条）
        article_pattern = re.compile(r'^第[一二三四五六七八九十零百千]+条\s+(.+)$')
        
        # 添加计数器用于调试
        para_count = 0
        chapter_detected = False
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            para_count += 1
            if not text:
                continue
            
            # 如果包含"附"，记录下来
            if '附' in text:
                logger.info(f"第{para_count}段包含'附': {text[:50]}")
                chapter_detected = True
            
            # 匹配编
            chapter_match = chapter_pattern.match(text)
            if chapter_match:
                if current_chapter:
                    chapters.append(current_chapter)
                current_chapter = {
                    'name': text,
                    'sections': []
                }
                current_section = None
                continue
            
            # 匹配特殊章节（附则等）
            special_match = special_section_pattern.match(text)
            if special_match:
                if not current_chapter:
                    current_chapter = {
                        'name': '其他规定',
                        'sections': []
                    }
                
                current_section = {
                    'id': f"section-{len(current_chapter['sections'])}",
                    'name': text,
                    'articles': []
                }
                current_chapter['sections'].append(current_section)
                logger.info(f"识别到特殊章节: {text}, section_id={current_section['id']}")
                continue
            
            # 匹配章
            section_match = section_pattern.match(text)
            if section_match:
                if not current_chapter:
                    current_chapter = {
                        'name': '总则',
                        'sections': []
                    }
                
                # 提取章节名称（去掉"第X章"前缀）
                section_name = section_match.group(1).strip() if section_match.groups() else text
                
                current_section = {
                    'id': f"section-{len(current_chapter['sections'])}",
                    'name': section_name,  # 使用提取的章节名称
                    'articles': []
                }
                current_chapter['sections'].append(current_section)
                logger.info(f"第{para_count}段 - 识别到章节: {text} -> 名称: {section_name}, 当前section: {id(current_section)}")
                
                # 如果是附则，记录当前状态
                if '附' in section_name:
                    logger.info(f"=== 附则章节已创建，section_id={id(current_section)}, 当前articles数: {len(current_section['articles'])} ===")
                continue
            
            # 匹配条
            article_match = article_pattern.match(text)
            if article_match:
                if not current_chapter:
                    current_chapter = {
                        'name': '总则',
                        'sections': []
                    }
                if not current_section:
                    current_section = {
                        'id': f"section-{len(current_chapter['sections'])}",
                        'name': '一般规定',
                        'articles': []
                    }
                    current_chapter['sections'].append(current_section)
                    logger.info(f"创建默认section: {current_section['name']}")
                
                # 提取条款号和内容
                article_number = text.split()[0] if ' ' in text else text[:text.index('条')+1]
                article_content = text[len(article_number):].strip()
                
                current_section['articles'].append({
                    'number': article_number,
                    'content': article_content
                })
                
                # 记录特殊章节的条款（用于调试附则等问题）
                section_name_clean = current_section['name'].replace(' ', '').replace('　', '')  # 去除所有空格
                if section_name_clean in ['附则', '总则', '说明', '补充规定']:
                    logger.info(f"第{para_count}段 - 添加条款到特殊章节 '{current_section['name']}' (section_id={id(current_section)}): {article_number}, 当前该节共{len(current_section['articles'])}条")
        
        # 添加最后一个章节
        if current_chapter:
            chapters.append(current_chapter)
        
        # 打印解析摘要（用于调试）
        logger.info(f"解析完成，共 {len(chapters)} 个章/编:")
        for chapter in chapters:
            logger.info(f"  {chapter['name']}: {len(chapter.get('sections', []))} 个节")
            for section in chapter.get('sections', []):
                article_count = len(section.get('articles', []))
                logger.info(f"    - {section['name']}: {article_count} 条")
        
        return chapters
    
    def _count_articles(self, chapters: List[Dict[str, Any]]) -> int:
        """统计条款总数"""
        total = 0
        for chapter in chapters:
            for section in chapter.get('sections', []):
                total += len(section.get('articles', []))
        return total


def parse_regulation_document(file_path: str) -> Dict[str, Any]:
    """
    解析法规文档
    
    Args:
        file_path: 文档路径
        
    Returns:
        解析结果
    """
    parser = RegulationDocumentParser(file_path)
    return parser.parse()

