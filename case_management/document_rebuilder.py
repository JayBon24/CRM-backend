"""
文档重建器模块 - 将填充后的内容重新生成为Word文档
"""

import os
import re
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xml.etree.ElementTree as ET
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentRebuilder:
    """文档重建器 - 将填充后的内容重新生成为Word文档"""
    
    def __init__(self, template_path: str):
        """
        初始化文档重建器
        
        Args:
            template_path: 原始模板文件路径
        """
        self.template_path = template_path
        self.doc = Document(template_path)
    
    def rebuild_from_markdown(self, filled_markdown: str, output_path: str) -> bool:
        """
        从填充后的Markdown重建文档
        
        Args:
            filled_markdown: 填充后的Markdown内容
            output_path: 输出文件路径
            
        Returns:
            是否重建成功
        """
        try:
            # 创建新文档
            new_doc = Document()
            
            # 解析Markdown内容
            lines = filled_markdown.split('\n')
            current_table = None
            table_rows = []
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    if current_table:
                        # 处理表格
                        self._add_table_to_doc(new_doc, table_rows)
                        current_table = None
                        table_rows = []
                    continue
                
                # 处理标题
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    new_doc.add_heading(title, level)
                
                # 处理表格
                elif line.startswith('|') and '|' in line[1:]:
                    if not current_table:
                        current_table = True
                        table_rows = []
                    
                    # 解析表格行
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if cells and not all(cell == '---' for cell in cells):
                        table_rows.append(cells)
                
                # 处理普通段落
                else:
                    if current_table:
                        # 先处理表格
                        self._add_table_to_doc(new_doc, table_rows)
                        current_table = None
                        table_rows = []
                    
                    # 添加段落
                    para = new_doc.add_paragraph()
                    self._format_paragraph_text(para, line)
            
            # 处理最后的表格
            if current_table:
                self._add_table_to_doc(new_doc, table_rows)
            
            # 保存文档
            new_doc.save(output_path)
            logger.info(f"文档已重建: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"从Markdown重建文档失败: {e}")
            return False
    
    def rebuild_from_structured_markdown(self, filled_markdown: str, output_path: str) -> bool:
        """
        从填充后的结构化Markdown重建文档（保留格式信息）
        
        Args:
            filled_markdown: 填充后的结构化Markdown内容
            output_path: 输出文件路径
            
        Returns:
            是否重建成功
        """
        try:
            # 创建新文档
            new_doc = Document()
            
            # 按行分割内容
            lines = filled_markdown.split('\n')
            
            for line in lines:
                if not line.strip():
                    # 空行
                    new_doc.add_paragraph()
                    continue
                
                # 提取格式标记
                format_info = self._extract_format_info(line)
                
                # 处理标题
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('# ').strip()
                    # 清理格式标签
                    title_text = self._clean_format_tags(title_text)
                    
                    heading = new_doc.add_heading(title_text, level=min(level, 6))
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 设置标题字体
                    for run in heading.runs:
                        run.font.name = '黑体'
                        if 'size' in format_info:
                            run.font.size = Pt(format_info['size'])
                        else:
                            run.font.size = Pt(16 - level * 2)
                        run.bold = True
                    continue
                
                # 处理表格行
                if line.startswith('|') and line.endswith('|'):
                    # 跳过表格分隔行
                    if '---' in line:
                        continue
                    # 处理表格数据行
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if cells and not any('列' in cell for cell in cells):
                        # 创建表格（如果还没有）
                        if not hasattr(new_doc, '_current_table') or new_doc._current_table is None:
                            new_doc._current_table = new_doc.add_table(rows=1, cols=len(cells))
                            new_doc._current_table.style = 'Table Grid'
                        else:
                            # 添加新行
                            row = new_doc._current_table.add_row()
                            for i, cell_content in enumerate(cells):
                                if i < len(row.cells):
                                    # 清理格式标签
                                    clean_content = self._clean_format_tags(cell_content)
                                    row.cells[i].text = clean_content
                                    
                                    # 设置表格单元格字体
                                    for paragraph in row.cells[i].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = '宋体'
                                            run.font.size = Pt(12)
                    continue
                
                # 处理列表项
                if line.strip().startswith('- '):
                    list_text = line.strip()[2:].strip()
                    # 清理格式标签
                    list_text = self._clean_format_tags(list_text)
                    
                    para = new_doc.add_paragraph(list_text, style='List Bullet')
                    for run in para.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)
                    continue
                
                # 处理普通段落
                para = new_doc.add_paragraph()
                
                # 应用段落格式
                self._apply_paragraph_format(para, format_info)
                
                # 处理缩进
                indent_level = 0
                while line.startswith('  '):
                    line = line[2:]
                    indent_level += 1
                
                if indent_level > 0:
                    para.paragraph_format.left_indent = Inches(0.5 * indent_level)
                
                # 处理格式标签
                self._process_formatted_text(para, line, format_info)
            
            # 设置页面边距和字体
            self._set_document_styles(new_doc)
            
            # 保存文档
            new_doc.save(output_path)
            logger.info(f"文档已重建: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"从结构化Markdown重建文档失败: {e}")
            return False
    
    def rebuild_from_xml(self, filled_xml: str, output_path: str) -> bool:
        """
        从填充后的XML重建文档
        
        Args:
            filled_xml: 填充后的XML内容
            output_path: 输出文件路径
            
        Returns:
            是否重建成功
        """
        try:
            root = ET.fromstring(filled_xml)
            new_doc = Document()
            
            for element in root:
                if element.tag == 'paragraph':
                    text = element.text.strip() if element.text else ""
                    style = element.get('style', 'Normal')
                    
                    para = new_doc.add_paragraph()
                    para.text = text
                    
                    # 应用样式
                    if 'Heading' in style:
                        level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
                        para.style = f'Heading {level}'
                    elif 'Title' in style:
                        para.style = 'Title'
                    elif 'Subtitle' in style:
                        para.style = 'Subtitle'
                
                elif element.tag == 'table':
                    self._add_xml_table_to_doc(new_doc, element)
            
            # 设置文档样式
            self._set_document_styles(new_doc)
            
            new_doc.save(output_path)
            logger.info(f"文档已重建: {output_path}")
            return True
            
        except ET.ParseError as e:
            logger.error(f"XML解析错误: {e}")
            return False
        except Exception as e:
            logger.error(f"从XML重建文档失败: {e}")
            return False
    
    def _extract_format_info(self, line: str) -> Dict[str, Any]:
        """
        从行中提取格式信息
        
        Args:
            line: 文本行
            
        Returns:
            格式信息字典
        """
        format_info = {}
        
        # 字体大小
        size_match = re.search(r'<size:(\d+)>', line)
        if size_match:
            format_info['size'] = int(size_match.group(1))
        
        # 字体名称
        font_match = re.search(r'<font:([^>]+)>', line)
        if font_match:
            format_info['font'] = font_match.group(1)
        
        # 段落缩进
        indent_match = re.search(r'<indent:(\d+)>', line)
        if indent_match:
            format_info['indent'] = int(indent_match.group(1))
        
        # 段落间距
        space_before_match = re.search(r'<space_before:(\d+)>', line)
        if space_before_match:
            format_info['space_before'] = int(space_before_match.group(1))
        
        space_after_match = re.search(r'<space_after:(\d+)>', line)
        if space_after_match:
            format_info['space_after'] = int(space_after_match.group(1))
        
        # 行距
        line_spacing_match = re.search(r'<line_spacing:([\d.]+)>', line)
        if line_spacing_match:
            format_info['line_spacing'] = float(line_spacing_match.group(1))
        
        # 对齐方式
        align_match = re.search(r'<align:(\w+)>', line)
        if align_match:
            format_info['align'] = align_match.group(1)
        
        return format_info
    
    def _clean_format_tags(self, text: str) -> str:
        """
        清理格式标签
        
        Args:
            text: 包含格式标签的文本
            
        Returns:
            清理后的文本
        """
        # 移除格式标记
        text = re.sub(r'<size:\d+>', '', text)
        text = re.sub(r'<font:[^>]+>', '', text)
        text = re.sub(r'<indent:\d+>', '', text)
        text = re.sub(r'<space_before:\d+>', '', text)
        text = re.sub(r'<space_after:\d+>', '', text)
        text = re.sub(r'<line_spacing:[\d.]+>', '', text)
        text = re.sub(r'<align:\w+>', '', text)
        
        # 清理加粗、斜体、下划线标记
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'<u>(.*?)</u>', r'\1', text)
        
        return text.strip()
    
    def _apply_paragraph_format(self, para, format_info: Dict[str, Any]):
        """
        应用段落格式
        
        Args:
            para: 段落对象
            format_info: 格式信息
        """
        if 'indent' in format_info:
            # 限制缩进值在合理范围内
            indent_value = min(max(format_info['indent'], 0), 1000)
            para.paragraph_format.first_line_indent = Inches(indent_value / 72)
        
        if 'space_before' in format_info:
            space_before_value = min(max(format_info['space_before'], 0), 1000)
            para.paragraph_format.space_before = Pt(space_before_value)
        
        if 'space_after' in format_info:
            space_after_value = min(max(format_info['space_after'], 0), 1000)
            para.paragraph_format.space_after = Pt(space_after_value)
        
        if 'line_spacing' in format_info:
            line_spacing_value = min(max(format_info['line_spacing'], 0.5), 3.0)
            para.paragraph_format.line_spacing = line_spacing_value
        
        if 'align' in format_info:
            if format_info['align'] == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif format_info['align'] == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif format_info['align'] == 'justify':
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _process_formatted_text(self, para, line: str, format_info: Dict[str, Any]):
        """
        处理格式化文本
        
        Args:
            para: 段落对象
            line: 文本行
            format_info: 格式信息
        """
        # 先按格式标签分割文本
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|<u>.*?</u>)', line)
        
        for part in parts:
            if not part:
                continue
            
            run = para.add_run(part)
            
            # 设置字体
            if 'font' in format_info:
                run.font.name = format_info['font']
            else:
                run.font.name = '仿宋'
            
            # 设置字体大小
            if 'size' in format_info:
                run.font.size = Pt(format_info['size'])
            else:
                run.font.size = Pt(18)
            
            # 应用格式
            if part.startswith('**') and part.endswith('**'):
                run.bold = True
                run.text = part[2:-2]
            elif part.startswith('*') and part.endswith('*'):
                run.italic = True
                run.text = part[1:-1]
            elif part.startswith('<u>') and part.endswith('</u>'):
                run.underline = True
                run.text = part[3:-4]
    
    def _add_table_to_doc(self, doc: Document, table_rows: List[List[str]]):
        """
        添加表格到文档
        
        Args:
            doc: 文档对象
            table_rows: 表格行数据
        """
        if not table_rows:
            return
        
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        
        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_data
    
    def _add_xml_table_to_doc(self, doc: Document, table_element):
        """
        从XML添加表格到文档
        
        Args:
            doc: 文档对象
            table_element: XML表格元素
        """
        rows = table_element.findall('row')
        if not rows:
            return
        
        # 确定列数
        max_cols = max(len(row.findall('cell')) for row in rows)
        
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for row_idx, row in enumerate(rows):
            cells = row.findall('cell')
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    table.rows[row_idx].cells[col_idx].text = cell.text or ""
    
    def _set_document_styles(self, doc: Document):
        """
        设置文档样式
        
        Args:
            doc: 文档对象
        """
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)

def rebuild_document_from_markdown(template_path: str, filled_markdown: str, output_path: str) -> bool:
    """
    从Markdown重建文档的便捷函数
    
    Args:
        template_path: 模板文件路径
        filled_markdown: 填充后的Markdown内容
        output_path: 输出文件路径
        
    Returns:
        是否重建成功
    """
    try:
        rebuilder = DocumentRebuilder(template_path)
        return rebuilder.rebuild_from_structured_markdown(filled_markdown, output_path)
    except Exception as e:
        logger.error(f"重建文档失败: {e}")
        return False

def rebuild_document_from_xml(template_path: str, filled_xml: str, output_path: str) -> bool:
    """
    从XML重建文档的便捷函数
    
    Args:
        template_path: 模板文件路径
        filled_xml: 填充后的XML内容
        output_path: 输出文件路径
        
    Returns:
        是否重建成功
    """
    try:
        rebuilder = DocumentRebuilder(template_path)
        return rebuilder.rebuild_from_xml(filled_xml, output_path)
    except Exception as e:
        logger.error(f"重建文档失败: {e}")
        return False

if __name__ == "__main__":
    # 测试文档重建器
    template_path = "backend/template/起诉状.docx"
    if os.path.exists(template_path):
        test_markdown = """
# 起诉状

**原告**：张三
**被告**：李四

## 诉讼请求
1. 要求被告支付合同款项100000元
2. 要求被告承担诉讼费用

## 事实与理由
双方于2023年1月签订买卖合同，被告未按约定履行义务...
"""
        
        rebuilder = DocumentRebuilder(template_path)
        success = rebuilder.rebuild_from_structured_markdown(test_markdown, "test_output.docx")
        
        if success:
            print("文档重建成功")
        else:
            print("文档重建失败")
    else:
        print(f"模板文件不存在: {template_path}")
